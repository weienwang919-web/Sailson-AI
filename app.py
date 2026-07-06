import os
import sys
import datetime
import time
import logging
import re
import smtplib
import json
import atexit
import subprocess
from collections import Counter, defaultdict
from email.mime.text import MIMEText
from urllib.parse import urlencode
import pandas as pd
import uuid
import threading
import requests
from io import BytesIO
from PIL import Image
from flask import Flask, render_template, request, jsonify, redirect, url_for, session, send_file, Response
from flask_bcrypt import Bcrypt
from werkzeug.exceptions import HTTPException
from apify_client import ApifyClient
from dotenv import load_dotenv
from openai import OpenAI
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill
from openpyxl.utils.dataframe import dataframe_to_rows
from functools import wraps
import html
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
import bleach
import database as db
import etl_tools
import etl_jobs
import video_metrics_etl
import profile_video_scheduler
import sentiment_insight
import agent_service
import competitor_radar
import tiktok_official_service
import usage_service
from apscheduler.schedulers.background import BackgroundScheduler
from apscheduler.jobstores.sqlalchemy import SQLAlchemyJobStore
from apscheduler.executors.pool import ThreadPoolExecutor

# XSS 防护：报告 HTML 允许的标签与属性
ALLOWED_TAGS = {
    'div', 'span', 'p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6',
    'table', 'thead', 'tbody', 'tr', 'th', 'td',
    'ul', 'ol', 'li', 'br', 'strong', 'em', 'b', 'i',
    'a', 'img', 'blockquote', 'code', 'pre'
}
ALLOWED_ATTRS = {
    '*': ['class', 'style', 'id'],
    'a': ['href', 'title', 'target'],
    'img': ['src', 'alt', 'width', 'height'],
    'td': ['colspan', 'rowspan'],
    'th': ['colspan', 'rowspan'],
}
ALLOWED_PROTOCOLS = ['http', 'https', 'data', 'mailto']


def sanitize_html(html_content):
    """对返回前端的 HTML 做 XSS 清洗"""
    if not html_content:
        return ''
    return bleach.clean(html_content, tags=ALLOWED_TAGS, attributes=ALLOWED_ATTRS, protocols=ALLOWED_PROTOCOLS)
import rag
import tasks
from video_vision import get_video_vision_section, analyze_all_videos_for_export, _build_vision_html_from_results

# 加载 .env 文件
load_dotenv()

# ============================================
# 日志配置 - 确保输出到 stdout
# ============================================
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(message)s',
    handlers=[logging.StreamHandler(sys.stdout)]
)
logger = logging.getLogger(__name__)

# ============================================
# 环境配置 - 优化版
# ============================================

# 清除代理设置（云端环境不需要代理）
for proxy_var in ['HTTP_PROXY', 'HTTPS_PROXY', 'http_proxy', 'https_proxy']:
    if proxy_var in os.environ:
        del os.environ[proxy_var]
        logger.info(f"🧹 已清除代理设置: {proxy_var}")

# 加载环境变量
DASHSCOPE_API_KEY = os.environ.get('DASHSCOPE_API_KEY')
APIFY_TOKEN = os.environ.get('APIFY_TOKEN')
PORT = int(os.environ.get('PORT', 5001))

# 长任务处理模式配置（预留开关，默认保持现状：由 Web 线程执行）
USE_DB_WORKER = os.environ.get('USE_DB_WORKER', 'false').lower() == 'true'

# 反馈邮件配置（可选）
SMTP_HOST = os.environ.get('SMTP_HOST')
SMTP_PORT = int(os.environ.get('SMTP_PORT', '587'))
SMTP_USER = os.environ.get('SMTP_USER')
SMTP_PASS = os.environ.get('SMTP_PASS')
FEEDBACK_EMAIL_TO = os.environ.get('FEEDBACK_EMAIL_TO')
FEEDBACK_EMAIL_FROM = os.environ.get('FEEDBACK_EMAIL_FROM', SMTP_USER or FEEDBACK_EMAIL_TO or '')

# 启动时输出配置状态
logger.info("=" * 60)
logger.info("🚀 Sailson AI 工作台启动中...")
logger.info(f"🔑 DASHSCOPE_API_KEY: {'✅ 已配置' if DASHSCOPE_API_KEY else '❌ 未配置'}")
logger.info(f"🔑 APIFY_TOKEN: {'✅ 已配置' if APIFY_TOKEN else '❌ 未配置'}")
logger.info(f"🌐 PORT: {PORT}")
logger.info(f"🧵 Long-task mode: {'DB worker' if USE_DB_WORKER else 'in-process threads'}")
logger.info(f"🐍 Python 版本: {sys.version}")
logger.info("=" * 60)

# 初始化 AI 引擎
if DASHSCOPE_API_KEY:
    try:
        qwen_client = OpenAI(
            api_key=DASHSCOPE_API_KEY,
            base_url="https://dashscope.aliyuncs.com/compatible-mode/v1"
        )
        logger.info("✅ 通义千问 API 初始化成功")
    except Exception as e:
        logger.error(f"❌ 通义千问 API 初始化失败: {e}")
        qwen_client = None
else:
    logger.warning("⚠️ 警告: DASHSCOPE_API_KEY 未配置，AI 功能将不可用")
    qwen_client = None

# 不再初始化全局 Apify 客户端，改用 REST API
# 只检查 token 是否存在
if APIFY_TOKEN:
    logger.info("✅ APIFY_TOKEN 已配置")
else:
    logger.warning("⚠️ 警告: APIFY_TOKEN 未配置，爬虫功能将不可用")

# Flask 应用初始化
app = Flask(__name__)

secret_key = os.environ.get('SECRET_KEY')
if not secret_key:
    # 为了不影响现有功能，在缺少 SECRET_KEY 时自动生成一次性开发密钥
    # 生产环境必须通过环境变量显式配置 SECRET_KEY
    logger.warning("⚠️ SECRET_KEY 未配置，将使用临时开发密钥。请在生产环境中通过环境变量设置 SECRET_KEY！")
    import secrets
    secret_key = "dev-" + secrets.token_hex(32)

app.secret_key = secret_key
bcrypt = Bcrypt(app)
KOL_STATIC_DIR = os.path.join(app.root_path, 'static', 'kol')
KOL_API_BASE_URL = os.environ.get('KOL_API_BASE_URL', '')
if not KOL_API_BASE_URL and os.environ.get('RENDER'):
    KOL_API_BASE_URL = 'http://127.0.0.1:8001'
KOL_API_BASE_URL = KOL_API_BASE_URL.rstrip('/')
KOL_PROXY_TOKEN = os.environ.get('KOL_PROXY_TOKEN', '')

# ============================================
# APScheduler 初始化
# Worker 进程不需要启动定时调度器
# ============================================
_IS_WORKER = os.environ.get('_IS_WORKER', 'false').lower() == 'true'


def start_kol_api_if_needed():
    """Render Dashboard 若仍使用旧 start command，也能随 Flask 自动拉起 KOL API。"""
    if _IS_WORKER:
        return None
    if not os.environ.get('RENDER'):
        return None
    if os.environ.get('KOL_AUTOSTART', 'true').lower() != 'true':
        return None
    if not KOL_API_BASE_URL.startswith(('http://127.0.0.1', 'http://localhost')):
        return None

    kol_backend_dir = os.path.join(app.root_path, 'kol_web', 'backend')
    if not os.path.isdir(kol_backend_dir):
        logger.error(f"KOL 后端目录不存在: {kol_backend_dir}")
        return None

    kol_port = os.environ.get('KOL_PORT', '8001')
    logger.info(f"🚀 自动启动 KOL FastAPI: 127.0.0.1:{kol_port}")
    process = subprocess.Popen(
        [
            sys.executable,
            '-m',
            'uvicorn',
            'app.main:app',
            '--host',
            '127.0.0.1',
            '--port',
            kol_port
        ],
        cwd=kol_backend_dir,
        env=os.environ.copy()
    )
    atexit.register(process.terminate)
    return process


_kol_api_process = start_kol_api_if_needed()


def run_scheduled_profile_video_sync():
    """APScheduler 可序列化入口：主页视频基础数据定时同步。"""
    enqueue_due_profile_video_sync()


def run_scheduled_feishu_profile_video_sync():
    """APScheduler 可序列化入口：从飞书配置表读取主页并同步视频四表。"""
    enqueue_due_feishu_profile_video_sync()


if not _IS_WORKER:
    jobstores = {
        'default': SQLAlchemyJobStore(url='sqlite:///jobs.sqlite')
    }
    executors = {
        'default': ThreadPoolExecutor(5)
    }
    job_defaults = {
        'coalesce': False,
        'max_instances': 1
    }
    scheduler = BackgroundScheduler(
        jobstores=jobstores,
        executors=executors,
        job_defaults=job_defaults,
        timezone='Asia/Shanghai'
    )

    # 注册定时任务
    # FB 评论抓取：每 6 小时执行一次（0, 6, 12, 18 点）
    scheduler.add_job(
        func=tasks.scrape_fb_comments,
        trigger='cron',
        hour='0,6,12,18',
        minute=0,
        id='fb_scrape_job',
        replace_existing=True
    )

    # TikTok 热点刷新：每天凌晨 2 点执行
    scheduler.add_job(
        func=tasks.refresh_tiktok_hotspots,
        trigger='cron',
        hour=2,
        minute=0,
        id='tiktok_hotspot_job',
        replace_existing=True
    )

    if profile_video_scheduler.profile_video_sync_hard_disabled():
        scheduler.start(paused=True)
        for stale_job_id in ('profile_video_sync_job', 'feishu_profile_video_sync_job'):
            try:
                scheduler.remove_job(stale_job_id)
            except Exception:
                pass
        scheduler.resume()
        logger.warning("⛔ 主页视频定时同步已硬禁用，未注册 profile_video_sync / feishu_profile_video_sync 定时任务")
    else:
        # 主页视频基础数据定时同步：每小时检查一次当前小时配置
        scheduler.add_job(
            func=run_scheduled_profile_video_sync,
            trigger='cron',
            minute=5,
            id='profile_video_sync_job',
            replace_existing=True
        )

        # 飞书四表主页视频自动化：每小时检查一次配置表中的抓取小时
        scheduler.add_job(
            func=run_scheduled_feishu_profile_video_sync,
            trigger='cron',
            minute=12,
            id='feishu_profile_video_sync_job',
            replace_existing=True
        )

        scheduler.start()
        logger.info("✅ APScheduler 已启动，定时任务已注册")
else:
    scheduler = None
    logger.info("⏭️ Worker 模式，跳过 APScheduler 初始化")


# 内存存储（保留用于向后兼容）
HISTORY_DB = []
LATEST_ANALYSIS_RESULTS = {}  # 存储最新的分析结果，用于导出
# TASK_QUEUE 已迁移到数据库，不再使用内存字典

# task_queue 表结构状态（用于向后兼容老数据库）
TASK_QUEUE_HAS_FUNCTION_TYPE = True
ANALYSIS_RESULTS_HAS_JSON = True

# 项目与提示词外置（多项目接入）
VALID_PROJECTS = ('CFL', 'PUBGM', 'HOK')
_PROMPTS_CACHE = None

def load_prompts():
    """加载 config/prompts.json；缺失时回退到 CFL 硬编码，保证现有行为不变。"""
    global _PROMPTS_CACHE
    if _PROMPTS_CACHE is not None:
        return _PROMPTS_CACHE
    fallback = {
        'sentiment': {
            'CFL': (
                "Analyze these comments and categorize them. Output ONLY a JSON array.\n\n"
                "Comments:\n{batch_content}\n\n"
                "Categories (Chinese only):\n1. 外挂作弊 - hackers, cheating\n2. 游戏优化 - lag, crashes\n"
                "3. 游戏Bug - glitches, errors\n4. 充值退款 - payment issues\n"
                "5. 新模式/地图/平衡性建议 - new content requests\n6. 其他 - spam, praise\n\n"
                "Output format (JSON array only, no markdown):\n[\n  {{\n    \"text\": \"comment text\",\n    \"category\": \"外挂作弊\",\n"
                "    \"sentiment\": \"负面\",\n    \"language\": \"英语\",\n    \"analysis\": \"详细分析内容\"\n  }},\n  ...\n]\n\n"
                "IMPORTANT:\n- Output ONLY valid JSON array\n"
                "- Output exactly one object per comment; do NOT skip any comment (use category \"其他\" for spam/praise if needed)\n"
                "- Use Chinese for category, sentiment, language, and analysis\n"
                "- Language options (MUST be one of these): 英语, 菲律宾语, 泰语, 越南语, 印尼语, 马来语\n"
                "- Identify the language accurately based on the text\n"
                "- 本报告供 CFL 品牌方/客户查看，简要分析需便于快速把握玩家诉求与情绪。\n"
                "- Analysis（简要分析）字数要求，必须严格执行：\n"
                "  * 短评论（原文 < 30 字）：一句话概括，15-20 个中文字。\n"
                "  * 中等评论（30-80 字）：两至三句话，包含问题点与情绪，50-70 个中文字。\n"
                "  * 长评论（≥ 80 字）：展开分析，包含主要诉求、玩家情绪、关键细节及对品牌的参考点，80-120 个中文字，不得仅用一句话总结。\n"
                "  * 内容须包含：主要问题、玩家情绪、关键细节；长评论务必多句展开，不可全部统一为一句话总结。\n"
            ),
            'PUBGM': '',
            'HOK': '',
        },
        'competitor': {
            'CFL': (
                "You are a Data Entry Assistant. Please fill the following TikTok data into the PROVIDED HTML TEMPLATE.\n\n"
                "【Data Source】: {cleaned}\n【Period】: {start_dt_str} to {end_dt_str}\n\n"
                "【STRICT TEMPLATE (Use this EXACT structure)】:\n<div style=\"width:100%; font-family:sans-serif;\">\n"
                "    <h3 style=\"color:#D32F2F; border-bottom:2px solid #eee; padding-bottom:10px;\">📊 数据概览表 ({start_dt_str} 至 {end_dt_str})</h3>\n"
                "    <table class=\"table\" style=\"width:100%; margin-bottom:30px; text-align:center; font-size:0.9rem;\">\n"
                "        <tr style=\"background:#f8f9fa;\">\n"
                "            <th>总播放</th><th>总互动</th><th>总点赞</th><th>总评论</th><th>总收藏</th><th>总转发</th>\n"
                "        </tr>\n        <tr>\n"
                "            <td>[总播放数]</td><td>[总互动数]</td><td>[总点赞数]</td><td>[总评论数]</td><td>[总收藏数]</td><td>[总转发数]</td>\n"
                "        </tr>\n    </table>\n\n"
                "    <h3 style=\"color:#D32F2F; border-bottom:2px solid #eee; padding-bottom:10px;\">🔥 爆款视频精选</h3>\n"
                "    <div style=\"background:#FFF9F9; border-left:5px solid #D32F2F; padding:20px; margin-bottom:15px; border-radius:8px;\">\n"
                "        <p><strong>视频描述：</strong> [描述内容]</p>\n"
                "        <p><strong>核心指标：</strong> 播放: [播放数] | 点赞: [点赞数] | 互动: [评论数]评论 / [分享数]分享</p>\n"
                "        <p><strong>查看详情：</strong> <a href=\"[webVideoUrl]\" target=\"_blank\" style=\"color:#2962FF;\">点击进入 TikTok 观看原文链接</a></p>\n"
                "    </div>\n</div>\n\n"
                "【Requirements】:\n- 必须使用中文填充模板。\n- 总互动 = 点赞 + 评论 + 收藏 + 转发的总和。\n"
                "- 严禁添加模板之外的任何文字（包括分析、建议、前言、结语）。\n- 仅输出 Raw HTML 代码，禁止 Markdown 代码块。\n"
            ),
            'PUBGM': '',
            'HOK': '',
        },
    }
    config_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'config', 'prompts.json')
    try:
        if os.path.isfile(config_path):
            with open(config_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            for feat in ('sentiment', 'competitor'):
                if feat not in data:
                    data[feat] = fallback[feat]
                else:
                    for proj in VALID_PROJECTS:
                        if proj not in data[feat]:
                            data[feat][proj] = fallback[feat].get(proj, '')
            _PROMPTS_CACHE = data
            logger.info("✅ 已加载 config/prompts.json")
        else:
            _PROMPTS_CACHE = fallback
            logger.info("⚠️ config/prompts.json 不存在，使用内置 CFL 提示词")
    except Exception as e:
        logger.warning(f"⚠️ 加载 config/prompts.json 失败: {e}，使用内置 CFL 提示词")
        _PROMPTS_CACHE = fallback
    return _PROMPTS_CACHE

def get_prompt(feature, project):
    """取指定功能、项目的提示词模板。feature: sentiment/competitor/copywriting/video_script"""
    if project not in VALID_PROJECTS:
        return ''
    prompts = load_prompts()
    by_feature = prompts.get(feature, {})
    return (by_feature.get(project) or '').strip()


def ensure_task_queue_schema():
    """确保 task_queue 表包含 function_type 字段（向后兼容老版本数据库）

    - 正常情况下会执行一次 ALTER TABLE ADD COLUMN IF NOT EXISTS
    - 若当前数据库用户无权限，或表不存在，只记录 warning，不中断启动
    - create_task 会根据 TASK_QUEUE_HAS_FUNCTION_TYPE 自动降级为老的插入方式
    """
    global TASK_QUEUE_HAS_FUNCTION_TYPE
    try:
        db.execute("""
            ALTER TABLE task_queue
            ADD COLUMN IF NOT EXISTS function_type VARCHAR(50)
        """)
        logger.info("✅ 已确认 task_queue.function_type 列存在")
        TASK_QUEUE_HAS_FUNCTION_TYPE = True
    except Exception as e:
        TASK_QUEUE_HAS_FUNCTION_TYPE = False
        logger.warning(f"⚠️ 无法自动为 task_queue 添加 function_type 列，将使用兼容模式: {e}")
    try:
        db.execute("""
            ALTER TABLE task_queue
            ADD COLUMN IF NOT EXISTS record_id INTEGER
        """)
        logger.info("✅ 已确认 task_queue.record_id 列存在")
    except Exception as e:
        logger.warning(f"⚠️ 无法添加 task_queue.record_id 列: {e}")
    try:
        db.execute("""
            ALTER TABLE task_queue
            ADD COLUMN IF NOT EXISTS task_params TEXT
        """)
        logger.info("✅ 已确认 task_queue.task_params 列存在")
    except Exception as e:
        logger.warning(f"⚠️ 无法添加 task_queue.task_params 列: {e}")
    for column_name, ddl in (
        ("worker_id", "VARCHAR(128)"),
        ("started_at", "TIMESTAMP"),
        ("finished_at", "TIMESTAMP"),
        ("attempts", "INTEGER DEFAULT 0"),
    ):
        try:
            db.execute(f"ALTER TABLE task_queue ADD COLUMN IF NOT EXISTS {column_name} {ddl}")
            logger.info(f"✅ 已确认 task_queue.{column_name} 列存在")
        except Exception as e:
            logger.warning(f"⚠️ 无法添加 task_queue.{column_name} 列: {e}")
    try:
        db.execute("CREATE INDEX IF NOT EXISTS idx_task_queue_status_created ON task_queue(status, created_at)")
    except Exception as e:
        logger.warning(f"⚠️ 无法创建 task_queue 状态索引: {e}")


def ensure_analysis_results_schema():
    """确保 analysis_results 表包含 result_json 字段（用于导出结构化结果）

    - 正常情况下会执行一次 ALTER TABLE ADD COLUMN IF NOT EXISTS
    - 若当前数据库用户无权限，或表不存在，只记录 warning，不中断启动
    """
    global ANALYSIS_RESULTS_HAS_JSON
    try:
        db.execute("""
            ALTER TABLE analysis_results
            ADD COLUMN IF NOT EXISTS result_json TEXT
        """)
        logger.info("✅ 已确认 analysis_results.result_json 列存在")
        ANALYSIS_RESULTS_HAS_JSON = True
    except Exception as e:
        ANALYSIS_RESULTS_HAS_JSON = False
        logger.warning(f"⚠️ 无法自动为 analysis_results 添加 result_json 列，将暂不支持按任意历史记录导出: {e}")


def ensure_etl_file_outputs_schema():
    """ETL 异步任务生成的 Excel（BYTEA），供 Web 与 Worker 分离部署时下载。"""
    try:
        db.execute("""
            CREATE TABLE IF NOT EXISTS etl_file_outputs (
                id SERIAL PRIMARY KEY,
                task_id VARCHAR(128) NOT NULL,
                user_id INTEGER,
                filename VARCHAR(512) NOT NULL,
                content BYTEA NOT NULL,
                created_at TIMESTAMP DEFAULT NOW()
            )
        """)
        db.execute("CREATE INDEX IF NOT EXISTS idx_etl_file_outputs_task_id ON etl_file_outputs (task_id)")
        db.execute("CREATE INDEX IF NOT EXISTS idx_etl_file_outputs_user_id ON etl_file_outputs (user_id)")
        logger.info("✅ 已确认 etl_file_outputs 表存在")
    except Exception as e:
        logger.warning(f"⚠️ 无法创建 etl_file_outputs 表: {e}")


def ensure_agent_actions_schema():
    """AI 任务助手动作审计表。"""
    try:
        db.execute("""
            CREATE TABLE IF NOT EXISTS agent_actions (
                id SERIAL PRIMARY KEY,
                user_id INTEGER REFERENCES users(id),
                chat_session_id INTEGER,
                intent VARCHAR(80) NOT NULL,
                status VARCHAR(32) NOT NULL DEFAULT 'draft',
                params_json TEXT,
                card_json TEXT,
                reply TEXT,
                tool_name VARCHAR(80),
                tool_task_id VARCHAR(160),
                tool_job_id VARCHAR(160),
                result_json TEXT,
                error TEXT,
                confirmed_at TIMESTAMP,
                created_at TIMESTAMP DEFAULT NOW(),
                updated_at TIMESTAMP DEFAULT NOW()
            )
        """)
        db.execute("CREATE INDEX IF NOT EXISTS idx_agent_actions_user_created ON agent_actions(user_id, created_at)")
        db.execute("CREATE INDEX IF NOT EXISTS idx_agent_actions_session ON agent_actions(chat_session_id)")
        logger.info("✅ 已确认 agent_actions 表存在")
    except Exception as e:
        logger.warning(f"⚠️ 无法创建 agent_actions 表: {e}")


def ensure_fb_post_metrics_schema():
    """确保帖子级指标表存在（用于SPD报告按engagement口径计算）。"""
    try:
        db.execute("""
            CREATE TABLE IF NOT EXISTS fb_post_metrics (
                id SERIAL PRIMARY KEY,
                post_url VARCHAR(1024) UNIQUE NOT NULL,
                platform VARCHAR(16),
                author VARCHAR(256),
                post_date DATE,
                post_content TEXT,
                thumbnail_url VARCHAR(1024),
                views BIGINT DEFAULT 0,
                shares BIGINT DEFAULT 0,
                likes BIGINT DEFAULT 0,
                comments_count BIGINT DEFAULT 0,
                engagement BIGINT DEFAULT 0,
                updated_at TIMESTAMP DEFAULT NOW()
            )
        """)
        logger.info("✅ 已确认 fb_post_metrics 表存在")
    except Exception as e:
        logger.warning(f"⚠️ 无法创建 fb_post_metrics 表: {e}")

    # 泰国专题数据集标签表
    try:
        db.execute("""
            CREATE TABLE IF NOT EXISTS thai_report_datasets (
                id SERIAL PRIMARY KEY,
                dataset_name VARCHAR(128) NOT NULL,
                post_url VARCHAR(1024) NOT NULL,
                created_at TIMESTAMP DEFAULT NOW(),
                UNIQUE (dataset_name, post_url)
            )
        """)
        db.execute("CREATE INDEX IF NOT EXISTS idx_thai_datasets_name ON thai_report_datasets (dataset_name)")
        db.execute("CREATE INDEX IF NOT EXISTS idx_thai_datasets_url ON thai_report_datasets (post_url)")
        logger.info("✅ 已确认 thai_report_datasets 表存在")
    except Exception as e:
        logger.warning(f"⚠️ 无法创建 thai_report_datasets 表: {e}")

    # fb_comments 复合索引（加速泰国报告的 JOIN + 时间过滤）
    try:
        db.execute("CREATE INDEX IF NOT EXISTS idx_fb_comments_url_time ON fb_comments (post_url, created_at)")
        db.execute("CREATE INDEX IF NOT EXISTS idx_fb_comments_created_at ON fb_comments (created_at)")
        logger.info("✅ 已确认 fb_comments 索引存在")
    except Exception as e:
        logger.warning(f"⚠️ 无法创建 fb_comments 索引（表可能不存在）: {e}")

def send_feedback_email(project_name: str, feedback: str) -> bool:
    """发送用户反馈邮件到运维邮箱（可选功能）

    依赖环境变量：
    - SMTP_HOST / SMTP_PORT / SMTP_USER / SMTP_PASS
    - FEEDBACK_EMAIL_TO
    """
    if not (SMTP_HOST and SMTP_USER and SMTP_PASS and FEEDBACK_EMAIL_TO):
        logger.warning("⚠️ 反馈邮件未发送：SMTP 或收件人环境变量未完整配置")
        return False

    try:
        subject = f"新用户反馈 - {project_name}"
        body = f"项目名称: {project_name}\n\n反馈内容:\n{feedback}"

        msg = MIMEText(body, "plain", "utf-8")
        msg["Subject"] = subject
        msg["From"] = FEEDBACK_EMAIL_FROM or SMTP_USER
        msg["To"] = FEEDBACK_EMAIL_TO

        with smtplib.SMTP(SMTP_HOST, SMTP_PORT) as server:
            server.starttls()
            server.login(SMTP_USER, SMTP_PASS)
            server.send_message(msg)

        logger.info("✅ 反馈邮件发送成功")
        return True
    except Exception as e:
        logger.error(f"❌ 反馈邮件发送失败: {e}")
        return False


# 汇率配置
USD_TO_CNY = 7.2

# ============================================
# 任务恢复机制（定义，稍后调用）
# ============================================

# 启动时尽早检查相关表结构
ensure_task_queue_schema()
ensure_analysis_results_schema()
ensure_etl_file_outputs_schema()
ensure_agent_actions_schema()
ensure_fb_post_metrics_schema()
profile_video_scheduler.ensure_schema()
tiktok_official_service.ensure_schema()
usage_service.ensure_schema()
rag.ensure_tables()


def recover_interrupted_tasks():
    """恢复被中断的任务。

    判定规则：
    - status='processing' 且 updated_at 已超过 5 分钟没更新（worker 大概率已挂）
    - 同时把同样停滞的 'claimed' 任务回退为 pending，让 worker 重新捡起
    - 不再限制 created_at 窗口，避免老孤儿永远卡在 processing
    """
    try:
        # 1. 把停滞 >5min 的 processing 标 failed
        interrupted_tasks = db.query_all("""
            SELECT task_id FROM task_queue
            WHERE status = 'processing'
              AND updated_at < NOW() - INTERVAL '5 minutes'
        """)
        if interrupted_tasks:
            logger.warning(f"⚠️ 发现 {len(interrupted_tasks)} 个停滞的 processing 任务，标记为失败")
            for task in interrupted_tasks:
                update_task(
                    task['task_id'],
                    status='failed',
                    error='服务重启或 worker 中断导致任务停滞',
                    progress='任务已中断'
                )
        # 2. 把停滞 >2min 的 claimed 回退为 pending（worker 抢到但还没开始执行就死了）
        try:
            db.execute("""
                UPDATE task_queue
                SET status = 'pending', updated_at = NOW()
                WHERE status = 'claimed'
                  AND updated_at < NOW() - INTERVAL '2 minutes'
            """)
        except Exception as e:
            logger.warning(f"⚠️ 回退 claimed 任务失败: {e}")
    except Exception as e:
        logger.error(f"❌ 恢复任务失败: {e}")

# ============================================
# 装饰器：权限控制
# ============================================

def _wants_json_response():
    """Return JSON errors for in-page fetch/API calls instead of HTML redirects."""
    if request.path.startswith('/api/') or request.path in {'/analyze'}:
        return True
    if request.is_json:
        return True
    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        return True
    best = request.accept_mimetypes.best
    return (
        best == 'application/json'
        and request.accept_mimetypes[best] >= request.accept_mimetypes['text/html']
    )


def login_required(f):
    """需要登录才能访问"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not session.get('logged_in'):
            if _wants_json_response():
                return jsonify({
                    'status': 'error',
                    'message': '登录状态已失效，请重新登录后再提交任务',
                    'error': '登录状态已失效，请重新登录后再提交任务'
                }), 401
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function

def admin_required(f):
    """需要管理员权限才能访问"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not session.get('logged_in'):
            if _wants_json_response():
                return jsonify({
                    'status': 'error',
                    'message': '登录状态已失效，请重新登录后再操作',
                    'error': '登录状态已失效，请重新登录后再操作'
                }), 401
            return redirect(url_for('login'))
        if session.get('role') != 'admin':
            return jsonify({'error': '需要管理员权限'}), 403
        return f(*args, **kwargs)
    return decorated_function


@app.errorhandler(Exception)
def handle_fetch_exception(error):
    if not _wants_json_response():
        if isinstance(error, HTTPException):
            return error
        raise error

    status_code = error.code if isinstance(error, HTTPException) else 500
    if status_code >= 500:
        logger.exception("❌ API 请求异常: %s", error)
        message = '服务器处理请求失败，请稍后重试或联系管理员查看后台日志'
    else:
        message = getattr(error, 'description', None) or str(error)
    return jsonify({
        'status': 'error',
        'message': message,
        'error': message
    }), status_code

# ============================================
# 核心工具函数
# ============================================

def create_task(task_id, user_id, session_id, function_type=None):
    """创建任务记录

    为兼容旧库：
    - 优先尝试写入 function_type 字段
    - 若字段不存在或无权限，则退回老的插入方式
    """
    global TASK_QUEUE_HAS_FUNCTION_TYPE

    try:
        if TASK_QUEUE_HAS_FUNCTION_TYPE:
            db.execute("""
                INSERT INTO task_queue (task_id, user_id, session_id, function_type, status, progress)
                VALUES (%s, %s, %s, %s, %s, %s)
            """, (task_id, user_id, session_id, function_type, 'pending', '任务已创建，等待 worker 领取'))
        else:
            db.execute("""
                INSERT INTO task_queue (task_id, user_id, session_id, status, progress)
                VALUES (%s, %s, %s, %s, %s)
            """, (task_id, user_id, session_id, 'pending', '任务已创建，等待 worker 领取'))

        logger.info(f"✅ 任务 {task_id} 已写入数据库（type={function_type}）")
    except Exception as e:
        # 如果是第一次写入发现没有 function_type 列，则自动降级为旧模式
        if TASK_QUEUE_HAS_FUNCTION_TYPE and 'function_type' in str(e):
            logger.warning(f"⚠️ task_queue 缺少 function_type 列，降级为兼容模式: {e}")
            TASK_QUEUE_HAS_FUNCTION_TYPE = False
            try:
                db.execute("""
                    INSERT INTO task_queue (task_id, user_id, session_id, status, progress)
                    VALUES (%s, %s, %s, %s, %s)
                """, (task_id, user_id, session_id, 'pending', '任务已创建'))
                logger.info(f"✅ 任务 {task_id} 已在兼容模式下写入数据库")
            except Exception as e2:
                logger.error(f"❌ 创建任务失败（兼容模式）: {e2}")
        else:
            logger.error(f"❌ 创建任务失败: {e}")

def update_task(task_id, status=None, progress=None, result=None, error=None, record_id=None):
    """更新任务状态"""
    try:
        updates = []
        params = []

        if status is not None:
            updates.append("status = %s")
            params.append(status)
            if status in {'processing', 'claimed'}:
                updates.append("started_at = COALESCE(started_at, CURRENT_TIMESTAMP)")
                updates.append("finished_at = NULL")
            elif status in {'completed', 'failed'}:
                updates.append("finished_at = CURRENT_TIMESTAMP")
        if progress is not None:
            updates.append("progress = %s")
            params.append(progress)
        if result is not None:
            updates.append("result = %s")
            params.append(result)
        if error is not None:
            updates.append("error = %s")
            params.append(error)
        if record_id is not None:
            updates.append("record_id = %s")
            params.append(record_id)

        if updates:
            updates.append("updated_at = CURRENT_TIMESTAMP")
            params.append(task_id)

            sql = f"UPDATE task_queue SET {', '.join(updates)} WHERE task_id = %s"
            db.execute(sql, tuple(params))
    except Exception as e:
        logger.error(f"❌ 更新任务状态失败: {e}")


def set_task_params(task_id, params):
    """写入 task_queue.task_params，供 DB worker 拾取。"""
    db.execute(
        "UPDATE task_queue SET task_params = %s WHERE task_id = %s",
        (json.dumps(params, ensure_ascii=False, default=str), task_id),
    )


def enqueue_due_profile_video_sync(hour=None):
    """APScheduler 回调：按当前小时把启用的主页视频同步配置入队。"""
    try:
        def _after_enqueue(task_id, params):
            if USE_DB_WORKER:
                return

            def _run():
                profile_video_scheduler.run_profile_video_sync_task(task_id, params, update_task)

            threading.Thread(target=_run, daemon=True).start()

        task_ids = profile_video_scheduler.enqueue_due_profile_video_sync(
            create_task,
            update_task_params_fn=set_task_params if USE_DB_WORKER else (lambda _task_id, _params: None),
            hour=hour if hour is not None else datetime.datetime.now().hour,
            after_enqueue_fn=_after_enqueue,
        )
        if task_ids:
            logger.info(f"✅ 已创建主页视频定时同步任务: {task_ids}")
    except Exception as e:
        logger.error(f"❌ 创建主页视频定时同步任务失败: {e}")


def enqueue_due_feishu_profile_video_sync(hour=None):
    """APScheduler 回调：读取飞书配置表并创建主页视频四表同步任务。"""
    try:
        def _after_enqueue(task_id, params):
            if USE_DB_WORKER:
                return

            def _run():
                profile_video_scheduler.run_feishu_profile_video_sync_task(task_id, params, update_task)

            threading.Thread(target=_run, daemon=True).start()

        task_ids = profile_video_scheduler.enqueue_due_feishu_profile_video_sync(
            create_task,
            update_task_params_fn=set_task_params if USE_DB_WORKER else (lambda _task_id, _params: None),
            hour=hour if hour is not None else datetime.datetime.now().hour,
            after_enqueue_fn=_after_enqueue,
        )
        if task_ids:
            logger.info(f"✅ 已创建飞书主页视频自动同步任务: {task_ids}")
    except Exception as e:
        logger.error(f"❌ 创建飞书主页视频自动同步任务失败: {e}")


def get_task(task_id):
    """获取任务状态"""
    try:
        task = db.query_one("""
            SELECT task_id, status, progress, result, error, record_id, function_type,
                   user_id, session_id, worker_id, attempts, created_at, updated_at,
                   started_at, finished_at
            FROM task_queue
            WHERE task_id = %s
        """, (task_id,))
        if task:
            task = dict(task)
            task.update(get_task_queue_position(task_id, task.get('status')))
        return task
    except Exception as e:
        logger.error(f"❌ 获取任务状态失败: {e}")
        return None


def get_task_queue_position(task_id, status=None):
    """Return queue position information for one task."""
    try:
        if status != 'pending':
            running = db.query_one("""
                SELECT COUNT(*) AS count FROM task_queue
                WHERE status IN ('claimed', 'processing')
            """) or {}
            pending = db.query_one("SELECT COUNT(*) AS count FROM task_queue WHERE status = 'pending'") or {}
            return {
                'queue_position': None,
                'tasks_ahead': 0,
                'pending_count': int(pending.get('count') or 0),
                'running_count': int(running.get('count') or 0),
            }
        row = db.query_one("""
            SELECT created_at FROM task_queue WHERE task_id = %s
        """, (task_id,))
        if not row:
            return {}
        ahead = db.query_one("""
            SELECT COUNT(*) AS count FROM task_queue
            WHERE status = 'pending' AND created_at < %s
        """, (row.get('created_at'),)) or {}
        pending = db.query_one("SELECT COUNT(*) AS count FROM task_queue WHERE status = 'pending'") or {}
        running = db.query_one("""
            SELECT COUNT(*) AS count FROM task_queue
            WHERE status IN ('claimed', 'processing')
        """) or {}
        ahead_count = int(ahead.get('count') or 0)
        return {
            'queue_position': ahead_count + 1,
            'tasks_ahead': ahead_count,
            'pending_count': int(pending.get('count') or 0),
            'running_count': int(running.get('count') or 0),
        }
    except Exception as e:
        logger.warning(f"⚠️ 计算任务排队位置失败: {e}")
        return {}

# ============================================
# 启动时恢复被中断的任务（仅 Web 进程执行）
# ============================================
if not _IS_WORKER:
    recover_interrupted_tasks()


def call_gemini(prompt, image=None, timeout=60, model=None, temperature=0.7):
    """调用通义千问 API

    Args:
        model: 显式指定模型；若 None，按 QWEN_DEFAULT_MODEL 环境变量，默认 qwen-turbo。
    """
    if not qwen_client:
        error_msg = "❌ 错误：DASHSCOPE_API_KEY 未配置"
        logger.error(error_msg)
        return error_msg, 0

    model_name = model or os.environ.get('QWEN_DEFAULT_MODEL') or 'qwen-turbo'

    try:
        logger.info(f"🤖 正在调用通义千问模型: {model_name}")
        logger.info(f"📏 Prompt 长度: {len(prompt)} 字符")

        response = qwen_client.chat.completions.create(
            model=model_name,
            messages=[
                {"role": "user", "content": prompt}
            ],
            temperature=temperature,
            timeout=timeout,
        )

        result = response.choices[0].message.content
        tokens = response.usage.total_tokens if hasattr(response, 'usage') else 0
        logger.info(f"✅ 通义千问调用成功，返回 {len(result)} 字符，消耗 {tokens} tokens")
        return result, tokens

    except Exception as e:
        error_msg = f"⚠️ 通义千问 API 调用失败: {str(e)}"
        logger.error(error_msg)
        import traceback
        traceback.print_exc()
        return error_msg, 0


def process_uploaded_file(file_data):
    """处理上传的文件（图片或表格）

    Args:
        file_data: 字典，包含 filename, content, content_type
    """
    try:
        fname = file_data['filename'].lower()
        content = file_data['content']
        logger.info(f"📁 处理文件: {fname}")

        if fname.endswith(('.png', '.jpg', '.jpeg', '.webp')):
            logger.info("🖼️ 识别为图片文件")
            return "IMAGE", Image.open(BytesIO(content))

        if fname.endswith(('.xlsx', '.csv')):
            logger.info("📊 识别为表格文件")
            if fname.endswith('.csv'):
                df = pd.read_csv(BytesIO(content))
            else:
                df = pd.read_excel(BytesIO(content))
            return "TEXT", df.to_string(index=False, max_rows=50)

        return "ERROR", "不支持的文件格式"

    except Exception as e:
        error_msg = f"文件处理失败: {str(e)}"
        logger.info(f"❌ {error_msg}")
        return "ERROR", error_msg


def save_history(user_id, title, result, type_tag, structured=None):
    """保存到历史记录（数据库 + 内存），并可选保存结构化结果

    - user_id: 用户 ID（可为空，空时仅保存到内存）
    - title/result/type_tag: 展示用的标题与 HTML 结果
    - structured: 可选的结构化结果（Python 列表/字典），会序列化到 result_json

    注意：不要在此函数内部访问 Flask session，
    需要在调用方把 user_id 显式传入，以便在线程中安全调用。
    返回：数据库记录 ID（成功且有 user_id 且表结构支持时），否则 None
    """
    try:
        record_id = None

        if not user_id:
            logger.warning("⚠️ 未提供 user_id，仅保存内存历史记录")
        else:
            # 保存到数据库，并返回记录 ID
            try:
                record_id = db.execute_and_fetch_id("""
                    INSERT INTO analysis_results (user_id, title, result, type)
                    VALUES (%s, %s, %s, %s)
                    RETURNING id
                """, (user_id, title, result, type_tag))
                logger.info(f"💾 已保存历史记录到数据库: {title} (id={record_id})")
            except Exception as e:
                logger.error(f"❌ 保存历史记录到数据库失败: {e}")

            # 若提供了结构化结果且表结构支持，尝试写入 result_json
            if record_id and structured is not None and ANALYSIS_RESULTS_HAS_JSON:
                try:
                    db.execute("""
                        UPDATE analysis_results
                        SET result_json = %s
                        WHERE id = %s
                    """, (json.dumps(structured, ensure_ascii=False), record_id))
                    logger.info(f"💾 已为记录 {record_id} 写入结构化结果 result_json")
                except Exception as e:
                    logger.warning(f"⚠️ 写入 result_json 失败，将继续使用 HTML 结果: {e}")

        # 同时保存到内存（向后兼容）
        record = {
            'id': len(HISTORY_DB) + 1,
            'title': f"{title} [{datetime.datetime.now().strftime('%H:%M')}]",
            'result': result,
            'type': type_tag
        }
        HISTORY_DB.append(record)

        return record_id

    except Exception as e:
        logger.error(f"❌ 保存历史记录失败: {e}")
        # 失败时至少保存到内存
        record = {
            'id': len(HISTORY_DB) + 1,
            'title': f"{title} [{datetime.datetime.now().strftime('%H:%M')}]",
            'result': result,
            'type': type_tag
        }
        HISTORY_DB.append(record)
        return None


def _truthy_form_value(value):
    if isinstance(value, bool):
        return value
    if value is None:
        return False
    return str(value).strip().lower() in {'1', 'true', 'yes', 'y', 'on', 'unlimited', 'none', 'no_limit', '不设上限', '不限'}


def normalize_insight_comment_limit(value=None, unlimited=False):
    return sentiment_insight.normalize_comments_per_post_limit(value, unlimited=unlimited)


def call_veo_api(prompt):
    """调用 Google Veo API（模拟）"""
    logger.info(f"🎬 模拟 Veo API 调用: {prompt[:50]}...")
    time.sleep(3)
    return "https://cdn.pixabay.com/video/2023/10/22/186115-877653483_large.mp4"


def log_usage(user_id, username, department, function_type, comments_count, ai_tokens, task_id=None, record_id=None):
    """记录使用情况和成本。

    兼容旧 usage_logs，同时写入 usage_events。爬虫费用统一按每 1000 条 3 美金计算。
    """
    try:
        ai_cost = ai_tokens * 0.008 / 1000  # 通义千问估算价：人民币/千 token
        apify_cost_usd = comments_count * 3.00 / 1000
        apify_cost = apify_cost_usd * USD_TO_CNY
        total_cost = ai_cost + apify_cost

        db.execute("""
            INSERT INTO usage_logs
            (user_id, username, department, function_type, comments_count,
             ai_tokens, ai_cost, apify_cost, total_cost)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
        """, (user_id, username, department, function_type, comments_count,
              ai_tokens, ai_cost, apify_cost, total_cost))

        usage_service.record_usage_event(
            module=function_type,
            user_id=user_id,
            username=username,
            department=department,
            task_id=task_id,
            record_id=record_id,
            item_count=comments_count,
            crawler_items=comments_count,
            ai_tokens=ai_tokens,
            source='actual',
            detail={'legacy_usage_logs': True, 'pricing_note': 'crawler USD 3 / 1000 rows'},
        )

        logger.info(f"💰 成本记录: AI={ai_cost:.4f}元 + 爬虫=${apify_cost_usd:.4f}/{apify_cost:.4f}元 = 总计{total_cost:.4f}元")

        return total_cost

    except Exception as e:
        logger.error(f"❌ 记录使用情况失败: {e}")
        return 0


def _format_insight_scrape_failure(summary):
    parts = []
    for item in summary[:5]:
        platform = item.get('platform') or 'UNKNOWN'
        item_count = item.get('item_count', 0)
        comment_count = item.get('comment_count', 0)
        error = (item.get('error') or '').strip()
        url = (item.get('url') or '').strip()
        resolved_url = (item.get('resolved_url') or '').strip()
        actor_run_id = (item.get('actor_run_id') or '').strip()
        actor_dataset_id = (item.get('actor_dataset_id') or '').strip()
        short_url = url[:80] + ('...' if len(url) > 80 else '')
        extras = []
        if resolved_url and resolved_url != url:
            short_resolved = resolved_url[:80] + ('...' if len(resolved_url) > 80 else '')
            extras.append(f"解析后: {short_resolved}")
        if actor_run_id:
            extras.append(f"run: {actor_run_id}")
        if actor_dataset_id:
            extras.append(f"dataset: {actor_dataset_id}")
        extra_text = f"，{'，'.join(extras)}" if extras else ""
        if error:
            parts.append(f"{platform} 返回 {item_count} 条原始数据/{comment_count} 条评论，错误: {error}，链接: {short_url}{extra_text}")
        else:
            parts.append(f"{platform} 返回 {item_count} 条原始数据/{comment_count} 条评论，链接: {short_url}{extra_text}")
    if len(summary) > 5:
        parts.append(f"另有 {len(summary) - 5} 条链接未展示")
    return "；".join(parts)[:1000]


def process_analysis_task(task_id, url=None, file_data=None, session_id='default', user_id=None,
                          username='unknown', department='未知', project='CFL', urls=None,
                          comments_per_post_limit=None):
    """异步处理分析任务。project 为 CFL/PUBGM/HOK，用于选择提示词。

    入参兼容：
      - urls: 多链接列表（新链路，舆情洞察 v2）
      - url:  单链接（兼容旧任务，会被规整成 [url] 走新链路）
      - file_data: 文件上传（仍走旧的 6 列 AI 分类链路）
    """
    if urls is None:
        urls = []
    elif isinstance(urls, str):
        urls = [urls]
    if not urls and url:
        urls = [url]
    comments_per_post_limit = normalize_insight_comment_limit(comments_per_post_limit)

    logger.info(f"🔄 后台线程已启动，任务ID: {task_id}，项目: {project}")
    logger.info(f"👤 用户信息: user_id={user_id}, username={username}, department={department}")
    logger.info(f"📋 任务参数: urls={len(urls)} 条, has_file={file_data is not None}, comments_per_post_limit={comments_per_post_limit}")

    # 在线程中创建新的 Apify 客户端（避免线程安全问题）
    thread_apify_client = None
    if APIFY_TOKEN:
        try:
            logger.info("🔧 在后台线程中初始化 Apify 客户端...")
            thread_apify_client = ApifyClient(APIFY_TOKEN)
            logger.info("✅ 线程 Apify 客户端初始化成功")
        except Exception as e:
            logger.error(f"❌ 线程 Apify 客户端初始化失败: {e}")
            update_task(task_id, status='failed', error=f"Apify 客户端初始化失败: {e}")
            return

    # 追踪成本数据
    total_tokens = 0
    total_comments = 0

    try:
        logger.info(f"📝 更新任务状态为 processing...")
        update_task(task_id, status='processing', progress='正在初始化...')
        logger.info(f"✅ 任务状态更新成功")

        content = ""
        img = None
        source_title = "未知"

        # 路径 A: 文件上传分析
        if file_data:
            update_task(task_id, progress='正在处理文件...')
            mode, res = process_uploaded_file(file_data)

            if mode == "ERROR":
                update_task(task_id, status='failed', error=res)
                return

            if mode == "IMAGE":
                img = res
                content = "分析图片中的反馈内容"
            else:
                content = res

            source_title = f"文件: {file_data.filename[:15]}"

        # 路径 B: 多平台社媒链接抓取分析（v2：FB/IG/TT/YTB）
        elif urls:
            logger.info(f"🌐 开始处理 {len(urls)} 条链接（v2 多平台舆情洞察）")
            update_task(task_id, progress=f'正在抓取 {len(urls)} 条社媒数据...')

            if not APIFY_TOKEN:
                logger.error("❌ APIFY_TOKEN 未配置")
                update_task(task_id, status='failed', error="APIFY_TOKEN 未配置")
                return

            try:
                def _progress(msg):
                    update_task(task_id, progress=msg)

                insight_model = os.environ.get('INSIGHT_AI_MODEL') or 'qwen-plus'

                def _ai_call(p, t=60):
                    return call_gemini(p, timeout=t, model=insight_model, temperature=0.3)

                pipeline_result = sentiment_insight.run_insight_pipeline(
                    urls=urls,
                    apify_token=APIFY_TOKEN,
                    ai_call=_ai_call,
                    progress=_progress,
                    comments_per_post_limit=comments_per_post_limit,
                )

                all_results = pipeline_result['structured']
                result = pipeline_result['html']
                total_tokens = pipeline_result['total_tokens']
                total_comments = pipeline_result['total_comments']

                if not all_results:
                    detail = _format_insight_scrape_failure(pipeline_result.get('scrape_summary') or [])
                    update_task(
                        task_id,
                        status='failed',
                        error='未抓取到任何评论' + (f'：{detail}' if detail else '（请检查链接或 Actor 兼容性）'),
                    )
                    return

                LATEST_ANALYSIS_RESULTS[session_id] = all_results

                platforms = sorted({r.get('platform', 'UNKNOWN') for r in all_results})
                source_title = f"舆情洞察 · {'/'.join(platforms)} · {len(urls)}链接"

                # 复用 try/except 收尾逻辑（保持原代码结构最少改动）
                _legacy_url_branch_done = True
                logger.info(f"✅ v2 链路完成：评论 {total_comments} 条，平台 {platforms}")

            except Exception as e:
                error_msg = f"舆情洞察任务失败: {e}"
                logger.error(f"❌ {error_msg}")
                import traceback
                logger.error(f"❌ 完整堆栈:\n{traceback.format_exc()}")
                update_task(task_id, status='failed', error=error_msg)
                return

        # ↓↓↓ 旧 FB 单链链路已下线，下面这段全部为不可达死代码（保留以最少改动）↓↓↓
        elif False:
            try:
                logger.info("legacy branch placeholder")
                start_time = time.time()
                try:
                    api_url = "https://api.apify.com/v2/acts/apify~facebook-comments-scraper/runs"
                    headers = {
                        "Authorization": f"Bearer {APIFY_TOKEN}",
                        "Content-Type": "application/json"
                    }

                    logger.info(f"   API URL: {api_url}")
                    logger.info(f"   使用 requests 库，超时: 30秒")

                    response = requests.post(
                        api_url,
                        json=run_input,
                        headers=headers,
                        timeout=30  # 30 秒超时
                    )

                    elapsed = time.time() - start_time
                    logger.info(f"✅ HTTP 请求完成（耗时: {elapsed:.2f}秒）")
                    logger.info(f"   状态码: {response.status_code}")

                    if response.status_code != 201:
                        raise ValueError(f"Apify API 返回错误状态码: {response.status_code}, 响应: {response.text}")

                    run = response.json()['data']
                    logger.info(f"✅ Apify API 返回成功")
                    logger.info(f"   返回类型: {type(run)}")
                    logger.info(f"   Run ID: {run.get('id') if run else 'None'}")

                    if not run or 'id' not in run:
                        raise ValueError(f"Apify 返回无效: {run}")

                    logger.info(f"✅ 爬虫任务已启动，Run ID: {run['id']}")

                except requests.Timeout:
                    error_msg = "Apify API 调用超时（30秒）"
                    logger.error(f"❌ {error_msg}")
                    update_task(task_id, status='failed', error=error_msg)
                    return
                except Exception as start_error:
                    error_msg = f"启动爬虫失败: {str(start_error)}"
                    logger.error(f"❌ {error_msg}")
                    logger.error(f"   错误类型: {type(start_error).__name__}")
                    import traceback
                    logger.error(f"   堆栈:\n{traceback.format_exc()}")
                    update_task(task_id, status='failed', error=error_msg)
                    return

                logger.info("⏳ 等待爬虫完成（最长 480 秒）...")
                update_task(task_id, progress='等待爬虫完成（约30-60秒）...')

                try:
                    logger.info("📡 开始轮询 Apify 任务状态...")
                    start_time = time.time()
                    max_wait_time = 480  # 最多等待 480 秒
                    poll_interval = 5  # 每 5 秒轮询一次

                    run_id = run['id']
                    api_url = f"https://api.apify.com/v2/actor-runs/{run_id}"
                    headers = {"Authorization": f"Bearer {APIFY_TOKEN}"}

                    while True:
                        elapsed = time.time() - start_time
                        if elapsed > max_wait_time:
                            raise TimeoutError(f"等待爬虫完成超时（{max_wait_time}秒）")

                        # 轮询任务状态
                        logger.info(f"   轮询状态... (已等待 {elapsed:.0f}秒)")
                        response = requests.get(api_url, headers=headers, timeout=10)

                        if response.status_code != 200:
                            raise ValueError(f"获取任务状态失败: {response.status_code}")

                        run_data = response.json()['data']
                        status = run_data['status']

                        logger.info(f"   当前状态: {status}")

                        if status in ['SUCCEEDED', 'FAILED', 'ABORTED', 'TIMED-OUT']:
                            # 任务完成
                            run = run_data
                            break

                        # 等待后继续轮询
                        time.sleep(poll_interval)

                    elapsed = time.time() - start_time
                    logger.info(f"✅ 爬虫完成，状态: {run['status']}，耗时: {elapsed:.1f}秒")

                except requests.Timeout:
                    error_msg = "轮询任务状态超时"
                    logger.error(f"❌ {error_msg}")
                    update_task(task_id, status='failed', error=error_msg)
                    return
                except TimeoutError as timeout_error:
                    error_msg = str(timeout_error)
                    logger.error(f"❌ {error_msg}")
                    update_task(task_id, status='failed', error=error_msg)
                    return
                except Exception as wait_error:
                    elapsed = time.time() - start_time if 'start_time' in locals() else 0
                    error_msg = f"等待爬虫完成失败（耗时 {elapsed:.1f}秒）: {str(wait_error)}"
                    logger.error(f"❌ {error_msg}")
                    update_task(task_id, status='failed', error=error_msg)
                    return

                if run['status'] != 'SUCCEEDED':
                    logger.error(f"❌ 爬虫任务失败: {run['status']}")
                    update_task(task_id, status='failed', error=f"爬虫任务失败: {run['status']}")
                    return

                # 获取数据
                logger.info("📦 开始获取爬虫数据...")
                dataset_id = run.get("defaultDatasetId")
                if not dataset_id:
                    error_msg = "未找到 dataset ID"
                    logger.error(f"❌ {error_msg}")
                    update_task(task_id, status='failed', error=error_msg)
                    return

                # 使用 REST API 获取数据
                dataset_url = f"https://api.apify.com/v2/datasets/{dataset_id}/items"
                try:
                    response = requests.get(dataset_url, headers=headers, timeout=30)
                    if response.status_code != 200:
                        raise ValueError(f"获取数据失败: {response.status_code}")
                    items = response.json()
                    logger.info(f"✅ 总共获取到 {len(items)} 条数据")
                except Exception as e:
                    error_msg = f"获取数据失败: {str(e)}"
                    logger.error(f"❌ {error_msg}")
                    update_task(task_id, status='failed', error=error_msg)
                    return
                total_comments = len(items)  # 记录评论数

                if not items:
                    update_task(task_id, status='failed', error="未发现公开评论")
                    return

                # 按项目取提示词模板；空则直接失败
                sentiment_template = get_prompt('sentiment', project)
                if not sentiment_template:
                    update_task(task_id, status='failed', error='该项目提示词尚未配置')
                    return

                # 分批处理评论
                batch_size = 50
                all_results = []
                total_batches = (len(items) + batch_size - 1) // batch_size

                for i in range(0, len(items), batch_size):
                    batch = items[i:i+batch_size]
                    batch_num = i // batch_size + 1

                    update_task(task_id, progress=f'AI 分析中：第 {batch_num}/{total_batches} 批...')
                    logger.info(f"🔄 处理第 {batch_num}/{total_batches} 批（{len(batch)} 条评论）...")

                    batch_content = "\n".join([f"用户{j}: {it.get('text', '')}" for j, it in enumerate(batch)])
                    batch_prompt = sentiment_template.format(batch_content=batch_content)

                    result, tokens = call_gemini(batch_prompt, timeout=60)
                    total_tokens += tokens

                    # 解析 JSON 结果
                    try:
                        import json
                        import re
                        clean_result = re.sub(r'```json\\s*|\\s*```', '', result).strip()
                        batch_data = json.loads(clean_result)
                        all_results.extend(batch_data)
                        logger.info(f"✅ 第 {batch_num} 批完成，获得 {len(batch_data)} 条有效结果")
                    except Exception as e:
                        logger.error(f"❌ 第 {batch_num} 批解析失败: {e}")
                        continue

                # 生成 HTML 表格
                update_task(task_id, progress='生成报告...')
                logger.info(f"📝 生成最终报告，共 {len(all_results)} 条有效评论...")

                # 按分类排序
                category_order = ["外挂作弊", "游戏优化", "游戏Bug", "充值退款", "新模式/地图/平衡性建议", "其他"]
                all_results.sort(key=lambda x: category_order.index(x.get('category', '其他')) if x.get('category') in category_order else len(category_order))

                # 生成 HTML
                html_rows = []
                for idx, item in enumerate(all_results, 1):
                    html_rows.append(f"""
                    <tr>
                        <td>{idx}</td>
                        <td style="white-space: pre-wrap; word-break: break-word;">{item.get('text', '')}</td>
                        <td><strong>{item.get('category', '')}</strong></td>
                        <td>{item.get('sentiment', '')}</td>
                        <td>{item.get('language', '')}</td>
                        <td style="white-space: pre-wrap; word-break: break-word;">{item.get('analysis', '')}</td>
                    </tr>
                    """)

                result = f"""
                <table class="table table-hover">
                    <thead>
                        <tr>
                            <th style="width:40px;">#</th>
                            <th style="width:25%;">原始评论</th>
                            <th style="width:100px;">归类</th>
                            <th style="width:70px;">情感</th>
                            <th style="width:60px;">语言</th>
                            <th>简要分析</th>
                        </tr>
                    </thead>
                    <tbody>
                        {''.join(html_rows)}
                    </tbody>
                </table>
                """

                # 保存结果用于导出（后续将改为从数据库读取）
                LATEST_ANALYSIS_RESULTS[session_id] = all_results
                source_title = f"FB: {url[:15]}..."

            except Exception as e:
                error_msg = f"爬虫任务失败: {str(e)}"
                logger.error(f"❌ {error_msg}")
                logger.error(f"❌ 错误类型: {type(e).__name__}")
                import traceback
                logger.error(f"❌ 完整堆栈:\n{traceback.format_exc()}")

                update_task(task_id, status='failed', error=error_msg)
                return

        else:
            update_task(task_id, status='failed', error="请提供链接或文件")
            return

        # 保存历史记录（同时写入结构化结果，便于后续导出任意历史记录）
        record_id = save_history(user_id, source_title, result, 'sentiment', structured=all_results)

        # 记录使用成本
        if user_id:
            log_usage(user_id, username, department, 'sentiment', total_comments, total_tokens, task_id=task_id, record_id=record_id)

        # 任务完成
        update_task(task_id, status='completed', result=result, progress='分析完成！', record_id=record_id)
        logger.info(f"✅ 任务 {task_id} 完成")

    except Exception as e:
        error_msg = f"系统错误: {str(e)}"
        logger.error(f"❌ 任务 {task_id} 失败: {e}")
        logger.error(f"❌ 错误类型: {type(e).__name__}")
        import traceback
        logger.error(f"❌ 完整堆栈:\n{traceback.format_exc()}")

        update_task(task_id, status='failed', error=error_msg, progress='任务失败')

# ============================================
# 基础路由
# ============================================

@app.route('/login', methods=['GET', 'POST'])
def login():
    """登录页面"""
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')

        try:
            # 从数据库查询用户
            user = db.query_one(
                "SELECT * FROM users WHERE username = %s",
                (username,)
            )

            if user and bcrypt.check_password_hash(user['password_hash'], password):
                session['logged_in'] = True
                session['user_id'] = user['id']
                session['username'] = user['username']
                session['real_name'] = user['real_name']
                session['department'] = user['department']
                session['role'] = user['role']
                session['session_id'] = f"{username}_{int(time.time())}"
                logger.info(f"✅ 用户登录成功: {username} ({user['real_name']})")
                return redirect(url_for('home'))
            else:
                logger.info(f"❌ 登录失败: {username}")
                return render_template('login.html', error='用户名或密码错误')

        except Exception as e:
            logger.error(f"❌ 登录异常: {e}")
            return render_template('login.html', error='系统错误，请稍后重试')

    return render_template('login.html')


@app.route('/logout')
def logout():
    """登出"""
    username = session.get('username', '未知用户')
    session.clear()
    logger.info(f"👋 用户已登出: {username}")
    return redirect(url_for('login'))


@app.route('/')
@login_required
def home():
    """首页"""
    return render_template('index.html', user=session)


@app.route('/agent-tool')
@login_required
def agent_tool():
    """AI 任务助手页面"""
    return render_template('agent.html', user=session)


@app.route('/tiktok/callback/', strict_slashes=False)
def tiktok_callback():
    """公开 TikTok OAuth 回调占位页，用于开发者后台 URL 校验。"""
    return _render_tiktok_oauth_callback('legacy')


@app.route('/tiktok/business/callback/', strict_slashes=False)
def tiktok_business_callback():
    """公开 TikTok 广告主授权回调页。不要加登录保护，供 TikTok 校验和 OAuth 返回。"""
    return _render_tiktok_oauth_callback('business')


@app.route('/tiktok/account/callback/', strict_slashes=False)
def tiktok_account_callback():
    """公开 TikTok 账号持有人授权回调页。不要加登录保护，供 TikTok 校验和 OAuth 返回。"""
    return _render_tiktok_oauth_callback('account')


def _tiktok_public_base_url():
    public_base = (os.environ.get('PUBLIC_BASE_URL') or request.url_root).rstrip('/')
    if public_base in {'https://www.sailson.com', 'https://sailson.com'}:
        return 'https://sailson-ai.onrender.com'
    return public_base


def _render_tiktok_oauth_callback(callback_type: str):
    """渲染 TikTok OAuth 回调结果，便于复制 code 换 token。"""
    code = request.args.get('code')
    state = request.args.get('state')
    error = request.args.get('error')
    error_description = request.args.get('error_description')
    logger.info(f"TikTok OAuth callback received: type={callback_type}, state={state or '-'}, has_code={bool(code)}")
    title = f"TikTok {callback_type} callback"
    if error:
        body = f"""
        <p style="color:#b91c1c;">授权失败</p>
        <p><strong>error:</strong> <code>{html.escape(error)}</code></p>
        <p><strong>description:</strong> <code>{html.escape(error_description or '')}</code></p>
        """
    elif code:
        if callback_type == 'account':
            try:
                public_base = _tiktok_public_base_url()
                redirect_uri = f'{public_base}/tiktok/account/callback'
                token_data = tiktok_official_service.exchange_account_code(code, redirect_uri)
                open_id = token_data.get('open_id') or ''
                scope = token_data.get('scope') or ''
                body = f"""
                <p style="color:#166534;">授权成功，access token 已自动保存。</p>
                <p><strong>open_id / business_id:</strong> <code>{html.escape(open_id)}</code></p>
                <p><strong>scope:</strong> <code>{html.escape(str(scope))}</code></p>
                <p><a href="/tiktok-official" style="color:#1a7fd4;">返回 TikTok 官号监控</a></p>
                """
            except Exception as e:
                logger.error(f"TikTok account token exchange failed: {e}")
                body = f"""
                <p style="color:#b91c1c;">收到 code，但自动换 token 失败。</p>
                <p><strong>错误：</strong><code>{html.escape(str(e))}</code></p>
                <p><strong>code:</strong></p>
                <textarea readonly style="width:100%;height:90px;">{html.escape(code)}</textarea>
                <p><strong>state:</strong> <code>{html.escape(state or '')}</code></p>
                """
        else:
            body = f"""
            <p style="color:#166534;">授权成功，复制下面的 code 给开发者换取 access token。</p>
            <p><strong>type:</strong> <code>{html.escape(callback_type)}</code></p>
            <p><strong>code:</strong></p>
            <textarea readonly style="width:100%;height:90px;">{html.escape(code)}</textarea>
            <p><strong>state:</strong> <code>{html.escape(state or '')}</code></p>
            """
    else:
        body = f"""
        <p>TikTok callback ready.</p>
        <p>当前回调类型：<code>{html.escape(callback_type)}</code></p>
        """
    return f"""
    <!doctype html>
    <html lang="zh-CN">
    <head>
      <meta charset="utf-8">
      <meta name="viewport" content="width=device-width, initial-scale=1">
      <title>{html.escape(title)}</title>
    </head>
    <body style="font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',sans-serif;background:#f8fafc;color:#111827;padding:32px;">
      <div style="max-width:760px;margin:0 auto;background:#fff;border-radius:12px;padding:24px;box-shadow:0 2px 10px rgba(0,0,0,.08);">
        <h1 style="font-size:20px;margin-bottom:16px;">{html.escape(title)}</h1>
        {body}
      </div>
    </body>
    </html>
    """, 200


@app.route('/kol-tool')
@login_required
def kol_tool():
    """KOL 名单管理工具"""
    index_path = os.path.join(KOL_STATIC_DIR, 'index.html')
    if not os.path.exists(index_path):
        logger.error(f"KOL 前端资源不存在: {index_path}")
        return "KOL 前端资源尚未构建", 503
    return send_file(index_path)


@app.route('/kol-api', defaults={'path': ''}, methods=['GET', 'POST', 'PUT', 'PATCH', 'DELETE', 'OPTIONS'])
@app.route('/kol-api/<path:path>', methods=['GET', 'POST', 'PUT', 'PATCH', 'DELETE', 'OPTIONS'])
@login_required
def kol_api_proxy(path):
    """登录保护的 KOL FastAPI 代理"""
    if not KOL_API_BASE_URL:
        return jsonify({'error': 'KOL_API_BASE_URL 未配置'}), 503

    api_base = KOL_API_BASE_URL.rstrip('/')
    if not api_base.startswith(('http://', 'https://')):
        api_base = f"http://{api_base}"
    if not api_base.endswith('/api'):
        api_base = f"{api_base}/api"
    target_url = f"{api_base}/{path}" if path else api_base

    excluded_headers = {'host', 'content-length', 'connection', 'accept-encoding'}
    if request.files:
        excluded_headers.add('content-type')
    headers = {
        key: value
        for key, value in request.headers
        if key.lower() not in excluded_headers
    }
    if KOL_PROXY_TOKEN:
        headers['X-KOL-Proxy-Token'] = KOL_PROXY_TOKEN

    files = None
    data = None
    if request.files:
        files = []
        for field_name in request.files:
            for storage in request.files.getlist(field_name):
                files.append((
                    field_name,
                    (storage.filename, storage.stream, storage.mimetype or 'application/octet-stream')
                ))
        data = request.form
    elif request.method not in {'GET', 'HEAD'}:
        data = request.get_data()

    try:
        upstream = requests.request(
            request.method,
            target_url,
            params=request.args,
            headers=headers,
            data=data,
            files=files,
            timeout=300,
            allow_redirects=False
        )
    except requests.RequestException as e:
        logger.error(f"KOL 服务代理失败: {e}")
        return jsonify({'error': 'KOL 服务暂时不可用'}), 502

    excluded_response_headers = {
        'content-encoding',
        'content-length',
        'transfer-encoding',
        'connection'
    }
    response_headers = [
        (key, value)
        for key, value in upstream.headers.items()
        if key.lower() not in excluded_response_headers
    ]
    return Response(upstream.content, status=upstream.status_code, headers=response_headers)


# ============================================
# AI 任务助手
# ============================================

AGENT_ACTION_EXECUTABLES = {
    'sentiment_comments',
    'video_metrics',
    'profile_video_sync',
    'kol_data_refresh_links',
    'kol_data_refresh_excel',
    'task_query',
}


def _agent_json(value):
    return json.dumps(value, ensure_ascii=False, default=str)


def _agent_load_json(value, default=None):
    if not value:
        return default if default is not None else {}
    try:
        return json.loads(value) if isinstance(value, str) else value
    except Exception:
        return default if default is not None else {}


def _agent_session(session_id_chat=None, title='AI 任务助手'):
    user_id = session.get('user_id')
    if session_id_chat:
        row = db.query_one(
            "SELECT id FROM chat_sessions WHERE id = %s AND user_id = %s",
            (session_id_chat, user_id),
        )
        if row:
            return int(row['id'])
    return db.execute_and_fetch_id(
        """
        INSERT INTO chat_sessions (user_id, mode, project, title)
        VALUES (%s, %s, %s, %s) RETURNING id
        """,
        (user_id, 'agent', 'CFL', title[:256]),
    )


def _agent_insert_message(session_id_chat, role, content):
    db.execute(
        "INSERT INTO chat_messages (session_id, role, content) VALUES (%s, %s, %s)",
        (session_id_chat, role, content),
    )


def _agent_create_action(chat_session_id, draft):
    user_id = session.get('user_id')
    return db.execute_and_fetch_id(
        """
        INSERT INTO agent_actions (
            user_id, chat_session_id, intent, status, params_json, card_json, reply
        )
        VALUES (%s, %s, %s, %s, %s, %s, %s)
        RETURNING id
        """,
        (
            user_id,
            chat_session_id,
            draft.intent,
            'draft' if draft.needs_confirmation else 'ready',
            _agent_json(draft.params),
            _agent_json(draft.card) if draft.card else None,
            draft.reply,
        ),
    )


def _agent_get_action(action_id):
    row = db.query_one(
        """
        SELECT *
        FROM agent_actions
        WHERE id = %s AND user_id = %s
        """,
        (action_id, session.get('user_id')),
    )
    return dict(row) if row else None


def _agent_update_action(action_id, **fields):
    allowed = {
        'status', 'tool_name', 'tool_task_id', 'tool_job_id',
        'result_json', 'error', 'confirmed_at', 'card_json', 'reply',
    }
    updates = []
    params = []
    for key, value in fields.items():
        if key not in allowed:
            continue
        updates.append(f"{key} = %s")
        params.append(value)
    if not updates:
        return
    updates.append("updated_at = NOW()")
    params.append(action_id)
    params.append(session.get('user_id'))
    db.execute(
        f"UPDATE agent_actions SET {', '.join(updates)} WHERE id = %s AND user_id = %s",
        tuple(params),
    )


def _agent_action_out(row):
    result = {
        'id': row.get('id'),
        'session_id': row.get('chat_session_id'),
        'intent': row.get('intent'),
        'status': row.get('status'),
        'reply': row.get('reply'),
        'params': _agent_load_json(row.get('params_json'), {}),
        'card': _agent_load_json(row.get('card_json'), None),
        'tool_name': row.get('tool_name'),
        'tool_task_id': row.get('tool_task_id'),
        'tool_job_id': row.get('tool_job_id'),
        'result': _agent_load_json(row.get('result_json'), None),
        'error': row.get('error'),
        'created_at': row.get('created_at'),
        'updated_at': row.get('updated_at'),
        'confirmed_at': row.get('confirmed_at'),
    }
    return _json_safe(result)


def _agent_kol_api_url(path):
    if not KOL_API_BASE_URL:
        raise RuntimeError('KOL_API_BASE_URL 未配置')
    api_base = KOL_API_BASE_URL.rstrip('/')
    if not api_base.startswith(('http://', 'https://')):
        api_base = f"http://{api_base}"
    if not api_base.endswith('/api'):
        api_base = f"{api_base}/api"
    return f"{api_base}/{path.lstrip('/')}"


def _agent_kol_headers():
    headers = {}
    if KOL_PROXY_TOKEN:
        headers['X-KOL-Proxy-Token'] = KOL_PROXY_TOKEN
    return headers


def _agent_create_sentiment_task(params):
    urls = [u for u in (params.get('urls') or []) if sentiment_insight.detect_platform(u) != 'UNKNOWN']
    if not urls:
        raise ValueError('未识别到可分析的社媒评论链接')
    task_id = str(uuid.uuid4())
    session_id = session.get('session_id', 'default')
    user_id = session.get('user_id')
    username = session.get('username', 'unknown')
    department = session.get('department', '未知')
    create_task(task_id, user_id, session_id, function_type='sentiment')
    task_params = {
        'urls': urls,
        'session_id': session_id,
        'user_id': user_id,
        'username': username,
        'department': department,
        'project': params.get('project', 'CFL'),
        'comments_per_post_limit': normalize_insight_comment_limit(
            params.get('comments_per_post_limit') or params.get('comment_limit'),
            unlimited=params.get('comment_limit_unlimited'),
        ),
        'comment_limit_unlimited': _truthy_form_value(params.get('comment_limit_unlimited')),
    }
    set_task_params(task_id, task_params)
    if not USE_DB_WORKER:
        threading.Thread(
            target=process_analysis_task,
            kwargs={
                'task_id': task_id,
                'urls': urls,
                'session_id': session_id,
                'user_id': user_id,
                'username': username,
                'department': department,
                'project': params.get('project', 'CFL'),
                'comments_per_post_limit': task_params['comments_per_post_limit'],
            },
            daemon=True,
        ).start()
    return {'tool_name': 'sentiment_comments', 'task_id': task_id, 'status': 'queued', 'url_count': len(urls)}


def _agent_create_video_metrics_task(params):
    urls = [u for u in (params.get('urls') or []) if video_metrics_etl.detect_platform(u) != 'UNKNOWN']
    if not urls:
        raise ValueError('未识别到可拉取基础数据的视频/帖子链接')
    user_id = session.get('user_id')
    session_id = session.get('session_id', 'default')
    queue_task_id = str(uuid.uuid4())
    df = pd.DataFrame({'视频链接': urls})
    buf = BytesIO()
    with pd.ExcelWriter(buf, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='视频链接')
    raw = buf.getvalue()
    parsed = etl_tools.parse_excel_urls(raw, '视频链接')
    input_row = db.execute_and_fetch_one(
        """
        INSERT INTO etl_file_outputs (task_id, user_id, filename, content)
        VALUES (%s, %s, %s, %s)
        RETURNING id
        """,
        (queue_task_id, user_id, '_input_video_metrics.xlsx', raw),
    )
    input_file_id = input_row['id'] if input_row else None
    if not input_file_id:
        raise RuntimeError('保存输入文件失败')
    selected_fields = params.get('selected_fields') or list(DEFAULT_VIDEO_METRIC_FIELDS)
    task_params = {
        'source': 'agent_video_metrics',
        'input_file_id': input_file_id,
        'url_column': '视频链接',
        'resolved_url_column': parsed.url_column,
        'sheet_name': parsed.sheet_name,
        'header_row': parsed.header_row,
        'urls': parsed.urls,
        'selected_fields': selected_fields,
        'user_id': user_id,
        'session_id': session_id,
    }
    create_task(queue_task_id, user_id, session_id, function_type='etl_video_metrics')
    set_task_params(queue_task_id, task_params)
    if not USE_DB_WORKER:
        threading.Thread(
            target=lambda: etl_jobs.run_etl_video_metrics_task(queue_task_id, task_params, update_task),
            daemon=True,
        ).start()
    return {'tool_name': 'video_metrics', 'task_id': queue_task_id, 'status': 'queued', 'url_count': len(urls)}


def _agent_create_profile_sync_task(params):
    if profile_video_scheduler.profile_video_sync_hard_disabled():
        raise ValueError('主页视频同步已被硬禁用，当前不允许创建主页抓取任务')
    profile_urls = [u for u in (params.get('urls') or []) if agent_service.is_profile_url(u)]
    if not profile_urls:
        raise ValueError('未识别到可同步的达人主页链接')
    sync_scope = params.get('sync_scope') or 'recent'
    if sync_scope not in {'recent', 'range', 'all'}:
        sync_scope = 'recent'
    start_date = params.get('start_date') or None
    end_date = params.get('end_date') or None
    if sync_scope == 'recent':
        start_date, end_date = agent_service.recent_window_dates(int(params.get('recent_days') or 7))
        sync_scope = 'range'
    max_videos = int(params.get('max_videos') or profile_video_scheduler.default_max_videos_per_profile())
    schedule_hour = int(params.get('schedule_hour') or profile_video_scheduler.default_sync_hour())
    config_ids = []
    inline_configs = []
    if params.get('schedule'):
        result = profile_video_scheduler.upsert_configs(
            profile_urls,
            user_id=session.get('user_id'),
            enabled=True,
            sync_scope=sync_scope,
            start_date=start_date,
            end_date=end_date,
            max_videos=max_videos,
            schedule_hour=schedule_hour,
            feishu_app_token=params.get('feishu_app_token') or None,
            feishu_table_id=params.get('feishu_table_id') or None,
        )
        config_ids = result.get('ids') or []
    else:
        inline_configs = [
            {
                'profile_url': url,
                'platform': video_metrics_etl.detect_platform(url),
                'sync_scope': sync_scope,
                'start_date': start_date,
                'end_date': end_date,
                'max_videos': max_videos,
                'schedule_hour': schedule_hour,
                'feishu_app_token': params.get('feishu_app_token') or None,
                'feishu_table_id': params.get('feishu_table_id') or None,
            }
            for url in profile_urls
        ]
    user_id = session.get('user_id')
    session_id = session.get('session_id', 'default')
    queue_task_id = str(uuid.uuid4())
    task_params = {
        'source': 'agent_profile_video_sync',
        'trigger_type': 'agent',
        'config_ids': config_ids,
        'profile_urls': [],
        'inline_configs': inline_configs,
        'user_id': user_id,
        'session_id': session_id,
    }
    create_task(queue_task_id, user_id, session_id, function_type='profile_video_sync')
    set_task_params(queue_task_id, task_params)
    if not USE_DB_WORKER:
        threading.Thread(
            target=lambda: profile_video_scheduler.run_profile_video_sync_task(queue_task_id, task_params, update_task),
            daemon=True,
        ).start()
    return {
        'tool_name': 'profile_video_sync',
        'task_id': queue_task_id,
        'status': 'queued',
        'profile_count': len(config_ids) + len(inline_configs),
        'config_ids': config_ids,
    }


def _agent_create_kol_link_task(params):
    urls = [u for u in (params.get('urls') or []) if video_metrics_etl.detect_platform(u) in {'TT', 'IG', 'YTB'}]
    if not urls:
        raise ValueError('未识别到 TikTok / Instagram / YouTube 达人主页链接')
    payload = {
        'text': '\n'.join(urls),
        'sync_to_pool': bool(params.get('sync_to_pool')),
        'include_acv': bool(params.get('include_acv', True)),
        'videos_per_profile': int(params.get('videos_per_profile') or 10),
    }
    response = requests.post(
        _agent_kol_api_url('/kols/data-refresh/link-task'),
        json=payload,
        headers=_agent_kol_headers(),
        timeout=30,
    )
    if response.status_code >= 400:
        raise RuntimeError(response.text[:500])
    data = response.json()
    return {'tool_name': 'kol_data_refresh_links', 'job_id': data.get('id'), 'status': data.get('status'), 'job': data}


def _agent_create_kol_excel_task(params):
    attachment_id = params.get('attachment_id')
    if not attachment_id:
        raise ValueError('请先上传 Excel 附件')
    row = db.query_one(
        "SELECT content, filename FROM etl_file_outputs WHERE id = %s AND user_id = %s",
        (attachment_id, session.get('user_id')),
    )
    if not row or not row.get('content'):
        raise ValueError('Excel 附件不存在或无权限访问')
    content = row['content']
    if isinstance(content, memoryview):
        content = content.tobytes()
    filename = row.get('filename') or 'agent_kol_data_refresh.xlsx'
    files = {
        'file': (
            filename,
            BytesIO(content),
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        )
    }
    query = {
        'sync_to_pool': str(bool(params.get('sync_to_pool'))).lower(),
        'include_acv': str(bool(params.get('include_acv', True))).lower(),
        'videos_per_profile': int(params.get('videos_per_profile') or 10),
    }
    response = requests.post(
        _agent_kol_api_url('/kols/data-refresh/excel-task'),
        params=query,
        files=files,
        headers=_agent_kol_headers(),
        timeout=60,
    )
    if response.status_code >= 400:
        raise RuntimeError(response.text[:500])
    data = response.json()
    return {'tool_name': 'kol_data_refresh_excel', 'job_id': data.get('id'), 'status': data.get('status'), 'job': data}


def _agent_list_recent_tasks():
    sentiment_items = []
    try:
        uid = session.get('user_id')
        rows = db.query_all(
            """
            SELECT task_id, function_type, status, progress, result, error, record_id,
                   created_at, updated_at, finished_at
            FROM task_queue
            WHERE (user_id = %s OR (user_id IS NULL AND function_type = 'feishu_profile_video_sync'))
              AND function_type IN ('sentiment', 'etl_video_metrics', 'profile_video_sync', 'feishu_profile_video_sync')
            ORDER BY created_at DESC
            LIMIT 20
            """,
            (uid,),
        ) or []
        for row in rows:
            result_payload = _agent_load_json(row.get('result'), {})
            sentiment_items.append({
                'source': 'sailson',
                'id': row.get('task_id'),
                'type': _sentiment_task_name(row.get('function_type')),
                'status': row.get('status'),
                'progress': row.get('progress'),
                'success_count': result_payload.get('success_count'),
                'failed_count': result_payload.get('failed_count'),
                'download_id': result_payload.get('download_id'),
                'error': row.get('error'),
                'created_at': row.get('created_at'),
            })
    except Exception as e:
        sentiment_items.append({'source': 'sailson', 'status': 'error', 'error': str(e)})

    kol_items = []
    if KOL_API_BASE_URL:
        try:
            response = requests.get(
                _agent_kol_api_url('/kols/data-refresh/jobs'),
                params={'limit': 10},
                headers=_agent_kol_headers(),
                timeout=15,
            )
            if response.status_code < 400:
                for job in response.json() or []:
                    kol_items.append({
                        'source': 'kol',
                        'id': job.get('id'),
                        'type': '达人数据更新',
                        'status': job.get('status'),
                        'progress': f"成功 {job.get('success_count', 0)}/{job.get('total', 0)}，失败 {job.get('failed_count', 0)}",
                        'success_count': job.get('success_count'),
                        'failed_count': job.get('failed_count'),
                        'download_url': f"/kol-api/kols/data-refresh/download/{job.get('id')}" if job.get('output_filename') else None,
                        'error': job.get('error'),
                        'created_at': job.get('created_at'),
                    })
        except Exception as e:
            kol_items.append({'source': 'kol', 'status': 'error', 'error': str(e)})
    return {'items': _json_safe_rows(sentiment_items + kol_items)}


def _agent_execute_action(row):
    intent = row.get('intent')
    params = _agent_load_json(row.get('params_json'), {})
    if intent not in AGENT_ACTION_EXECUTABLES:
        raise ValueError('该任务类型暂不支持执行')
    if intent == 'sentiment_comments':
        return _agent_create_sentiment_task(params)
    if intent == 'video_metrics':
        return _agent_create_video_metrics_task(params)
    if intent == 'profile_video_sync':
        return _agent_create_profile_sync_task(params)
    if intent == 'kol_data_refresh_links':
        return _agent_create_kol_link_task(params)
    if intent == 'kol_data_refresh_excel':
        return _agent_create_kol_excel_task(params)
    if intent == 'task_query':
        return _agent_list_recent_tasks()
    raise ValueError('无法执行该任务')


@app.route('/api/agent/message', methods=['POST'])
@login_required
def agent_message_api():
    try:
        user_message = ''
        session_id_chat = None
        attachment_id = None
        if request.files:
            user_message = (request.form.get('message') or '').strip()
            session_id_chat = request.form.get('session_id')
            upload = request.files.get('file')
            if upload and upload.filename:
                if not upload.filename.lower().endswith('.xlsx'):
                    return jsonify({'status': 'error', 'message': 'Agent 附件第一版仅支持 .xlsx'}), 400
                raw = upload.read()
                row = db.execute_and_fetch_one(
                    """
                    INSERT INTO etl_file_outputs (task_id, user_id, filename, content)
                    VALUES (%s, %s, %s, %s)
                    RETURNING id
                    """,
                    (f"agent_upload_{uuid.uuid4().hex}", session.get('user_id'), upload.filename, raw),
                )
                attachment_id = row['id'] if row else None
        else:
            data = request.get_json(silent=True) or {}
            user_message = (data.get('message') or '').strip()
            session_id_chat = data.get('session_id')
            attachment_id = data.get('attachment_id')
        if not user_message and not attachment_id:
            return jsonify({'status': 'error', 'message': '请输入任务需求或上传 Excel'}), 400
        if attachment_id and not user_message:
            user_message = '帮我用这个 Excel 做达人数据更新，默认只导出 Excel'

        chat_session_id = _agent_session(session_id_chat, user_message[:40] or 'AI 任务助手')
        _agent_insert_message(chat_session_id, 'user', user_message)
        draft = agent_service.build_draft(user_message, qwen_client=qwen_client)
        if attachment_id:
            draft.params['attachment_id'] = attachment_id
            if draft.intent in {'unknown', 'kol_data_refresh_links'}:
                draft.intent = 'kol_data_refresh_excel'
            draft.card = agent_service.build_action_card(draft.intent, draft.params)
            draft.needs_confirmation = True
            draft.reply = f"我识别到一个「{draft.card.get('task_type')}」任务，请确认后开始执行。"
        action_id = None
        action_result = None
        if draft.intent != 'unknown':
            action_id = _agent_create_action(chat_session_id, draft)
            if draft.intent == 'task_query':
                row = _agent_get_action(action_id)
                action_result = _agent_execute_action(row)
                _agent_update_action(
                    action_id,
                    status='executed',
                    tool_name='task_query',
                    result_json=_agent_json(action_result),
                )
                row = _agent_get_action(action_id)
                draft.card = _agent_load_json(row.get('card_json'), draft.card)
        _agent_insert_message(chat_session_id, 'assistant', draft.reply)
        return jsonify({
            'status': 'success',
            'session_id': chat_session_id,
            'reply': draft.reply,
            'intent': draft.intent,
            'needs_confirmation': draft.needs_confirmation,
            'action': {
                'id': action_id,
                'intent': draft.intent,
                'status': 'executed' if action_result else ('draft' if draft.needs_confirmation else 'ready'),
                'card': draft.card,
                'params': draft.params,
                'result': action_result,
            } if action_id else None,
        })
    except Exception as e:
        logger.error(f"Agent message failed: {e}")
        return jsonify({'status': 'error', 'message': str(e)}), 500


@app.route('/api/agent/actions/<int:action_id>/confirm', methods=['POST'])
@login_required
def agent_confirm_action_api(action_id):
    row = _agent_get_action(action_id)
    if not row:
        return jsonify({'status': 'error', 'message': '任务动作不存在或无权限访问'}), 404
    if row.get('status') in {'executing', 'executed'}:
        return jsonify({'status': 'success', 'action': _agent_action_out(row)})
    try:
        _agent_update_action(action_id, status='executing', confirmed_at=datetime.datetime.now())
        fresh = _agent_get_action(action_id)
        result = _agent_execute_action(fresh)
        tool_task_id = result.get('task_id')
        tool_job_id = result.get('job_id')
        _agent_update_action(
            action_id,
            status='executed',
            tool_name=result.get('tool_name') or fresh.get('intent'),
            tool_task_id=str(tool_task_id) if tool_task_id else None,
            tool_job_id=str(tool_job_id) if tool_job_id else None,
            result_json=_agent_json(result),
        )
        return jsonify({'status': 'success', 'action': _agent_action_out(_agent_get_action(action_id))})
    except Exception as e:
        logger.error(f"Agent confirm failed: {e}")
        _agent_update_action(action_id, status='failed', error=str(e)[:1000])
        return jsonify({'status': 'error', 'message': str(e), 'action': _agent_action_out(_agent_get_action(action_id))}), 500


@app.route('/api/agent/actions/<int:action_id>')
@login_required
def agent_action_api(action_id):
    row = _agent_get_action(action_id)
    if not row:
        return jsonify({'status': 'error', 'message': '任务动作不存在或无权限访问'}), 404
    return jsonify({'status': 'success', 'action': _agent_action_out(row)})


@app.route('/api/agent/tasks')
@login_required
def agent_tasks_api():
    return jsonify({'status': 'success', **_agent_list_recent_tasks()})


@app.route('/api/agent/sessions')
@login_required
def agent_sessions_api():
    rows = db.query_all(
        """
        SELECT id, title, created_at
        FROM chat_sessions
        WHERE user_id = %s AND mode = 'agent'
        ORDER BY created_at DESC
        LIMIT 30
        """,
        (session.get('user_id'),),
    ) or []
    return jsonify({'status': 'success', 'items': _json_safe_rows(rows)})


@app.route('/api/agent/sessions/<int:session_id_chat>')
@login_required
def agent_session_detail_api(session_id_chat):
    row = db.query_one(
        "SELECT id, title, created_at FROM chat_sessions WHERE id = %s AND user_id = %s AND mode = 'agent'",
        (session_id_chat, session.get('user_id')),
    )
    if not row:
        return jsonify({'status': 'error', 'message': '会话不存在'}), 404
    messages = db.query_all(
        "SELECT role, content, created_at FROM chat_messages WHERE session_id = %s ORDER BY created_at ASC",
        (session_id_chat,),
    ) or []
    actions = db.query_all(
        "SELECT * FROM agent_actions WHERE chat_session_id = %s AND user_id = %s ORDER BY created_at ASC",
        (session_id_chat, session.get('user_id')),
    ) or []
    return jsonify({
        'status': 'success',
        'session': _json_safe(dict(row)),
        'messages': _json_safe_rows(messages),
        'actions': [_agent_action_out(dict(a)) for a in actions],
    })


@app.route('/dashboard_stats')
@login_required
def dashboard_stats():
    """首页数据看板 API。

    业务口径：
    - Total Reviews：每条被处理的数据算一条，优先用 usage_events.item_count，老数据回退 usage_logs.comments_count。
    - Total Analyses：每次任务算一次，优先用 task_queue，老数据回退 usage_logs。
    - AI Tokens：只有触发 AI 分析且有 token 记录的任务才累计。
    """
    try:
        now = datetime.datetime.now()
        current_month = now.strftime('%Y-%m')
        last_month = (now.replace(day=1) - datetime.timedelta(days=1)).strftime('%Y-%m')

        first_usage_event = db.query_one("SELECT MIN(created_at) AS first_at FROM usage_events") or {}
        first_event_at = first_usage_event.get('first_at')

        usage_total = db.query_one("""
            SELECT COALESCE(SUM(item_count), 0) AS reviews,
                   COALESCE(SUM(ai_tokens), 0) AS tokens,
                   COUNT(*) AS events
            FROM usage_events
        """) or {}
        usage_current = db.query_one("""
            SELECT COALESCE(SUM(item_count), 0) AS reviews,
                   COALESCE(SUM(ai_tokens), 0) AS tokens,
                   COUNT(*) AS events
            FROM usage_events
            WHERE TO_CHAR(created_at, 'YYYY-MM') = %s
        """, (current_month,)) or {}
        usage_last = db.query_one("""
            SELECT COALESCE(SUM(item_count), 0) AS reviews,
                   COALESCE(SUM(ai_tokens), 0) AS tokens,
                   COUNT(*) AS events
            FROM usage_events
            WHERE TO_CHAR(created_at, 'YYYY-MM') = %s
        """, (last_month,)) or {}

        legacy_where = ""
        legacy_params = []
        if first_event_at:
            legacy_where = "WHERE created_at < %s"
            legacy_params.append(first_event_at)
        legacy_total = db.query_one(f"""
            SELECT COALESCE(SUM(comments_count), 0) AS reviews,
                   COALESCE(SUM(ai_tokens), 0) AS tokens,
                   COUNT(*) AS events
            FROM usage_logs
            {legacy_where}
        """, tuple(legacy_params)) or {}
        legacy_current = db.query_one(f"""
            SELECT COALESCE(SUM(comments_count), 0) AS reviews,
                   COALESCE(SUM(ai_tokens), 0) AS tokens,
                   COUNT(*) AS events
            FROM usage_logs
            WHERE TO_CHAR(created_at, 'YYYY-MM') = %s
            {('AND created_at < %s' if first_event_at else '')}
        """, tuple([current_month] + ([first_event_at] if first_event_at else []))) or {}
        legacy_last = db.query_one(f"""
            SELECT COALESCE(SUM(comments_count), 0) AS reviews,
                   COALESCE(SUM(ai_tokens), 0) AS tokens,
                   COUNT(*) AS events
            FROM usage_logs
            WHERE TO_CHAR(created_at, 'YYYY-MM') = %s
            {('AND created_at < %s' if first_event_at else '')}
        """, tuple([last_month] + ([first_event_at] if first_event_at else []))) or {}

        task_total = db.query_one("SELECT COUNT(*) AS count FROM task_queue") or {}
        task_current = db.query_one("""
            SELECT COUNT(*) AS count FROM task_queue
            WHERE TO_CHAR(created_at, 'YYYY-MM') = %s
        """, (current_month,)) or {}
        task_last = db.query_one("""
            SELECT COUNT(*) AS count FROM task_queue
            WHERE TO_CHAR(created_at, 'YYYY-MM') = %s
        """, (last_month,)) or {}
        queue = db.query_one("""
            SELECT
              COUNT(*) FILTER (WHERE status = 'pending') AS pending,
              COUNT(*) FILTER (WHERE status IN ('claimed', 'processing')) AS running,
              COUNT(*) FILTER (WHERE status = 'failed') AS failed
            FROM task_queue
            WHERE created_at >= NOW() - INTERVAL '7 days'
        """) or {}

        total_reviews = int(usage_total.get('reviews') or 0) + int(legacy_total.get('reviews') or 0)
        current_reviews = int(usage_current.get('reviews') or 0) + int(legacy_current.get('reviews') or 0)
        last_reviews = int(usage_last.get('reviews') or 0) + int(legacy_last.get('reviews') or 0)
        total_tokens = int(usage_total.get('tokens') or 0) + int(legacy_total.get('tokens') or 0)
        current_tokens = int(usage_current.get('tokens') or 0) + int(legacy_current.get('tokens') or 0)

        task_count = int(task_total.get('count') or 0)
        fallback_events = int(usage_total.get('events') or 0) + int(legacy_total.get('events') or 0)
        total_analyses = task_count if task_count else fallback_events
        current_analyses = int(task_current.get('count') or 0) if task_count else int(usage_current.get('events') or 0) + int(legacy_current.get('events') or 0)
        last_analyses = int(task_last.get('count') or 0) if task_count else int(usage_last.get('events') or 0) + int(legacy_last.get('events') or 0)

        growth = _percent_change(current_reviews, last_reviews)
        analyses_growth = _percent_change(current_analyses, last_analyses)
        running = int(queue.get('running') or 0)
        pending = int(queue.get('pending') or 0)
        failed = int(queue.get('failed') or 0)

        return jsonify({
            'comments': total_reviews,
            'analyses': total_analyses,
            'tokens': total_tokens,
            'growth': growth,
            'current_month': {
                'comments': current_reviews,
                'analyses': current_analyses,
                'tokens': current_tokens,
            },
            'queue': {
                'pending': pending,
                'running': running,
                'failed': failed,
            },
            'badges': {
                'reviews': _growth_badge(growth),
                'analyses': _analyses_badge(running, pending, failed, analyses_growth),
                'tokens': _tokens_badge(current_tokens),
            }
        })

    except Exception as e:
        logger.error(f"❌ 获取数据看板失败: {e}")
        return jsonify({
            'comments': 0,
            'analyses': 0,
            'tokens': 0,
            'growth': 0,
            'badges': {
                'reviews': {'text': '0%', 'style': 'neutral'},
                'analyses': {'text': 'Idle', 'style': 'neutral'},
                'tokens': {'text': 'No AI', 'style': 'neutral'},
            }
        })


def _percent_change(current_value, last_value):
    if not last_value:
        return 100.0 if current_value else 0.0
    return round(((current_value - last_value) / last_value) * 100, 1)


def _growth_badge(growth):
    if growth > 0:
        return {'text': f'+{growth:g}%', 'style': 'positive'}
    if growth < 0:
        return {'text': f'{growth:g}%', 'style': 'negative'}
    return {'text': '0%', 'style': 'neutral'}


def _analyses_badge(running, pending, failed, growth):
    if running:
        return {'text': f'Running {running}', 'style': 'active'}
    if pending:
        return {'text': f'Queued {pending}', 'style': 'warning'}
    if failed:
        return {'text': f'Failed {failed}', 'style': 'negative'}
    if growth > 0:
        return {'text': f'+{growth:g}%', 'style': 'positive'}
    return {'text': 'Idle', 'style': 'neutral'}


def _tokens_badge(current_tokens):
    if current_tokens > 0:
        return {'text': 'AI Used', 'style': 'active'}
    return {'text': 'No AI', 'style': 'neutral'}


@app.route('/api/usage/summary')
@login_required
def usage_summary_api():
    """统一消耗汇总。普通用户默认只看自己；管理员可传 all_users=1。"""
    try:
        include_estimated = request.args.get('include_estimated', '1') != '0'
        month = request.args.get('month') or None
        all_users = request.args.get('all_users') == '1' and session.get('role') == 'admin'
        user_id = None if all_users else session.get('user_id')
        rows = usage_service.get_usage_summary(user_id=user_id, include_estimated=include_estimated, month=month)
        totals = {
            'events': sum(int(row.get('events') or 0) for row in rows),
            'item_count': sum(int(row.get('item_count') or 0) for row in rows),
            'crawler_items': sum(int(row.get('crawler_items') or 0) for row in rows),
            'ai_tokens': sum(int(row.get('ai_tokens') or 0) for row in rows),
            'api_calls': sum(int(row.get('api_calls') or 0) for row in rows),
            'crawler_cost_usd': round(sum(float(row.get('crawler_cost_usd') or 0) for row in rows), 4),
            'crawler_cost_cny': round(sum(float(row.get('crawler_cost_cny') or 0) for row in rows), 4),
            'ai_cost_cny': round(sum(float(row.get('ai_cost_cny') or 0) for row in rows), 4),
            'total_cost_cny': round(sum(float(row.get('total_cost_cny') or 0) for row in rows), 4),
        }
        return jsonify({
            'pricing': {
                'crawler': 'USD 3 / 1000 crawler rows',
                'ai': 'CNY 0.008 / 1000 tokens',
                'official_api': 'tracked as api_calls, crawler cost not applied',
            },
            'totals': totals,
            'rows': rows,
        })
    except Exception as e:
        logger.error(f"❌ 获取消耗汇总失败: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/usage/events')
@login_required
def usage_events_api():
    """统一消耗明细。"""
    try:
        include_estimated = request.args.get('include_estimated', '1') != '0'
        all_users = request.args.get('all_users') == '1' and session.get('role') == 'admin'
        user_id = None if all_users else session.get('user_id')
        limit = int(request.args.get('limit') or 100)
        return jsonify({
            'rows': usage_service.get_usage_events(user_id=user_id, include_estimated=include_estimated, limit=limit)
        })
    except Exception as e:
        logger.error(f"❌ 获取消耗明细失败: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/usage/history-estimates')
@login_required
def usage_history_estimates_api():
    """历史消耗估算，不写库。用于说明哪些历史账只能回算。"""
    try:
        limit = int(request.args.get('limit') or 200)
        rows = usage_service.estimate_history(limit=limit)
        if session.get('role') != 'admin':
            uid = session.get('user_id')
            rows = [row for row in rows if row.get('user_id') == uid]
        return jsonify({
            'note': 'source=recorded_legacy 来自旧 usage_logs；source=estimated 来自历史结果/任务参数回算，无法代表实际 Apify 账单。',
            'pricing': {'crawler': 'USD 3 / 1000 crawler rows'},
            'rows': rows,
        })
    except Exception as e:
        logger.error(f"❌ 获取历史消耗估算失败: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/debug')
@login_required
def debug_page():
    """调试页面"""
    debug_info = {
        "status": "Online",
        "qwen_key": bool(DASHSCOPE_API_KEY),
        "apify_key": bool(APIFY_TOKEN),
        "port": PORT,
        "history_count": len(HISTORY_DB)
    }
    logger.info(f"🔍 调试信息: {debug_info}")
    return jsonify(debug_info)


@app.route('/health')
def health_check():
    """健康检查端点 - 用于 Render 监控"""
    return jsonify({"status": "ok", "service": "Sailson AI"}), 200


@app.route('/submit_feedback', methods=['POST'])
def submit_feedback():
    """接收用户反馈并发送邮件"""
    try:
        data = request.json
        project_name = data.get('project_name')
        feedback = data.get('feedback')

        if not project_name or not feedback:
            return jsonify({'error': '请填写完整信息'}), 400

        # 记录到日志
        logger.info(f"📧 收到用户反馈")
        logger.info(f"   项目名称: {project_name}")
        logger.info(f"   反馈内容: {feedback}")

        # 保存到数据库（可选）
        try:
            db.execute("""
                INSERT INTO feedback (user_email, content, created_at)
                VALUES (%s, %s, NOW())
            """, (project_name, feedback))
        except Exception as db_error:
            # 如果表不存在，只记录日志
            logger.warning(f"⚠️ 保存反馈到数据库失败（表可能不存在）: {db_error}")

        # 发送邮件通知管理员（如果配置了 SMTP）
        email_sent = send_feedback_email(project_name, feedback)

        msg = '感谢您的反馈！'
        if not email_sent:
            # 不打扰用户，只在日志中记录邮件失败
            logger.warning("⚠️ 反馈已保存，但邮件通知未成功发送")

        return jsonify({'success': True, 'message': msg})

    except Exception as e:
        logger.error(f"❌ 处理反馈失败: {e}")
        return jsonify({'error': '系统错误，请稍后重试'}), 500


# ============================================
# 功能 1: 舆情分析
# ============================================

@app.route('/sentiment-tool')
@login_required
def sentiment_tool():
    """舆情分析工具页面"""
    user_id = session.get('user_id')

    has_used_sentiment = False
    if user_id:
        try:
            row = db.query_one(
                """
                SELECT 1 FROM analysis_results
                WHERE user_id = %s AND type = %s
                LIMIT 1
                """,
                (user_id, 'sentiment')
            )
            has_used_sentiment = bool(row)
        except Exception as e:
            logger.error(f"❌ 检查舆情历史记录失败: {e}")

    return render_template(
        'analysis.html',
        has_used_sentiment=has_used_sentiment,
        video_metrics_batch_size=video_metrics_etl.BATCH_SIZE,
        video_metrics_max_upload_bytes=ETL_VIDEO_METRICS_MAX_UPLOAD_BYTES,
        video_metrics_max_upload_mb=round(ETL_VIDEO_METRICS_MAX_UPLOAD_BYTES / 1024 / 1024, 1),
        profile_video_max_profiles=ETL_PROFILE_VIDEO_MAX_PROFILES,
        profile_video_max_per_profile=ETL_PROFILE_VIDEO_MAX_PER_PROFILE,
        insight_comments_default=sentiment_insight.DEFAULT_COMMENTS_PER_POST_LIMIT,
        insight_comments_unlimited=sentiment_insight.UNLIMITED_COMMENTS_PER_POST_LIMIT,
    )


from thai_utils import (
    is_thai_content, has_thai_chars,
    contains_any_tag_or_term, normalize_tag_token,
    thai_datasets_config, thai_matching_datasets,
    MLBB_DISCOVER_KEYWORDS, SPD_KEYWORDS, ROV_DISCOVER_KEYWORDS,
    THAI_GAME_TAGS,
)

# 布尔规则：后端 _eval_boolean_expr 匹配（小写归一），非 AI。
TASK_BOOLEAN_RULES = {
    "MLBB": (
        '(#mlbb OR #mlbb* OR "mlbb" OR #mobilelegends OR "mobile legends" OR "mobilelegendsbangbang" '
        'OR #mobilelegendsbangbang OR #mobalegends OR #moba55) '
        'AND NOT ("freefire" OR #freefire OR #garenafreefire OR #freefirebgid OR #mlb OR #mlbbaseball OR #mlbbets)'
    ),
    "SPD": (
        '((#mlbb OR #mlbb* OR "mlbb" OR #mobilelegends OR "mobile legends" OR "mobilelegendsbangbang" '
        'OR #mobilelegendsbangbang OR #mobalegends OR #moba55) '
        'AND NOT ("freefire" OR #freefire OR #garenafreefire OR #freefirebgid OR #mlb OR #mlbbaseball OR #mlbbets) '
        'AND ("naruto" OR "itachi" OR "minato" OR "gusion" OR "julian" OR "valir" OR "sasuke" OR "sakura" OR "kakashi" '
        'OR "gaara" OR "madara" OR "lukas" OR "kalea" OR "suyou" OR "hayabusa" OR "vale" OR "uchiha" OR "火影" OR "疾风传" OR #naruto)) '
        'OR (#mlbbxnaruto OR "mlbb naruto" OR #mlbbnewskin OR #mlbbfreegaaraskin OR #valirfreeskin OR #goldsongkran OR #goldhunt)'
    ),
    "ROV": (
        '("rov" OR #rov OR "realm of valor" OR "garena rov" OR "garena realm of valor" OR "อาโอวี" OR '
        '#realmofvalor OR #garenarov OR #rovthailand OR #rovvietnam OR "#อาโอวี" OR '
        '"garena rov thailand" OR #garenarovthailand OR #ambassadorofvalor OR '
        '"#อยากตีป้อมโว้ย" OR "#rovlnwสาด" OR "#4เมษาlnwมาตีป้อม" OR #missrov2026 OR #missrovtournament OR #rovtourbillion) '
        'AND NOT ("freefire" OR #freefire OR #garenafreefire OR #freefirebgid OR #hok OR "hok" OR #mlbb OR "mlbb" '
        'OR #mobilelegends OR "mobilelegends" OR "mobilelegendsbangbang" OR #mobilelegendsbangbang)'
    ),
}


def _normalize_monitor_text(content, brief_analysis):
    merged = f"{content or ''} {brief_analysis or ''}".lower()
    merged = re.sub(r'\s+', ' ', merged).strip()
    return merged


def _tokenize_boolean_expr(expr):
    # 支持括号、引号短语、AND/OR/NOT、普通词
    token_re = re.compile(r'"[^"]*"|\(|\)|\bAND\b|\bOR\b|\bNOT\b|[^\s()]+', flags=re.IGNORECASE)
    return token_re.findall(expr or "")


def _boolean_to_rpn(tokens):
    precedence = {"NOT": 3, "AND": 2, "OR": 1}
    output = []
    ops = []
    for token in tokens:
        upper = token.upper()
        if upper in precedence:
            while ops and ops[-1] != "(" and precedence.get(ops[-1], 0) >= precedence[upper]:
                output.append(ops.pop())
            ops.append(upper)
        elif token == "(":
            ops.append(token)
        elif token == ")":
            while ops and ops[-1] != "(":
                output.append(ops.pop())
            if ops and ops[-1] == "(":
                ops.pop()
        else:
            output.append(token)
    while ops:
        output.append(ops.pop())
    return output


def _match_boolean_term(text, raw_term):
    term = (raw_term or "").strip()
    if not term:
        return False
    if term.startswith('"') and term.endswith('"'):
        term = term[1:-1].strip()
    term = term.lower()
    if not term:
        return False

    # 通配后缀：mlbb* / #mlbb*
    if term.endswith('*'):
        prefix = term[:-1]
        if not prefix:
            return False
        return prefix in text
    return term in text


def _eval_boolean_expr(text, expr):
    tokens = _tokenize_boolean_expr(expr)
    if not tokens:
        return False
    rpn = _boolean_to_rpn(tokens)
    stack = []
    for token in rpn:
        if token in {"AND", "OR", "NOT"}:
            if token == "NOT":
                value = stack.pop() if stack else False
                stack.append(not value)
            else:
                right = stack.pop() if stack else False
                left = stack.pop() if stack else False
                stack.append(left and right if token == "AND" else left or right)
        else:
            stack.append(_match_boolean_term(text, token))
    return bool(stack[-1]) if stack else False


def _match_task(task_name, content, brief_analysis):
    rule = TASK_BOOLEAN_RULES.get(task_name)
    if not rule:
        return False
    text = _normalize_monitor_text(content, brief_analysis)
    return _eval_boolean_expr(text, rule)


def _sentiment_bucket(score):
    if score is None:
        return 'neutral'
    if score > 0.3:
        return 'positive'
    if score < -0.2:
        return 'negative'
    return 'neutral'


def _region_by_language(language_code):
    code = (language_code or '').lower()
    if code.startswith('th'):
        return 'TH'
    if code.startswith('id'):
        return 'ID'
    if code.startswith('pt'):
        return 'PT'
    return 'EN'


def _to_percentage(value, total):
    if total <= 0:
        return 0.0
    return round((value / total) * 100, 1)


def _extract_opinion_text(row):
    """从单条评论中提炼观点文本（不依赖旧 category 枚举）。"""
    def _clean(raw):
        text = re.sub(r'\s+', ' ', (raw or '')).strip()
        if not text:
            return ''
        # 去掉纯 hashtag/@mention/emoji/活动口令类噪声
        if re.fullmatch(r'[#@\w\W]{0,3}', text):
            return ''
        if re.fullmatch(r'(#\w+\s*){1,6}', text, flags=re.IGNORECASE):
            return ''
        if re.fullmatch(r'([@#]\w+\s*){1,8}', text):
            return ''
        return text

    def _is_low_value(text):
        t = (text or '').lower()
        if not t:
            return True
        low_value_patterns = [
            '该评论仅为话题标签', '仅为话题标签', '仅为标签', '无实质内容',
            '无明确情感倾向', '属于中性宣传', '联动期待类文本', '参与活动',
            'tag friends', 'comment done', 'done ✅'
        ]
        if any(p in t for p in low_value_patterns):
            return True
        # 纯短词/标点内容
        compact = re.sub(r'[\W_]+', '', t)
        return len(compact) < 4

    # 优先 brief_analysis；若被判低价值则回退原文
    brief_text = _clean(row.get('_rt_brief') or row.get('brief_analysis'))
    if brief_text and not _is_low_value(brief_text):
        return brief_text[:80]

    content_text = _clean(row.get('content'))
    if content_text and not _is_low_value(content_text):
        return content_text[:80]

    return ""


def _opinion_signature(text):
    normalized = (text or "").lower()
    normalized = re.sub(r'[^\w\u4e00-\u9fff]+', ' ', normalized)
    normalized = re.sub(r'\s+', ' ', normalized).strip()
    if not normalized:
        return set()

    tokens = set()
    for token in normalized.split(' '):
        if len(token) >= 2:
            tokens.add(token)

    # 增加中文二元组，提升“语义相近短句”聚类命中
    zh = ''.join(ch for ch in normalized if '\u4e00' <= ch <= '\u9fff')
    for i in range(len(zh) - 1):
        tokens.add(zh[i:i + 2])
    return tokens


def _cluster_opinions(rows, top_k=3):
    """对同一情感下评论做轻量聚类，返回 top 观点簇。"""
    clusters = []
    for row in rows:
        opinion_text = _extract_opinion_text(row)
        if not opinion_text:
            continue
        sig = _opinion_signature(opinion_text)
        best_idx = -1
        best_score = 0.0
        for i, cluster in enumerate(clusters):
            inter = len(sig & cluster['signature'])
            union = len(sig | cluster['signature']) or 1
            score = inter / union
            if score > best_score:
                best_score = score
                best_idx = i

        if best_idx >= 0 and best_score >= 0.32:
            c = clusters[best_idx]
            c['rows'].append(row)
            c['signature'] = c['signature'] | sig
            if len(opinion_text) < len(c['opinion']):
                c['opinion'] = opinion_text
        else:
            clusters.append({
                'opinion': opinion_text,
                'signature': sig,
                'rows': [row]
            })

    clusters.sort(key=lambda x: len(x['rows']), reverse=True)
    return clusters[:top_k]


def _is_spd_relaxed_text(content, brief_analysis):
    text = _normalize_monitor_text(content, brief_analysis)
    if not text:
        return False
    relaxed_terms = [
        "naruto", "火影", "疾风传", "itachi", "julian", "valir",
        "ember gaze", "mlbbxnaruto", "#mlbbxnaruto", "uchiha", "晓组织"
    ]
    return any(term in text for term in relaxed_terms)


def _is_mlbb_relaxed_text(content, brief_analysis):
    text = _normalize_monitor_text(content, brief_analysis)
    if not text:
        return False
    mlbb_terms = [
        "mlbb", "#mlbb", "mobile legends", "mobilelegends",
        "mobilelegendsbangbang", "#mobilelegends", "#mobilelegendsbangbang", "#moba55"
    ]
    return any(term in text for term in mlbb_terms)


def _contains_zh(text):
    return bool(re.search(r'[\u4e00-\u9fff]', (text or '')))


def _translate_text_to_zh(text):
    source = (text or '').strip()
    if not source:
        return ''
    if _contains_zh(source):
        return source
    # 报告接口默认不做在线翻译，避免超时导致整页空白
    if os.environ.get("SPD_REPORT_ENABLE_TRANSLATE", "false").lower() != "true":
        return source
    if not qwen_client:
        return source
    try:
        prompt = (
            "把下面这句社媒评论翻译成简体中文，只返回翻译结果，不要解释：\n"
            f"{source}"
        )
        response = qwen_client.chat.completions.create(
            model="qwen-plus",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.1
        )
        translated = (response.choices[0].message.content or '').strip()
        return translated or source
    except Exception:
        return source


def _reanalyze_comments_for_template(rows, force_full=False):
    """仅用于 SPD 报告：对 Top5 评论池做重分析。"""
    if not rows:
        return []
    try:
        reanalyze_limit = int(os.environ.get("SPD_REPORT_REANALYZE_LIMIT", "0"))
    except Exception:
        reanalyze_limit = 0
    if reanalyze_limit < 0:
        reanalyze_limit = 0

    result = []
    reanalyzed = 0
    for row in rows:
        new_row = dict(row)
        content = (new_row.get('content') or '').strip()
        # 默认先复用历史结果，确保任何情况下都能返回报告
        new_row['_rt_sentiment'] = new_row.get('sentiment') or 'neutral'
        new_row['_rt_brief'] = new_row.get('brief_analysis') or ''

        if not content:
            result.append(new_row)
            continue

        if reanalyzed >= reanalyze_limit:
            result.append(new_row)
            continue

        # 非强制全量模式下，若已有分析结果则直接复用
        if (not force_full) and new_row.get('brief_analysis') and (new_row.get('sentiment') in {'positive', 'neutral', 'negative'}):
            result.append(new_row)
            continue

        try:
            score, _, language, brief = tasks.analyze_comment_sentiment(content)
            new_row['_rt_sentiment'] = _sentiment_bucket(score)
            new_row['_rt_language'] = (language or new_row.get('language') or '').lower()
            brief_text = (brief or '').strip()
            new_row['_rt_brief'] = brief_text if brief_text else (new_row.get('brief_analysis') or '')
            reanalyzed += 1
        except Exception:
            pass
        result.append(new_row)
    if len(rows) > reanalyze_limit:
        logger.info(f"ℹ️ SPD report reanalyze capped: {reanalyzed}/{len(rows)} | force_full={force_full}")
    return result


def _rewrite_opinion_for_template(sentiment, region, raw_opinion, examples=None):
    """
    将聚类观点改写为模板风格一句话：
    {区域}：{一句话可执行观点}（不在此函数内加百分比）
    """
    opinion = (raw_opinion or '').strip()
    if not opinion:
        return ''
    # 报告接口默认不做在线改写，避免超时导致整页空白
    if os.environ.get("SPD_REPORT_ENABLE_REWRITE", "false").lower() != "true":
        return opinion[:42]
    if not qwen_client:
        return opinion[:42]
    sample_text = "；".join([((x or {}).get('translation') or (x or {}).get('original') or '')[:80] for x in (examples or [])[:2]])
    try:
        prompt = f"""你是手游舆情分析师。请把下面观点改写成“可交付给品牌方”的一句话中文结论。

情感：{sentiment}
区域：{region}
原始观点：{opinion}
示例评论：{sample_text or '无'}

要求：
1) 只输出一句中文，不超过38个字；
2) 必须包含具体对象（如角色/皮肤/机制/福利/活动规则之一）；
3) 禁止出现“该评论仅为话题标签/无实义”这类空话；
4) 不要输出百分比、不要解释、不要编号。
"""
        resp = qwen_client.chat.completions.create(
            model="qwen-plus",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2
        )
        text = (resp.choices[0].message.content or '').strip()
        text = re.sub(r'\s+', ' ', text).strip()
        if not text:
            return opinion[:42]
        return text[:42]
    except Exception:
        return opinion[:42]


def _ai_aggregate_opinions(comment_pool, total_comments):
    """
    用 LLM 对 Top5 评论池做聚合舆情提炼 + 情感占比估算。
    返回 dict: {'opinions': [{sentiment, region, opinion, percentage}], 'sentiment_ratio': {positive, neutral, negative}}
    返回 None 表示失败（触发降级）；返回空 dict 的 opinions 为 [] 时同理。
    """
    if not comment_pool:
        return {'opinions': [], 'sentiment_ratio': {'positive': 0, 'neutral': 0, 'negative': 0}}
    if not qwen_client:
        return None

    region_comments = defaultdict(list)
    for row in comment_pool:
        region = row.get('region') or 'EN'
        brief = (row.get('_rt_brief') or row.get('brief_analysis') or '').strip()
        content = (row.get('content') or '').strip()
        text = brief if brief else content[:150]
        if text:
            region_comments[region].append(text)

    lines = []
    for region, texts in region_comments.items():
        sampled = texts[:100]
        lines.append(f"【{region}】共{len(texts)}条")
        for t in sampled:
            lines.append(f"  - {t[:120]}")

    if not lines:
        return None

    prompt = f"""你是资深手游舆情分析师。以下是某游戏联动活动 TOP5 热门帖子的全部评论（按地区分组）。

请完成两项任务：
1. 估算整体情感分布（正面/中性/负面各占百分比，三者之和=100）
2. 对正面、中性、负面各提炼 3-5 条最核心的舆情观点，标注来源地区

评论数据：
{chr(10).join(lines[:600])}

总评论量：{total_comments}

严格按以下 JSON 格式输出，不要输出任何其他文字：
{{
  "sentiment_ratio": {{"positive": 正面百分比数字, "neutral": 中性百分比数字, "negative": 负面百分比数字}},
  "positive": [
    {{"region": "地区代码如EN/TH/ID/PT", "opinion": "一句话中文观点（不超过40字）", "count_estimate": 估计提到该观点的评论数}}
  ],
  "neutral": [...],
  "negative": [...]
}}

要求：
1. 观点必须具体，包含明确的游戏元素（角色/皮肤/活动/机制等）
2. 禁止空泛描述如"玩家讨论了游戏"
3. 每种情绪 3-5 条，按讨论热度排序
4. count_estimate 是你估算该观点涉及的评论条数
5. sentiment_ratio 三项之和必须等于100"""

    try:
        resp = qwen_client.chat.completions.create(
            model="qwen-plus",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.15,
            timeout=30
        )
        raw = (resp.choices[0].message.content or '').strip()
        match = re.search(r'\{[\s\S]*\}', raw)
        if not match:
            return None
        data = json.loads(match.group())

        opinions = []
        for sentiment in ['positive', 'neutral', 'negative']:
            items = data.get(sentiment, [])
            for item in items[:5]:
                try:
                    count_est = int(re.sub(r'[^\d]', '', str(item.get('count_estimate', 0))) or 0)
                except (ValueError, TypeError):
                    count_est = 0
                opinions.append({
                    'sentiment': sentiment,
                    'region': item.get('region', 'EN'),
                    'opinion': (item.get('opinion', '') or '')[:60],
                    'percentage': _to_percentage(count_est, total_comments) if count_est > 0 else 0.0
                })

        raw_ratio = data.get('sentiment_ratio', {})
        try:
            ratio = {
                'positive': round(float(raw_ratio.get('positive', 0)), 1),
                'neutral': round(float(raw_ratio.get('neutral', 0)), 1),
                'negative': round(float(raw_ratio.get('negative', 0)), 1),
            }
        except (ValueError, TypeError):
            ratio = None

        if opinions:
            return {'opinions': opinions, 'sentiment_ratio': ratio}
        return None
    except Exception as e:
        logger.warning(f"⚠️ AI aggregate opinions failed, falling back to heuristic: {e}")
        return None


def _ai_analyze_single_post(comments, total_post_comments):
    """
    用 LLM 对单帖评论做正面/中性/负面观点提炼 + 情感占比估算。
    返回 dict: {'sentiments': [...], 'sentiment_ratio': {...}} 或 None。
    """
    if not comments or not qwen_client:
        return None

    lines = []
    for row in comments:
        region = row.get('region') or 'EN'
        brief = (row.get('_rt_brief') or row.get('brief_analysis') or '').strip()
        content = (row.get('content') or '').strip()
        text = brief if brief else content[:150]
        if text:
            lines.append(f"[{region}] {text[:120]}")

    if not lines:
        return None

    prompt = f"""你是资深手游舆情分析师。以下是某帖子的全部评论（标注了地区）。

请完成两项任务：
1. 估算该帖评论的情感分布（正面/中性/负面各占百分比，之和=100）
2. 对正面、中性、负面各提炼 Top3 主要观点，标注来源地区

评论数据：
{chr(10).join(lines[:400])}

总评论数：{total_post_comments}

严格按以下 JSON 格式输出，不要输出任何其他文字：
{{
  "sentiment_ratio": {{"positive": 正面百分比, "neutral": 中性百分比, "negative": 负面百分比}},
  "positive": [
    {{"region": "地区代码", "opinion": "一句话中文观点（不超过35字）", "count_estimate": 估计数}}
  ],
  "neutral": [...],
  "negative": [...]
}}

要求：
1. 每种情绪最多3条，按热度排序
2. 观点要具体，含游戏元素
3. 禁止空泛描述
4. sentiment_ratio 三项之和=100"""

    try:
        resp = qwen_client.chat.completions.create(
            model="qwen-plus",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.15,
            timeout=25
        )
        raw = (resp.choices[0].message.content or '').strip()
        match = re.search(r'\{[\s\S]*\}', raw)
        if not match:
            return None
        data = json.loads(match.group())

        sentiments = []
        for sentiment in ['positive', 'neutral', 'negative']:
            items = data.get(sentiment, [])
            for item in items[:3]:
                try:
                    count_val = int(re.sub(r'[^\d]', '', str(item.get('count_estimate', 0))) or 0)
                except (ValueError, TypeError):
                    count_val = 0
                sentiments.append({
                    'sentiment': sentiment,
                    'region': item.get('region', 'EN'),
                    'opinion': (item.get('opinion', '') or '')[:50],
                    'count': f"{count_val}+",
                    'examples': []
                })

        raw_ratio = data.get('sentiment_ratio', {})
        try:
            ratio = {
                'positive': round(float(raw_ratio.get('positive', 0)), 1),
                'neutral': round(float(raw_ratio.get('neutral', 0)), 1),
                'negative': round(float(raw_ratio.get('negative', 0)), 1),
            }
        except (ValueError, TypeError):
            ratio = None

        if sentiments:
            return {'sentiments': sentiments, 'sentiment_ratio': ratio}
        return None
    except Exception as e:
        logger.warning(f"⚠️ AI single post analysis failed, falling back to heuristic: {e}")
        return None


def _extract_entity_from_text(text):
    t = (text or '').lower()
    if not t:
        return ''
    entity_patterns = [
        ('Valir「Ember Gaze」皮肤', ['valir', 'ember gaze', '瓦里尔', 'valir skin']),
        ('Julian×Itachi 角色联动', ['julian', 'itachi', '朱利安', '鼬', 'uchiha itachi']),
        ('Tigreal×Marcel 角色玩法', ['tigreal', 'marcel', 'kiba', 'combo', '连招']),
        ('Naruto联动皮肤', ['naruto', '火影', '疾风传', '晓组织', 'uchiha', 'naruto skin']),
        ('免费联动皮肤活动', ['free skin', 'new skin', '#mlbbnewskin', 'free epic', '免费皮肤', '送皮肤']),
        ('MLBB联动活动', ['mlbbxnaruto', 'mlbb x naruto', '#mlbbxnaruto', '联动活动']),
    ]
    for entity, kws in entity_patterns:
        if any(k in t for k in kws):
            return entity
    return ''


def _extract_intent_from_text(text):
    t = (text or '').lower()
    if not t:
        return ''
    rules = [
        ('福利与获取门槛', ['free', '福利', '免费', '送', '抽卡', '钻石', 'diamond']),
        ('角色机制与平衡', ['skill', '技能', '机制', '平衡', 'buff', 'nerf', '强度', '太强', '太弱']),
        ('皮肤设计与还原度', ['skin', '皮肤', '特效', '设计', '还原', '同质化', '像不像']),
        ('活动规则与参与路径', ['how', 'where', 'when', '怎么', '哪里', '什么时候', '?', '规则']),
        ('性能与稳定性问题', ['bug', 'lag', 'crash', '卡顿', '闪退', '延迟', '崩溃']),
        ('购买与价格意愿', ['buy', '购买', '充值', '价格', '贵', 'too expensive']),
        ('辱骂与情绪宣泄', ['anjing', 'lonte', 'maling', 'kontol', 'bangsat', 'wtf', 'fuck', '垃圾', '傻逼']),
    ]
    for intent, kws in rules:
        if any(k in t for k in kws):
            return intent
    return ''


def _build_actionable_opinion(sentiment, entity, intent):
    if sentiment == 'positive':
        return f"玩家认可{entity}，正向反馈集中在{intent}"
    if sentiment == 'negative':
        return f"玩家质疑{entity}，负面反馈集中在{intent}"
    return f"玩家围绕{entity}进行信息讨论，主要关注{intent}"


def _comment_signal(row):
    brief = (row.get('_rt_brief') or row.get('brief_analysis') or '')
    content = (row.get('content') or '')
    text = f"{brief} {content}".strip()
    sentiment = row.get('_rt_sentiment') or row.get('sentiment') or 'neutral'
    region = row.get('region') or 'EN'
    entity = _extract_entity_from_text(text)
    intent = _extract_intent_from_text(text)
    return sentiment, region, entity, intent


def _is_generic_signal(entity, intent):
    generic_entities = {'', '联动活动', 'MLBB联动活动'}
    generic_intents = {'', '泛讨论反馈'}
    return entity in generic_entities or intent in generic_intents


def _is_actionable_signal(entity, intent):
    return (not _is_generic_signal(entity, intent))


def _extract_post_topic_from_text(text):
    entity = _extract_entity_from_text(text)
    if entity:
        return entity
    if _extract_intent_from_text(text):
        return "联动活动讨论"
    return "官方活动帖"


@app.route('/spd-report-tool')
@login_required
def spd_report_tool():
    """SPD 专题报告页面"""
    return render_template('spd_report.html')


@app.route('/api/spd_report_data', methods=['GET'])
@login_required
def spd_report_data():
    """SPD 报告数据（MVP）"""
    try:
        start_date = (request.args.get('start_date') or '').strip()
        end_date = (request.args.get('end_date') or '').strip()
        region = (request.args.get('region') or 'global').strip().lower()
        allowed_regions = {'global', 'th', 'id', 'en', 'pt'}
        if region not in allowed_regions:
            return jsonify({'status': 'error', 'message': 'region 参数非法'}), 400

        if not end_date:
            end_date = datetime.datetime.now().strftime('%Y-%m-%d')
        try:
            end_dt = datetime.datetime.strptime(end_date, '%Y-%m-%d')
        except ValueError:
            return jsonify({'status': 'error', 'message': 'end_date 格式错误，应为 YYYY-MM-DD'}), 400

        if not start_date:
            start_dt = end_dt - datetime.timedelta(days=13)
            start_date = start_dt.strftime('%Y-%m-%d')
        else:
            try:
                start_dt = datetime.datetime.strptime(start_date, '%Y-%m-%d')
            except ValueError:
                return jsonify({'status': 'error', 'message': 'start_date 格式错误，应为 YYYY-MM-DD'}), 400

        if start_dt > end_dt:
            return jsonify({'status': 'error', 'message': 'start_date 不能晚于 end_date'}), 400

        comment_rows = db.query_all("""
            SELECT post_url, comment_id, post_link, author, created_at, content, brief_analysis,
                   sentiment_score, language, category
            FROM fb_comments
            WHERE created_at >= %s
              AND created_at <= %s
            ORDER BY created_at ASC
        """, (start_date, f"{end_date} 23:59:59"))

        post_rows = db.query_all("""
            SELECT post_url, platform, author, post_date, post_content, thumbnail_url,
                   views, shares, likes, comments_count, engagement
            FROM fb_post_metrics
            WHERE post_date >= %s
              AND post_date <= %s
            ORDER BY post_date ASC
        """, (start_date, end_date))

        # 评论侧：用于情感与观点
        filtered_rows = []
        post_comments = defaultdict(list)
        post_region_votes = defaultdict(Counter)
        for row in comment_rows:
            row_dict = dict(row)
            row_region = _region_by_language(row_dict.get('language'))
            row_dict['region'] = row_region
            row_dict['is_mlbb'] = _match_task('MLBB', row_dict.get('content', ''), row_dict.get('brief_analysis', ''))
            row_dict['is_spd'] = _match_task('SPD', row_dict.get('content', ''), row_dict.get('brief_analysis', ''))
            row_dict['sentiment'] = _sentiment_bucket(row_dict.get('sentiment_score'))
            post_key = row_dict.get('post_url') or row_dict.get('post_link') or 'unknown'
            post_comments[post_key].append(row_dict)
            post_region_votes[post_key][row_region] += 1
            if region != 'global' and row_region != region.upper():
                continue
            filtered_rows.append(row_dict)

        def _post_region(post_key):
            votes = post_region_votes.get(post_key)
            if not votes:
                return 'EN'
            return votes.most_common(1)[0][0]

        # 帖子侧：用于趋势、峰值、Top5按 engagement 排序
        filtered_posts = []
        for prow in post_rows:
            p = dict(prow)
            post_key = p.get('post_url') or 'unknown'
            p_region = _post_region(post_key)
            if region != 'global' and p_region != region.upper():
                continue
            p['region'] = p_region
            p['is_mlbb'] = _match_task('MLBB', p.get('post_content', ''), '') or _is_mlbb_relaxed_text(p.get('post_content', ''), '')
            p['is_spd'] = _match_task('SPD', p.get('post_content', ''), '') or _is_spd_relaxed_text(p.get('post_content', ''), '')
            p['engagement'] = int(p.get('engagement') or 0)
            p['likes'] = int(p.get('likes') or 0)
            p['shares'] = int(p.get('shares') or 0)
            p['views'] = int(p.get('views') or 0)
            p['comments_count'] = int(p.get('comments_count') or 0)
            filtered_posts.append(p)

        # 历史数据兼容：评论文本严格命中通常偏少，先做宽松匹配
        strict_spd_comments = [r for r in filtered_rows if r.get('is_spd')]
        if not strict_spd_comments and filtered_rows:
            relaxed_rows = []
            for row in filtered_rows:
                if _is_spd_relaxed_text(row.get('content', ''), row.get('brief_analysis', '')):
                    row['is_spd'] = True
                    relaxed_rows.append(row)
            if relaxed_rows:
                logger.info(f"ℹ️ SPD report fallback: relaxed comment matching enabled ({len(relaxed_rows)} rows)")
            else:
                for row in filtered_rows:
                    row['is_spd'] = True
                logger.info(f"ℹ️ SPD report fallback: use all filtered comments as SPD ({len(filtered_rows)} rows)")

        # 关键修正：一旦某帖子被识别为 SPD，整帖评论都应纳入（否则会出现“帖子几千评，报告仅百条”）
        spd_post_urls = set()
        for row in filtered_rows:
            if row.get('is_spd'):
                pkey = row.get('post_url') or row.get('post_link')
                if pkey:
                    spd_post_urls.add(pkey)
        for post in filtered_posts:
            if post.get('is_spd') and post.get('post_url'):
                spd_post_urls.add(post.get('post_url'))

        if spd_post_urls:
            expanded = 0
            for row in filtered_rows:
                pkey = row.get('post_url') or row.get('post_link')
                if pkey in spd_post_urls and not row.get('is_spd'):
                    row['is_spd'] = True
                    expanded += 1
            if expanded:
                logger.info(f"ℹ️ SPD report expanded by post_url membership: +{expanded} comments")

        if filtered_posts and (not any(p.get('is_spd') for p in filtered_posts)):
            relaxed_post_hits = 0
            for post in filtered_posts:
                if _is_spd_relaxed_text(post.get('post_content', ''), ''):
                    post['is_spd'] = True
                    relaxed_post_hits += 1
            if relaxed_post_hits:
                logger.info(f"ℹ️ SPD report fallback: relaxed post matching enabled ({relaxed_post_hits} posts)")

        daily_counter = defaultdict(lambda: {'mlbb': 0, 'spd': 0})
        daily_post_counter = {
            'mlbb': defaultdict(lambda: defaultdict(int)),
            'spd': defaultdict(lambda: defaultdict(int))
        }

        for post in filtered_posts:
            date_key = (post.get('post_date').strftime('%Y-%m-%d')
                        if hasattr(post.get('post_date'), 'strftime')
                        else str(post.get('post_date') or start_date))
            post_key = post.get('post_url') or 'unknown'
            # 折线图口径：按 hashtag 抓取到的帖子，按日汇总“赞+评+转”
            engagement_val = int(post.get('likes') or 0) + int(post.get('shares') or 0) + int(post.get('comments_count') or 0)
            post['effective_engagement'] = engagement_val
            if post['is_mlbb']:
                daily_counter[date_key]['mlbb'] += engagement_val
                daily_post_counter['mlbb'][date_key][post_key] += engagement_val
            if post['is_spd']:
                daily_counter[date_key]['spd'] += engagement_val
                daily_post_counter['spd'][date_key][post_key] += engagement_val

        labels = sorted(daily_counter.keys())
        # 兜底：若帖子指标暂缺，回退到评论口径，避免报告空白
        if not labels and filtered_rows:
            for row in filtered_rows:
                if not row.get('created_at'):
                    continue
                date_key = row['created_at'].strftime('%Y-%m-%d')
                post_key = row.get('post_url') or row.get('post_link') or 'unknown'
                if row.get('is_mlbb'):
                    daily_counter[date_key]['mlbb'] += 1
                    daily_post_counter['mlbb'][date_key][post_key] += 1
                if row.get('is_spd'):
                    daily_counter[date_key]['spd'] += 1
                    daily_post_counter['spd'][date_key][post_key] += 1
            labels = sorted(daily_counter.keys())

        mlbb_series = [daily_counter[d]['mlbb'] for d in labels]
        spd_series = [daily_counter[d]['spd'] for d in labels]

        peak_drilldown = []
        for task_key, task_name in [('mlbb', 'MLBB'), ('spd', 'SPD')]:
            top_dates = sorted(labels, key=lambda d: daily_counter[d][task_key], reverse=True)[:2]
            for peak_date in top_dates:
                top_posts = []
                post_counts_on_date = daily_post_counter[task_key][peak_date]
                for post_key, count in sorted(post_counts_on_date.items(), key=lambda x: x[1], reverse=True)[:2]:
                    matched_post = next((p for p in filtered_posts if (p.get('post_url') or '') == post_key), {})
                    comments = [c for c in post_comments.get(post_key, []) if (task_key != 'spd' or c.get('is_spd'))]
                    first = comments[0] if comments else {}
                    topic_source = (matched_post.get('post_content') or first.get('content') or first.get('brief_analysis') or '')
                    topic = _extract_post_topic_from_text(topic_source)
                    topic = (topic or '官方活动帖')[:48]
                    top_posts.append({
                        'author': matched_post.get('author') or first.get('author') or 'unknown',
                        'author_region': matched_post.get('region') or first.get('region') or 'EN',
                        'post_topic': topic,
                        'post_url': post_key,
                        'engagement': count
                    })
                if top_posts:
                    peak_drilldown.append({
                        'date': peak_date,
                        'task': task_name,
                        'region': region.upper() if region != 'global' else 'GLOBAL',
                        'top_posts': top_posts
                    })

        # Top5：仅 SPD 相关帖子，按 engagement 排序
        spd_posts_sorted = sorted(
            [p for p in filtered_posts if p.get('is_spd')],
            key=lambda x: int(x.get('effective_engagement') or x.get('engagement') or 0),
            reverse=True
        )[:5]
        if not spd_posts_sorted:
            fallback_counts = Counter()
            for row in filtered_rows:
                if row.get('is_spd'):
                    pkey = row.get('post_url') or row.get('post_link') or 'unknown'
                    fallback_counts[pkey] += 1
            for post_key, count in fallback_counts.most_common(5):
                sample = (post_comments.get(post_key) or [{}])[0]
                spd_posts_sorted.append({
                    'post_url': post_key,
                    'platform': 'FB',
                    'author': sample.get('author') or 'unknown',
                    'region': sample.get('region') or 'EN',
                    'post_content': sample.get('content') or '',
                    'thumbnail_url': '',
                    'views': 0,
                    'shares': 0,
                    'likes': 0,
                    'comments_count': count,
                    'engagement': count,
                    'effective_engagement': count,
                })

        top5_post_urls = {(p.get('post_url') or 'unknown') for p in spd_posts_sorted[:5]}
        top5_comment_pool = [
            row for row in filtered_rows
            if (row.get('post_url') or row.get('post_link') or 'unknown') in top5_post_urls
        ]
        if not top5_comment_pool:
            top5_comment_pool = [row for row in filtered_rows if row.get('is_spd')]

        top5_comment_pool = _reanalyze_comments_for_template(top5_comment_pool)
        top5_pool_by_post = defaultdict(list)
        for row in top5_comment_pool:
            pkey = row.get('post_url') or row.get('post_link') or 'unknown'
            top5_pool_by_post[pkey].append(row)

        total_spd_comments = sum(int(p.get('comments_count') or 0) for p in spd_posts_sorted[:5]) or len(top5_comment_pool)
        sentiment_counter = Counter([row.get('_rt_sentiment') or row.get('sentiment') for row in top5_comment_pool])

        ai_result = _ai_aggregate_opinions(top5_comment_pool, total_spd_comments)
        ai_sentiment_ratio = None
        if ai_result and ai_result.get('opinions'):
            key_opinions = ai_result['opinions']
            ai_sentiment_ratio = ai_result.get('sentiment_ratio')
            logger.info(f"✅ SPD report: AI aggregate produced {len(key_opinions)} opinions")
        else:
            signal_counter = Counter()
            signal_examples = defaultdict(list)
            min_support = 2
            for row in top5_comment_pool:
                s_label, region_label, entity, intent = _comment_signal(row)
                key = (s_label, region_label, entity, intent)
                signal_counter[key] += 1
                if len(signal_examples[key]) < 2:
                    original = (row.get('content') or '')[:120]
                    translation = (row.get('_rt_brief') or row.get('brief_analysis') or '').strip() or original
                    if not _contains_zh(translation):
                        translation = _translate_text_to_zh(translation)
                    signal_examples[key].append({'original': original, 'translation': translation[:120]})

            key_opinions = []
            for sentiment in ['positive', 'neutral', 'negative']:
                bucket = [((s, r, e, i), c) for (s, r, e, i), c in signal_counter.items() if s == sentiment]
                bucket.sort(key=lambda x: x[1], reverse=True)
                selected = []
                used_ei = set()
                candidates = [item for item in bucket if item[1] >= min_support and _is_actionable_signal(item[0][2], item[0][3])]
                for item in candidates:
                    ei = (item[0][2], item[0][3])
                    if ei in used_ei:
                        continue
                    selected.append(item)
                    used_ei.add(ei)
                    if len(selected) >= 3:
                        break
                for (s, op_region, entity, intent), count in selected:
                    key_opinions.append({
                        'sentiment': s,
                        'region': op_region,
                        'opinion': _build_actionable_opinion(s, entity, intent),
                        'percentage': _to_percentage(count, total_spd_comments)
                    })
            logger.info(f"ℹ️ SPD report: heuristic fallback produced {len(key_opinions)} opinions")

        top5_posts = []
        min_support = 2
        translation_cache = {}
        for idx, post in enumerate(spd_posts_sorted, start=1):
            post_key = post.get('post_url') or 'unknown'
            comments = list(top5_pool_by_post.get(post_key, []))
            if not comments:
                comments = [row for row in filtered_rows if (row.get('post_url') or row.get('post_link') or 'unknown') == post_key]
            total = len(comments)
            sent_count = Counter([(row.get('_rt_sentiment') or row.get('sentiment') or 'neutral') for row in comments])

            ai_post_result = _ai_analyze_single_post(comments, total)
            post_ai_ratio = None
            if ai_post_result and ai_post_result.get('sentiments'):
                top_sentiments = ai_post_result['sentiments']
                post_ai_ratio = ai_post_result.get('sentiment_ratio')
            else:
                top_sentiments = []
                for sentiment in ['positive', 'neutral', 'negative']:
                    sentiment_rows = [row for row in comments if (row.get('_rt_sentiment') or row.get('sentiment') or 'neutral') == sentiment]
                    local_counter = Counter()
                    local_examples = defaultdict(list)
                    for row in sentiment_rows:
                        _, region_label, entity, intent = _comment_signal(row)
                        lk = (region_label, entity, intent)
                        local_counter[lk] += 1
                        if len(local_examples[lk]) < 2:
                            original = (row.get('content') or '')[:120]
                            translation = (row.get('_rt_brief') or row.get('brief_analysis') or '').strip() or original
                            if not _contains_zh(translation):
                                cache_key = translation.lower()
                                if cache_key not in translation_cache:
                                    translation_cache[cache_key] = _translate_text_to_zh(translation)
                                translation = translation_cache[cache_key]
                            local_examples[lk].append({'original': original, 'translation': translation[:120]})
                    ranked = sorted(local_counter.items(), key=lambda x: x[1], reverse=True)
                    picked = []
                    used_ei = set()
                    candidates = [item for item in ranked if item[1] >= min_support and _is_actionable_signal(item[0][1], item[0][2])]
                    for item in candidates:
                        ei = (item[0][1], item[0][2])
                        if ei in used_ei:
                            continue
                        picked.append(item)
                        used_ei.add(ei)
                        if len(picked) >= 3:
                            break
                    for (region_label, entity, intent), c_count in picked:
                        examples = local_examples.get((region_label, entity, intent), [])
                        top_sentiments.append({
                            'sentiment': sentiment,
                            'region': region_label,
                            'opinion': _build_actionable_opinion(sentiment, entity, intent),
                            'count': f"{c_count}+",
                            'examples': examples
                        })

            first = comments[0] if comments else {}
            topic_source = (post.get('post_content') or first.get('content') or first.get('brief_analysis') or '')
            topic = _extract_post_topic_from_text(topic_source)

            top5_posts.append({
                'rank': idx,
                'region': post.get('region') or first.get('region', 'EN'),
                'platform': post.get('platform') or 'FB',
                'post_content': {
                    'topic': (topic or '帖子讨论')[:60],
                    'text': ((post.get('post_content') or first.get('content') or '')[:220]),
                    'thumbnail_url': post.get('thumbnail_url') or '',
                },
                'sentiment_ratio': post_ai_ratio if post_ai_ratio else {
                    'positive': _to_percentage(sent_count.get('positive', 0), total),
                    'neutral': _to_percentage(sent_count.get('neutral', 0), total),
                    'negative': _to_percentage(sent_count.get('negative', 0), total),
                },
                'top_sentiments': top_sentiments,
                'shares': int(post.get('shares') or 0),
                'likes': int(post.get('likes') or 0),
                'comments': int(post.get('comments_count') or total),
                'engagement': int(post.get('effective_engagement') or post.get('engagement') or 0),
                'post_url': post_key
            })

        return jsonify({
            'status': 'success',
            'meta': {
                'task_name': 'SPD',
                'start_date': start_date,
                'end_date': end_date,
                'region': region
            },
            'trend': {
                'labels': labels,
                'mlbb': mlbb_series,
                'spd': spd_series,
            },
            'peak_drilldown': peak_drilldown,
            'summary': {
                'total_comments': total_spd_comments,
                'sentiment_ratio': ai_sentiment_ratio if ai_sentiment_ratio else {
                    'positive': _to_percentage(sentiment_counter.get('positive', 0), total_spd_comments),
                    'neutral': _to_percentage(sentiment_counter.get('neutral', 0), total_spd_comments),
                    'negative': _to_percentage(sentiment_counter.get('negative', 0), total_spd_comments),
                },
                'key_opinions': key_opinions
            },
            'top5_posts': top5_posts
        })
    except Exception as e:
        logger.error(f"❌ 生成 SPD 报告数据失败: {e}")
        return jsonify({'status': 'error', 'message': '内部错误，请稍后重试'}), 500


@app.route('/analyze', methods=['POST'])
@login_required
def analyze():
    """舆情分析 API - 异步版本（v2：支持多链接 FB/IG/TT/YTB）"""
    logger.info("\n" + "=" * 60)
    logger.info("📥 收到舆情分析请求")
    logger.info(f"🔑 DASHSCOPE_API_KEY: {'✅' if DASHSCOPE_API_KEY else '❌'}")
    logger.info(f"🔑 APIFY_TOKEN: {'✅' if APIFY_TOKEN else '❌'}")

    # 兼容旧字段 url（单链接），新字段 urls（多行/逗号分隔）
    url = (request.form.get('url') or '').strip()
    urls_raw = request.form.get('urls') or ''
    urls = sentiment_insight.parse_urls_text(urls_raw)
    if not urls and url:
        urls = sentiment_insight.parse_urls_text(url)

    if not urls:
        return jsonify({'error': '请至少提供一个有效链接'}), 400

    comment_limit_unlimited = _truthy_form_value(request.form.get('comment_limit_unlimited'))
    comments_per_post_limit = normalize_insight_comment_limit(
        request.form.get('comments_per_post_limit') or request.form.get('comment_limit'),
        unlimited=comment_limit_unlimited,
    )

    # 标注无法识别平台的链接
    unsupported = [u for u in urls if sentiment_insight.detect_platform(u) == 'UNKNOWN']
    if unsupported:
        logger.warning(f"⚠️ 以下链接平台无法识别，将被跳过: {unsupported}")

    # 生成任务 ID
    task_id = str(uuid.uuid4())
    session_id = session.get('session_id', 'default')

    # 在主线程中提取用户信息（避免线程安全问题）
    user_id = session.get('user_id')
    username = session.get('username', 'unknown')
    department = session.get('department', '未知')

    # 创建任务记录到数据库（标记类型为 sentiment）
    try:
        create_task(task_id, user_id, session_id, function_type='sentiment')
    except Exception as e:
        logger.exception(f"❌ 创建舆情任务失败: {e}")
        return jsonify({
            'status': 'error',
            'error': '任务创建失败，请稍后重试或联系管理员查看后台日志',
            'message': '任务创建失败，请稍后重试或联系管理员查看后台日志',
        }), 500

    task_params = {
        'urls': urls,
        'session_id': session_id,
        'user_id': user_id,
        'username': username,
        'department': department,
        'comments_per_post_limit': comments_per_post_limit,
        'comment_limit_unlimited': comment_limit_unlimited,
    }
    try:
        set_task_params(task_id, task_params)
    except Exception as e:
        logger.error(f"❌ 写入 task_params 失败: {e}")

    if not USE_DB_WORKER:
        thread = threading.Thread(
            target=process_analysis_task,
            kwargs={
                'task_id': task_id,
                'urls': urls,
                'session_id': session_id,
                'user_id': user_id,
                'username': username,
                'department': department,
                'comments_per_post_limit': comments_per_post_limit,
            }
        )
        try:
            thread.start()
        except Exception as e:
            logger.exception(f"❌ 启动舆情任务线程失败: {e}")
            try:
                update_task(task_id, status='failed', error=str(e), progress='任务启动失败')
            except Exception:
                pass
            return jsonify({
                'status': 'error',
                'error': '任务启动失败，请稍后重试或联系管理员查看后台日志',
                'message': '任务启动失败，请稍后重试或联系管理员查看后台日志',
            }), 500
        logger.info(f"✅ 任务 {task_id} 已创建并在本进程中启动（urls={len(urls)}）")
    else:
        # DB worker 模式：把参数序列化写入 task_params，由独立 worker 进程拾取
        logger.info(f"✅ 任务 {task_id} 已创建，等待外部 worker 处理")

    # 立即返回任务 ID
    return jsonify({
        'task_id': task_id,
        'status': 'pending',
        'message': '任务已提交，正在后台处理...'
    })


@app.route('/api/sentiment/precheck', methods=['POST'])
@login_required
def sentiment_precheck_api():
    """提交前预检：识别平台、异常链接和粗略耗时。"""
    try:
        data = request.get_json(silent=True) or {}
        urls_raw = data.get('urls') or request.form.get('urls') or request.form.get('url') or ''
        urls = sentiment_insight.parse_urls_text(urls_raw)
        comment_limit_unlimited = _truthy_form_value(
            data.get('comment_limit_unlimited') if 'comment_limit_unlimited' in data else request.form.get('comment_limit_unlimited')
        )
        comments_per_post_limit = normalize_insight_comment_limit(
            data.get('comments_per_post_limit') or data.get('comment_limit') or request.form.get('comments_per_post_limit'),
            unlimited=comment_limit_unlimited,
        )
        platform_counts = Counter(sentiment_insight.detect_platform(url) for url in urls)
        unsupported = [url for url in urls if sentiment_insight.detect_platform(url) == 'UNKNOWN']
        supported_count = len(urls) - len(unsupported)
        min_minutes = max(1, supported_count)
        max_minutes = max(2, supported_count * 2)
        return jsonify({
            'status': 'success',
            'url_count': len(urls),
            'supported_count': supported_count,
            'platform_counts': dict(platform_counts),
            'unsupported': unsupported,
            'estimated_seconds_min': min_minutes * 60,
            'estimated_seconds_max': max_minutes * 60,
            'estimated_text': f'约 {min_minutes}-{max_minutes} 分钟' if supported_count else '无法估算',
            'comments_per_post_limit': comments_per_post_limit,
            'comment_limit_unlimited': comment_limit_unlimited,
            'links': [{'url': url, 'platform': sentiment_insight.detect_platform(url)} for url in urls],
        })
    except Exception as e:
        logger.error(f"❌ 舆情预检失败: {e}")
        return jsonify({'status': 'error', 'message': str(e)}), 500


@app.route('/api/sentiment/tasks')
@login_required
def sentiment_tasks_api():
    """舆情工作台统一任务中心：评论分析、拉视频数据、主页同步。"""
    try:
        limit = max(1, min(int(request.args.get('limit') or 30), 100))
    except (TypeError, ValueError):
        limit = 30
    uid = session.get('user_id')
    rows = db.query_all(
        """
        SELECT task_id, function_type, status, progress, result, error, record_id,
               created_at, updated_at, started_at, finished_at
        FROM task_queue
        WHERE (user_id = %s OR (user_id IS NULL AND function_type = 'feishu_profile_video_sync'))
          AND function_type IN ('sentiment', 'etl_video_metrics', 'profile_video_sync', 'feishu_profile_video_sync')
        ORDER BY created_at DESC
        LIMIT %s
        """,
        (uid, limit),
    ) or []

    download_by_task = {}
    task_ids = [row.get('task_id') for row in rows if row.get('task_id')]
    if task_ids:
        placeholders = ','.join(['%s'] * len(task_ids))
        output_rows = db.query_all(
            f"""
            SELECT DISTINCT ON (task_id) task_id, id AS download_id, filename
            FROM etl_file_outputs
            WHERE user_id = %s
              AND filename <> '_input_video_metrics.xlsx'
              AND task_id IN ({placeholders})
            ORDER BY task_id, id DESC
            """,
            tuple([uid] + task_ids),
        ) or []
        download_by_task = {row.get('task_id'): row for row in output_rows}

    items = []
    for row in rows:
        result_payload = {}
        raw_result = row.get('result')
        if raw_result:
            try:
                result_payload = json.loads(raw_result) if isinstance(raw_result, str) else dict(raw_result)
            except Exception:
                result_payload = {}
        task_id = row.get('task_id')
        queue_info = get_task_queue_position(task_id, row.get('status')) if task_id else {}
        download_row = download_by_task.get(task_id) or {}
        download_id = result_payload.get('download_id') or download_row.get('download_id')
        filename = result_payload.get('filename') or download_row.get('filename')
        success_count = result_payload.get('success_count')
        total_count = (
            result_payload.get('url_count')
            or result_payload.get('profile_count')
            or result_payload.get('total')
        )
        items.append({
            'task_id': task_id,
            'function_type': row.get('function_type'),
            'task_name': (
                '主页视频导出'
                if row.get('function_type') == 'etl_video_metrics' and result_payload.get('mode') == 'profile_videos'
                else _sentiment_task_name(row.get('function_type'))
            ),
            'status': row.get('status'),
            'progress': row.get('progress'),
            'error': row.get('error'),
            'record_id': row.get('record_id'),
            'created_at': row.get('created_at'),
            'updated_at': row.get('updated_at'),
            'started_at': row.get('started_at'),
            'finished_at': row.get('finished_at'),
            'download_id': download_id,
            'filename': filename,
            'success_count': success_count,
            'total_count': total_count,
            'failed_count': result_payload.get('failed_count'),
            'queue_position': queue_info.get('queue_position'),
            'tasks_ahead': queue_info.get('tasks_ahead', 0),
            'pending_count': queue_info.get('pending_count', 0),
            'running_count': queue_info.get('running_count', 0),
        })
    return jsonify({'status': 'success', 'items': _json_safe_rows(items)})


def _sentiment_task_name(function_type):
    return {
        'sentiment': '评论分析',
        'etl_video_metrics': '拉视频数据',
        'profile_video_sync': '主页同步',
        'feishu_profile_video_sync': '飞书主页同步',
    }.get(function_type or '', function_type or '任务')


@app.route('/api/sentiment/tasks/<task_id>/retry', methods=['POST'])
@login_required
def sentiment_task_retry_api(task_id):
    """重试统一任务中心中的可恢复任务。"""
    uid = session.get('user_id')
    row = db.query_one(
        """
        SELECT task_id, function_type, task_params, session_id
        FROM task_queue
        WHERE task_id = %s
          AND (user_id = %s OR (user_id IS NULL AND function_type = 'feishu_profile_video_sync'))
        """,
        (task_id, uid),
    )
    if not row:
        return jsonify({'status': 'error', 'message': '任务不存在或无权限访问'}), 404
    function_type = row.get('function_type')
    if function_type not in {'sentiment', 'etl_video_metrics', 'profile_video_sync', 'feishu_profile_video_sync'}:
        return jsonify({'status': 'error', 'message': '该任务类型暂不支持重试'}), 400
    raw_params = row.get('task_params')
    if not raw_params:
        return jsonify({'status': 'error', 'message': '旧任务缺少原始参数，无法自动重试'}), 400
    try:
        params = json.loads(raw_params) if isinstance(raw_params, str) else dict(raw_params)
    except Exception:
        return jsonify({'status': 'error', 'message': '任务参数已损坏，无法自动重试'}), 400

    new_task_id = str(uuid.uuid4())
    session_id = params.get('session_id') or row.get('session_id') or session.get('session_id', 'default')
    create_task(new_task_id, uid, session_id, function_type=function_type)
    params['user_id'] = uid
    params['session_id'] = session_id
    if function_type == 'feishu_profile_video_sync':
        params['trigger_type'] = 'retry'
    set_task_params(new_task_id, params)

    if not USE_DB_WORKER:
        if function_type == 'sentiment':
            threading.Thread(
                target=process_analysis_task,
                kwargs={
                    'task_id': new_task_id,
                    'urls': params.get('urls'),
                    'url': params.get('url'),
                    'session_id': session_id,
                    'user_id': uid,
                    'username': params.get('username') or session.get('username', 'unknown'),
                    'department': params.get('department') or session.get('department', '未知'),
                    'project': params.get('project', 'CFL'),
                    'comments_per_post_limit': params.get('comments_per_post_limit'),
                },
                daemon=True,
            ).start()
        elif function_type == 'etl_video_metrics':
            threading.Thread(
                target=lambda: etl_jobs.run_etl_video_metrics_task(new_task_id, params, update_task),
                daemon=True,
            ).start()
        elif function_type == 'profile_video_sync':
            threading.Thread(
                target=lambda: profile_video_scheduler.run_profile_video_sync_task(new_task_id, params, update_task),
                daemon=True,
            ).start()
        elif function_type == 'feishu_profile_video_sync':
            threading.Thread(
                target=lambda: profile_video_scheduler.run_feishu_profile_video_sync_task(new_task_id, params, update_task),
                daemon=True,
            ).start()

    return jsonify({'status': 'queued', 'task_id': new_task_id})


@app.route('/api/tasks/summary')
@login_required
def tasks_summary_api():
    """任务队列汇总：用于看当前到底是排队、运行还是失败。"""
    try:
        rows = db.query_all("""
            SELECT status, function_type, COUNT(*) AS count
            FROM task_queue
            WHERE created_at >= NOW() - INTERVAL '7 days'
            GROUP BY status, function_type
            ORDER BY status, function_type
        """) or []
        totals = {}
        by_module = {}
        for row in rows:
            status = row.get('status') or 'unknown'
            module = row.get('function_type') or 'unknown'
            count = int(row.get('count') or 0)
            totals[status] = totals.get(status, 0) + count
            by_module.setdefault(module, {})[status] = count
        oldest_pending = db.query_one("""
            SELECT task_id, function_type, created_at, progress
            FROM task_queue
            WHERE status = 'pending'
            ORDER BY created_at ASC
            LIMIT 1
        """)
        running = db.query_all("""
            SELECT task_id, function_type, progress, worker_id, started_at, updated_at
            FROM task_queue
            WHERE status IN ('claimed', 'processing')
            ORDER BY COALESCE(started_at, updated_at, created_at) ASC
            LIMIT 20
        """) or []
        return jsonify({
            'totals': totals,
            'by_module': by_module,
            'oldest_pending': _serialize_task_row(oldest_pending) if oldest_pending else None,
            'running': [_serialize_task_row(row) for row in running],
        })
    except Exception as e:
        logger.error(f"❌ 获取任务队列汇总失败: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/tasks')
@login_required
def tasks_list_api():
    """任务队列明细。普通用户看自己的，管理员可 all_users=1。"""
    try:
        status = (request.args.get('status') or '').strip()
        all_users = request.args.get('all_users') == '1' and session.get('role') == 'admin'
        limit = max(1, min(int(request.args.get('limit') or 100), 500))
        where = []
        params = []
        if status:
            where.append('status = %s')
            params.append(status)
        if not all_users:
            where.append('user_id = %s')
            params.append(session.get('user_id'))
        sql_where = 'WHERE ' + ' AND '.join(where) if where else ''
        params.append(limit)
        rows = db.query_all(f"""
            SELECT task_id, user_id, function_type, status, progress, error, record_id,
                   worker_id, attempts, created_at, updated_at, started_at, finished_at
            FROM task_queue
            {sql_where}
            ORDER BY created_at DESC
            LIMIT %s
        """, tuple(params)) or []
        return jsonify({'rows': [_serialize_task_row(row) for row in rows]})
    except Exception as e:
        logger.error(f"❌ 获取任务队列明细失败: {e}")
        return jsonify({'error': str(e)}), 500


def _serialize_task_row(row):
    data = dict(row or {})
    for key in ('created_at', 'updated_at', 'started_at', 'finished_at'):
        if data.get(key) and hasattr(data[key], 'isoformat'):
            data[key] = data[key].isoformat()
    return data


@app.route('/task_status/<task_id>')
def task_status(task_id):
    """查询任务状态"""
    task = get_task(task_id)

    if not task:
        return jsonify({'error': '任务不存在'}), 404

    result = task.get('result') or ''
    if result and task['status'] == 'completed':
        result = sanitize_html(result)
    resp = {
        'task_id': task.get('task_id'),
        'status': task['status'],
        'function_type': task.get('function_type'),
        'progress': task['progress'],
        'result': result,
        'error': task['error'],
        'record_id': task.get('record_id'),
        'queue_position': task.get('queue_position'),
        'tasks_ahead': task.get('tasks_ahead', 0),
        'pending_count': task.get('pending_count', 0),
        'running_count': task.get('running_count', 0),
        'worker_id': task.get('worker_id'),
        'attempts': task.get('attempts') or 0,
        'created_at': task.get('created_at').isoformat() if task.get('created_at') else None,
        'updated_at': task.get('updated_at').isoformat() if task.get('updated_at') else None,
        'started_at': task.get('started_at').isoformat() if task.get('started_at') else None,
        'finished_at': task.get('finished_at').isoformat() if task.get('finished_at') else None,
    }
    return jsonify(resp)

# ============================================
# 功能 2: 竞品监控
# ============================================

@app.route('/competitor-tool')
@login_required
def competitor_tool():
    """竞品监控工具页面"""
    user_id = session.get('user_id')

    has_used_competitor = False
    if user_id:
        try:
            row = db.query_one(
                """
                SELECT 1 FROM analysis_results
                WHERE user_id = %s AND type = %s
                LIMIT 1
                """,
                (user_id, 'competitor')
            )
            has_used_competitor = bool(row)
        except Exception as e:
            logger.error(f"❌ 检查竞品历史记录失败: {e}")

    return render_template('competitor.html', has_used_competitor=has_used_competitor)


@app.route('/monitor_competitors', methods=['POST'])
def monitor_competitors():
    """竞品监控 API - 异步版本"""
    logger.info("\n" + "=" * 60)
    logger.info("📥 收到竞品监控请求")

    try:
        data = request.json or {}
        target_url = data.get('competitor_name')  # 兼容老前端
        # 新前端：urls 数组 / urls_text 多行文本
        urls_input = data.get('urls')
        urls_text = data.get('urls_text') or data.get('urlsText')
        urls: list[str] = []
        if isinstance(urls_input, list):
            for u in urls_input:
                if isinstance(u, str) and u.strip():
                    urls.append(u.strip())
        if urls_text:
            urls.extend(competitor_radar.parse_urls_text(urls_text))
        if not urls and target_url:
            urls = [target_url.strip()]
        # 去重
        urls = list(dict.fromkeys(urls))

        start_dt_str = data.get('startDate')
        end_dt_str = data.get('endDate')
        project = (data.get('project') or 'CFL').strip().upper()
        generate_report = bool(data.get('generateReport', False))
        enable_video_vision = bool(data.get('enableVideoVision', False))
        if project not in VALID_PROJECTS:
            project = 'CFL'

        logger.info(f"🎯 目标 URLs: {urls}")
        logger.info(f"📅 时间段: {start_dt_str} ~ {end_dt_str}，项目: {project}")
        logger.info(f"📊 生成分析报告: {generate_report}，启用看视频分析: {enable_video_vision}")

        if not urls:
            return jsonify({'error': '请至少提供一个 TikTok / Instagram 主页链接'}), 400

        # 平台校验
        unsupported = [u for u in urls if competitor_radar.detect_platform(u) == 'UNKNOWN']
        if unsupported:
            return jsonify({'error': f'仅支持 TikTok / Instagram 主页，以下链接无法识别：{unsupported[0]}'}), 400

        if not APIFY_TOKEN:
            error_msg = "❌ 错误：APIFY_TOKEN 未配置，无法使用爬虫功能"
            logger.error(error_msg)
            return jsonify({'error': error_msg}), 400

        # 获取用户信息
        user_id = session.get('user_id')
        username = session.get('username', 'unknown')
        department = session.get('department', '未知')
        session_id = session.get('session_id', str(uuid.uuid4()))

        # 创建任务 ID
        task_id = str(uuid.uuid4())

        # 创建任务记录到数据库（标记类型为 competitor）
        create_task(task_id, user_id, session_id, function_type='competitor')

        # 根据模式选择执行方式
        if not USE_DB_WORKER:
            # 进程内线程执行
            thread = threading.Thread(
                target=process_competitor_task,
                kwargs=dict(
                    task_id=task_id,
                    target_url=urls[0] if urls else None,  # 兼容旧签名
                    start_dt_str=start_dt_str,
                    end_dt_str=end_dt_str,
                    user_id=user_id,
                    username=username,
                    department=department,
                    session_id=session_id,
                    project=project,
                    generate_report=generate_report,
                    enable_video_vision=enable_video_vision,
                    urls=urls,
                ),
            )
            thread.start()
            logger.info(f"✅ 竞品监控任务 {task_id} 已创建并在本进程中启动")
        else:
            # DB worker 模式
            task_params = {
                'target_url': urls[0] if urls else None,
                'urls': urls,
                'start_dt_str': start_dt_str,
                'end_dt_str': end_dt_str,
                'user_id': user_id,
                'username': username,
                'department': department,
                'session_id': session_id,
                'project': project,
                'generate_report': generate_report,
                'enable_video_vision': enable_video_vision,
            }
            try:
                db.execute(
                    "UPDATE task_queue SET task_params = %s WHERE task_id = %s",
                    (json.dumps(task_params, ensure_ascii=False), task_id)
                )
            except Exception as e:
                logger.error(f"❌ 写入 task_params 失败: {e}")
            logger.info(f"✅ 竞品监控任务 {task_id} 已创建，等待外部 worker 处理")

        # 立即返回任务 ID
        return jsonify({
            'task_id': task_id,
            'status': 'pending',
            'message': '竞品监控任务已启动，请稍后查看结果'
        })

    except Exception as e:
        error_msg = f"❌ 创建任务失败: {str(e)}"
        logger.error(error_msg)
        import traceback
        traceback.print_exc()
        return jsonify({'error': error_msg}), 500


def _run_competitor_radar_task(task_id, urls, start_dt_str, end_dt_str,
                               user_id, username, department, session_id,
                               project='CFL', enable_video_vision=False):
    """新版竞品雷达流程：多主页 TT/IG 视频采集 + 可选视觉分析 + 12 列表格输出。"""
    logger.info(f"🔄 [Radar] 开始处理任务 {task_id} | urls={urls} | project={project}")
    update_task(task_id, status='processing', progress='正在初始化雷达任务...')

    if not APIFY_TOKEN:
        update_task(task_id, status='failed', error='APIFY_TOKEN 未配置，无法启动爬虫')
        return

    start_date = None
    end_date = None
    try:
        if start_dt_str:
            start_date = datetime.datetime.strptime(start_dt_str, '%Y-%m-%d').date()
        if end_dt_str:
            end_date = datetime.datetime.strptime(end_dt_str, '%Y-%m-%d').date()
    except Exception as e:
        update_task(task_id, status='failed', error=f'日期解析失败: {e}')
        return

    def _progress(msg):
        try:
            update_task(task_id, progress=msg)
        except Exception:
            pass

    # 视觉分析函数（按需启用）
    vision_call = None
    if enable_video_vision:
        try:
            import video_vision

            # 收集所有视频，统一批量分析，避免在 pipeline 内逐个 API 调用
            def _build_lazy_vision_call(_proj):
                _result_cache = {}
                _resolved = {"done": False}

                def _do_resolve(all_videos):
                    if _resolved["done"]:
                        return
                    _resolved["done"] = True
                    try:
                        mapping = video_vision.analyze_videos_for_radar(all_videos, project=_proj)
                        _result_cache.update(mapping or {})
                    except Exception as e:
                        logger.error(f"❌ 批量视觉分析失败: {e}")

                return _do_resolve, _result_cache

            # 由于 pipeline 内部是逐视频回调，这里改为：先跑一次抓取，再在调用端做视觉分析
            vision_call = None
        except Exception as e:
            logger.warning(f"⚠️ 加载 video_vision 失败，跳过视觉分析: {e}")
            vision_call = None

    # 第一步：先做不带视觉分析的抓取
    try:
        structured = competitor_radar.run_radar_pipeline(
            urls=urls,
            apify_token=APIFY_TOKEN,
            start_date=start_date,
            end_date=end_date,
            enable_vision=False,
            vision_call=None,
            progress=_progress,
        )
    except Exception as e:
        logger.error(f"❌ 雷达抓取失败: {e}")
        update_task(task_id, status='failed', error=f'抓取失败: {e}')
        return

    # 第二步：可选视觉分析（批量 + 写回）
    if enable_video_vision:
        try:
            import video_vision
            all_videos = []
            for prof in structured.get('profiles', []):
                for v in prof.get('videos', []) or []:
                    all_videos.append(v)
            if all_videos:
                total_v = len(all_videos)
                _progress(f'正在做视频内容分析（0/{total_v}）...')
                logger.info(f"🎬 [Radar] 开始批量视觉分析，共 {total_v} 条视频")

                def _vision_progress(done, total):
                    try:
                        _progress(f'正在做视频内容分析（{done}/{total}）...')
                    except Exception:
                        pass

                mapping = video_vision.analyze_videos_for_radar(
                    all_videos, project=project, progress=_vision_progress,
                )
                vis_done = 0
                for prof in structured.get('profiles', []):
                    for v in prof.get('videos', []) or []:
                        info = mapping.get(v.get('video_url') or '') or {}
                        if info.get('marketing_type'):
                            v['marketing_type'] = info['marketing_type']
                        if info.get('vision_summary'):
                            v['vision_summary'] = info['vision_summary']
                        if info.get('marketing_type') or info.get('vision_summary'):
                            vis_done += 1
                structured['total_vision_done'] = vis_done
                logger.info(f"✅ [Radar] 视觉分析完成 {vis_done}/{len(all_videos)}")
        except Exception as e:
            logger.error(f"❌ 视觉分析整体失败: {e}")
            import traceback
            logger.error(traceback.format_exc())

    # 兜底字段
    for prof in structured.get('profiles', []):
        for v in prof.get('videos', []) or []:
            v.setdefault('marketing_type', '')
            v.setdefault('vision_summary', '')

    # 构造 HTML
    summary_html = competitor_radar.build_html_table(structured)
    period_text = ''
    if start_dt_str and end_dt_str:
        period_text = f"<div style='color:#999;font-size:0.9rem;margin-bottom:8px;'>时间范围：{html.escape(start_dt_str)} ~ {html.escape(end_dt_str)}</div>"
    full_html = (
        f"<div style='font-family:sans-serif;'>"
        f"<h3 style='color:#D32F2F;border-bottom:2px solid #eee;padding-bottom:8px;'>📡 竞品雷达 · 数据明细</h3>"
        f"{period_text}{summary_html}</div>"
    )

    # 保存历史记录（含结构化数据，供导出使用）
    title_url = urls[0] if urls else ''
    title = f"竞品雷达:{title_url[20:50]}" if title_url else '竞品雷达'
    structured_payload = {
        'version': competitor_radar.SCHEMA_VERSION,
        'profiles': structured.get('profiles', []),
        'total_videos': structured.get('total_videos', 0),
        'total_vision_done': structured.get('total_vision_done', 0),
        'urls': urls,
        'start_date': start_dt_str,
        'end_date': end_dt_str,
        'project': project,
        'enable_video_vision': enable_video_vision,
    }
    record_id = save_history(user_id, title, full_html, 'competitor', structured=structured_payload)

    # 记录使用成本（按视频条数计）
    if user_id:
        try:
            log_usage(user_id, username, department, 'competitor',
                      structured.get('total_videos', 0), 0, task_id=task_id, record_id=record_id)
        except Exception:
            pass

    update_task(task_id, status='completed', result=full_html, progress='分析完成', record_id=record_id)
    logger.info(f"✅ [Radar] 任务 {task_id} 完成")


def process_competitor_task(task_id, target_url=None, start_dt_str=None, end_dt_str=None,
                            user_id=None, username='unknown', department='未知', session_id='default',
                            project='CFL', generate_report=False, enable_video_vision=False,
                            urls=None):
    """后台处理竞品监控任务。

    新版（urls 参数）走 competitor_radar 流程（TT + IG）；旧版（target_url）保留 TikTok 单链路径。
    """
    # ===== 新路径：竞品雷达 v2（TT + IG 多主页） =====
    if urls:
        try:
            _run_competitor_radar_task(
                task_id=task_id,
                urls=urls,
                start_dt_str=start_dt_str,
                end_dt_str=end_dt_str,
                user_id=user_id,
                username=username,
                department=department,
                session_id=session_id,
                project=project,
                enable_video_vision=enable_video_vision,
            )
        except Exception as e:
            logger.error(f"❌ 竞品雷达任务失败: {e}")
            import traceback
            logger.error(traceback.format_exc())
            update_task(task_id, status='failed', error=str(e))
        return

    # ===== 旧路径（保留，向后兼容） =====
    try:
        logger.info(f"🔄 开始处理竞品监控任务 {task_id}")
        update_task(task_id, status='processing', progress='正在初始化...')

        # 1. 日期转换
        target_start = datetime.datetime.strptime(start_dt_str, '%Y-%m-%d').date()
        target_end = datetime.datetime.strptime(end_dt_str, '%Y-%m-%d').date()
        logger.info(f"📆 解析日期: {target_start} ~ {target_end}")

        # 2. 云端抓取
        logger.info("🕵️ 启动 TikTok 爬虫...")
        update_task(task_id, progress='正在启动 TikTok 爬虫...')

        run_input = {
            "profiles": [target_url],
            "resultsPerPage": 35,
            "oldestPostDate": start_dt_str,
            "shouldDownloadVideos": False
        }

        try:
            # 使用 REST API 启动爬虫
            logger.info("📞 正在调用 Apify REST API...")
            logger.info(f"   Actor: clockworks/tiktok-scraper")
            logger.info(f"   Input: {run_input}")

            api_url = "https://api.apify.com/v2/acts/clockworks~tiktok-scraper/runs"
            headers = {
                "Authorization": f"Bearer {APIFY_TOKEN}",
                "Content-Type": "application/json"
            }

            response = requests.post(
                api_url,
                json=run_input,
                headers=headers,
                timeout=30
            )

            if response.status_code != 201:
                raise ValueError(f"Apify API 返回错误状态码: {response.status_code}, 响应: {response.text}")

            run = response.json()['data']
            logger.info(f"✅ 爬虫任务已启动，Run ID: {run['id']}")
        except requests.Timeout:
            error_msg = "Apify API 调用超时（30秒）"
            logger.error(f"❌ {error_msg}")
            update_task(task_id, status='failed', error=error_msg)
            return
        except Exception as start_error:
            error_msg = f"启动爬虫失败: {str(start_error)}"
            logger.error(f"❌ {error_msg}")
            update_task(task_id, status='failed', error=error_msg)
            return

        # 等待爬虫完成
        logger.info("⏳ 等待爬虫完成...")
        update_task(task_id, progress='等待爬虫完成（约30-60秒）...')

        try:
            logger.info("📡 开始轮询 TikTok 爬虫状态...")
            start_time = time.time()
            max_wait_time = 480  # 最多等待 480 秒
            poll_interval = 5  # 每 5 秒轮询一次

            run_id = run['id']
            api_url = f"https://api.apify.com/v2/actor-runs/{run_id}"
            headers = {"Authorization": f"Bearer {APIFY_TOKEN}"}

            while True:
                elapsed = time.time() - start_time
                if elapsed > max_wait_time:
                    raise TimeoutError(f"等待爬虫完成超时（{max_wait_time}秒）")

                # 轮询任务状态
                logger.info(f"   轮询状态... (已等待 {elapsed:.0f}秒)")
                response = requests.get(api_url, headers=headers, timeout=10)

                if response.status_code != 200:
                    raise ValueError(f"获取任务状态失败: {response.status_code}")

                run_data = response.json()['data']
                status = run_data['status']

                logger.info(f"   当前状态: {status}")

                if status in ['SUCCEEDED', 'FAILED', 'ABORTED', 'TIMED-OUT']:
                    # 任务完成
                    run = run_data
                    break

                # 等待后继续轮询
                time.sleep(poll_interval)

            elapsed = time.time() - start_time
            logger.info(f"✅ 爬虫任务完成，状态: {run['status']}，耗时: {elapsed:.1f}秒")

        except requests.Timeout:
            error_msg = "轮询任务状态超时"
            logger.error(f"❌ {error_msg}")
            update_task(task_id, status='failed', error=error_msg)
            return
        except TimeoutError as timeout_error:
            error_msg = str(timeout_error)
            logger.error(f"❌ {error_msg}")
            update_task(task_id, status='failed', error=error_msg)
            return
        except Exception as wait_error:
            error_msg = f"等待爬虫完成失败: {str(wait_error)}"
            logger.error(f"❌ {error_msg}")
            update_task(task_id, status='failed', error=error_msg)
            return

        if run['status'] != 'SUCCEEDED':
            error_msg = f"爬虫任务失败，状态: {run['status']}"
            logger.error(f"❌ {error_msg}")
            update_task(task_id, status='failed', error=error_msg)
            return

        # 获取数据
        dataset_id = run.get("defaultDatasetId")
        if not dataset_id:
            error_msg = "未找到 dataset ID"
            logger.error(f"❌ {error_msg}")
            update_task(task_id, status='failed', error=error_msg)
            return

        # 使用 REST API 获取数据
        dataset_url = f"https://api.apify.com/v2/datasets/{dataset_id}/items"
        try:
            response = requests.get(dataset_url, headers=headers, timeout=30)
            if response.status_code != 200:
                raise ValueError(f"获取数据失败: {response.status_code}")
            items = response.json()
            logger.info(f"📦 获取到 {len(items)} 条原始数据")
            update_task(task_id, progress=f'已获取 {len(items)} 条数据，正在过滤...')
        except Exception as e:
            error_msg = f"获取数据失败: {str(e)}"
            logger.error(f"❌ {error_msg}")
            update_task(task_id, status='failed', error=error_msg)
            return

        # 3. 本地时间过滤
        cleaned = []
        for it in items:
            raw_date = it.get("createTimeISO")
            if not raw_date:
                continue

            post_dt = datetime.datetime.fromisoformat(raw_date.replace('Z', '+00:00')).date()

            if target_start <= post_dt <= target_end:
                cleaned.append({
                    "desc": it.get("text") or it.get("desc") or "无描述",
                    "likes": it.get("diggCount", 0),
                    "views": it.get("playCount", 0),
                    "comments": it.get("commentCount", 0),
                    "shares": it.get("shareCount", 0),
                    "collects": it.get("collectCount", 0),
                    "url": it.get("webVideoUrl"),
                    "date": str(post_dt),
                    "author": (it.get("authorMeta") or {}).get("name") or (it.get("authorMeta") or {}).get("uniqueId") or "未知"
                })

        logger.info(f"✅ 时间过滤后剩余 {len(cleaned)} 条数据")

        if not cleaned:
            warning_msg = f"在此期间 ({start_dt_str} ~ {end_dt_str}) 未发现视频。"
            logger.info("⚠️ 未发现符合条件的视频")
            update_task(task_id, status='completed', result=f"<div class='alert alert-warning'>{warning_msg}</div>")
            return

        # 4. （可选）调用「看视频」分析：仅在生成报告且开关开启时启用
        video_vision_html = ""
        video_vision_summaries = ""
        video_analysis_results = []
        if generate_report and enable_video_vision:
            try:
                logger.info("🎬 已开启看视频分析，准备对全量视频进行视觉分析...")
                update_task(task_id, progress='正在分析视频内容（全量）...')
                video_analysis_results = analyze_all_videos_for_export(cleaned, project=project)
                if video_analysis_results:
                    video_vision_html, video_vision_summaries = _build_vision_html_from_results(video_analysis_results)
                    logger.info(f"✅ 全量看视频分析完成，共 {len(video_analysis_results)} 条结果")
                else:
                    logger.warning("⚠️ 全量分析未返回结果（可能是直链获取或视觉模型失败）")
            except Exception as vision_error:
                logger.error(f"❌ 看视频分析链路异常: {vision_error}")
                import traceback
                logger.error(traceback.format_exc())
                video_vision_html = ""
                video_vision_summaries = ""
                video_analysis_results = []

        # 5. 预计算聚合指标
        total_count = len(cleaned)
        total_views = sum(item.get("views", 0) for item in cleaned)
        total_likes = sum(item.get("likes", 0) for item in cleaned)
        total_comments = sum(item.get("comments", 0) for item in cleaned)
        total_collects = sum(item.get("collects", 0) for item in cleaned)
        total_shares = sum(item.get("shares", 0) for item in cleaned)
        total_engagement = total_likes + total_comments + total_collects + total_shares
        avg_views = total_views // total_count if total_count else 0
        avg_engagement = total_engagement // total_count if total_count else 0

        # 6. 生成最终输出
        if generate_report:
            competitor_template = get_prompt('competitor', project)
            if not competitor_template:
                update_task(task_id, status='failed', error='该项目提示词尚未配置')
                return

            update_task(task_id, progress='正在生成分析报告...')
            cleaned_str = json.dumps(cleaned, ensure_ascii=False)
            prompt = competitor_template.format(
                cleaned=cleaned_str,
                start_dt_str=start_dt_str,
                end_dt_str=end_dt_str,
                total_count=total_count,
                total_views=total_views,
                avg_views=avg_views,
                total_engagement=total_engagement,
                avg_engagement=avg_engagement,
                video_vision_summaries=video_vision_summaries or ""
            )

            logger.info("🤖 开始调用通义千问 API 生成竞品报告...")
            result, tokens = call_gemini(prompt)
            result = (result or "").replace('```html', '').replace('```', '').strip()

            full_html = f"{result}\n{video_vision_html}"
        else:
            tokens = 0
            overview_html = f"""
            <div style="width:100%; font-family:sans-serif;">
                <h3 style="color:#D32F2F; border-bottom:2px solid #eee; padding-bottom:10px;">
                    📊 数据概览表 ({html.escape(start_dt_str)} 至 {html.escape(end_dt_str)})
                </h3>
                <table class="table" style="width:100%; margin-bottom:30px; text-align:center; font-size:0.9rem;">
                    <tr style="background:#f8f9fa;">
                        <th>总条数</th><th>总播放</th><th>播放均值</th><th>总互动</th><th>互动均值</th><th>总点赞</th><th>总评论</th><th>总收藏</th><th>总转发</th>
                    </tr>
                    <tr>
                        <td>{total_count}</td><td>{total_views}</td><td>{avg_views}</td><td>{total_engagement}</td><td>{avg_engagement}</td><td>{total_likes}</td><td>{total_comments}</td><td>{total_collects}</td><td>{total_shares}</td>
                    </tr>
                </table>
            </div>
            """
            full_html = overview_html

        # 保存历史记录（HTML + 结构化 JSON）
        structured_data = {
            "cleaned": cleaned,
            "video_analysis": video_analysis_results or []
        }
        record_id = save_history(user_id, f"竞品数据:{target_url[20:30]}", full_html, 'competitor', structured=structured_data)

        # 记录使用成本
        if user_id:
            log_usage(
                user_id,
                username,
                department,
                'competitor',
                len(cleaned),  # TikTok 视频数量
                tokens,
                task_id=task_id,
                record_id=record_id
            )

        # 更新任务状态为完成（含 record_id 供前端导出）
        update_task(task_id, status='completed', result=full_html, progress='分析完成', record_id=record_id)

        logger.info("✅ 竞品监控完成")
        logger.info("=" * 60 + "\n")

    except Exception as e:
        error_msg = f"竞品监控失败: {str(e)}"
        logger.error(f"❌ {error_msg}")
        import traceback
        logger.error(traceback.format_exc())
        update_task(task_id, status='failed', error=error_msg)
        return jsonify({'result': f"<div class='alert alert-danger'>{error_msg}</div>"})

# ============================================
# 功能 4: 内容创作（RAG 对话）
# ============================================

@app.route('/chat-tool')
@login_required
def chat_tool():
    """内容创作对话页面"""
    return render_template('chat.html')


# 语料库入口已移除：/corpus-manage 及 /corpus/* 路由已删除


# --- 对话 API ---

@app.route('/chat/send', methods=['POST'])
@login_required
def chat_send():
    """发送消息并获取 AI 回复"""
    try:
        data = request.json or {}
        user_message = (data.get('message') or '').strip()
        if not user_message:
            return jsonify({'error': '消息不能为空'}), 400

        session_id_chat = data.get('session_id')
        mode = data.get('mode', 'copywriting')
        project = data.get('project', 'CFL')

        user_id = session.get('user_id')

        if not session_id_chat:
            title = user_message[:20]
            session_id_chat = db.execute_and_fetch_id(
                """
                INSERT INTO chat_sessions (user_id, mode, project, title)
                VALUES (%s, %s, %s, %s) RETURNING id
                """,
                (user_id, mode, project, title)
            )

        db.execute(
            "INSERT INTO chat_messages (session_id, role, content) VALUES (%s, %s, %s)",
            (session_id_chat, 'user', user_message)
        )

        # 内容创作模式不使用语料检索，仅用固定提示词
        if mode in ('copywriting', 'video_script'):
            chunks = []
        else:
            query_emb = rag.get_embedding(user_message)
            doc_type_filter = mode if mode in ('copywriting', 'video_script') else None
            chunks = rag.search_similar(query_emb, project, doc_type=doc_type_filter, top_k=5)

        history_rows = db.query_all(
            """
            SELECT role, content FROM chat_messages
            WHERE session_id = %s
            ORDER BY created_at ASC
            """,
            (session_id_chat,)
        )
        recent_history = [{'role': r['role'], 'content': r['content']} for r in (history_rows or [])]
        recent_history = recent_history[-10:]

        system_template = get_prompt(mode, project)
        if not system_template:
            if mode == 'copywriting':
                system_template = "你是一位专业的社媒文案策划。请根据用户需求生成优质文案。\n\n【参考语料】:\n{retrieved_context}\n"
            else:
                system_template = "你是一位专业的视频内容策划。请根据用户需求输出视频脚本大纲。\n\n【参考语料】:\n{retrieved_context}\n"

        messages = rag.build_rag_prompt(system_template, chunks, recent_history[:-1], user_message)

        if not qwen_client:
            return jsonify({'error': 'AI 服务未配置'}), 500

        response = qwen_client.chat.completions.create(
            model='qwen-turbo',
            messages=messages,
            temperature=0.7
        )
        ai_reply = response.choices[0].message.content

        db.execute(
            "INSERT INTO chat_messages (session_id, role, content) VALUES (%s, %s, %s)",
            (session_id_chat, 'assistant', ai_reply)
        )

        return jsonify({
            'session_id': session_id_chat,
            'reply': ai_reply,
        })
    except Exception as e:
        logger.error(f"❌ 对话失败: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/chat/history')
@login_required
def chat_history():
    """获取当前用户的所有对话会话"""
    user_id = session.get('user_id')
    try:
        sessions = db.query_all(
            """
            SELECT id, mode, project, title, created_at
            FROM chat_sessions
            WHERE user_id = %s
            ORDER BY created_at DESC
            """,
            (user_id,)
        )
        result = []
        for s in (sessions or []):
            result.append({
                'id': s['id'],
                'mode': s['mode'],
                'project': s['project'],
                'title': s['title'] or '新对话',
                'created_at': s['created_at'].strftime('%Y-%m-%d %H:%M') if s['created_at'] else '',
            })
        return jsonify(result)
    except Exception as e:
        logger.error(f"❌ 获取对话历史失败: {e}")
        return jsonify([])


@app.route('/chat/messages/<int:sid>')
@login_required
def chat_messages(sid):
    """获取某会话的全部消息"""
    user_id = session.get('user_id')
    try:
        owner = db.query_one(
            "SELECT user_id FROM chat_sessions WHERE id = %s", (sid,)
        )
        if not owner or owner['user_id'] != user_id:
            return jsonify({'error': '无权访问'}), 403

        msgs = db.query_all(
            "SELECT role, content, created_at FROM chat_messages WHERE session_id = %s ORDER BY created_at ASC",
            (sid,)
        )
        result = []
        for m in (msgs or []):
            result.append({
                'role': m['role'],
                'content': m['content'],
                'created_at': m['created_at'].strftime('%Y-%m-%d %H:%M') if m['created_at'] else '',
            })
        return jsonify(result)
    except Exception as e:
        logger.error(f"❌ 获取消息失败: {e}")
        return jsonify([])


@app.route('/chat/session/<int:sid>', methods=['DELETE'])
@login_required
def chat_delete_session(sid):
    """删除对话会话"""
    user_id = session.get('user_id')
    try:
        owner = db.query_one(
            "SELECT user_id FROM chat_sessions WHERE id = %s", (sid,)
        )
        if not owner or owner['user_id'] != user_id:
            return jsonify({'error': '无权访问'}), 403
        db.execute("DELETE FROM chat_sessions WHERE id = %s", (sid,))
        return jsonify({'success': True})
    except Exception as e:
        logger.error(f"❌ 删除会话失败: {e}")
        return jsonify({'error': str(e)}), 500


# ============================================
# 功能 3: 视频生成
# ============================================

@app.route('/video-tool')
@login_required
def video_tool():
    """视频生成工具页面"""
    return render_template('video.html')


@app.route('/generate_video', methods=['POST'])
def generate_video():
    """视频生成 API"""
    logger.info("\n" + "=" * 60)
    logger.info("📥 收到视频生成请求")

    try:
        prompt = request.json.get('prompt')
        logger.info(f"🎬 Prompt: {prompt[:50]}...")

        video_url = call_veo_api(prompt)
        save_history(session.get('user_id'), f"视频: {prompt[:10]}", video_url, 'video')

        logger.info("✅ 视频生成完成")
        logger.info("=" * 60 + "\n")

        return jsonify({'video_url': video_url})

    except Exception as e:
        error_msg = f"❌ 视频生成失败: {str(e)}"
        print(error_msg)
        import traceback
        traceback.print_exc()
        return jsonify({'error': error_msg})

# ============================================
# 历史记录管理
# ============================================

@app.route('/get_history')
@login_required
def get_history():
    """获取历史记录（从数据库）"""
    try:
        user_id = session.get('user_id')

        # 从数据库读取用户的历史记录
        records = db.query_all("""
            SELECT id, title, result, type, created_at
            FROM analysis_results
            WHERE user_id = %s
            ORDER BY created_at DESC
            LIMIT 50
        """, (user_id,))

        # 转换为前端需要的格式（含 created_at 供时间筛选）
        result = []
        for record in records:
            created_at = record['created_at']
            result.append({
                'id': record['id'],
                'title': f"{record['title']} [{created_at.strftime('%H:%M')}]",
                'result': record['result'],
                'type': record['type'],
                'created_at': created_at.isoformat() if created_at else None
            })

        return jsonify(result)

    except Exception as e:
        logger.error(f"❌ 获取历史记录失败: {e}")
        # 失败时返回内存中的记录
        return jsonify(HISTORY_DB[::-1])


@app.route('/get_record/<int:id>')
@login_required
def get_record(id):
    """获取单条记录（从数据库）"""
    try:
        user_id = session.get('user_id')

        # 从数据库读取（确保只能读取自己的记录）
        record = db.query_one("""
            SELECT id, title, result, type, created_at
            FROM analysis_results
            WHERE id = %s AND user_id = %s
        """, (id, user_id))

        if record:
            return jsonify({
                'id': record['id'],
                'title': record['title'],
                'result': sanitize_html(record['result'] or ''),
                'type': record['type']
            })
        else:
            return jsonify({'error': '记录不存在'}), 404

    except Exception as e:
        logger.error(f"❌ 获取记录失败: {e}")
        # 失败时从内存查找
        record = next((x for x in HISTORY_DB if x['id'] == id), None)
        return jsonify(record) if record else jsonify({'error': '记录不存在'}), 404


def build_competitor_docx(html_content):
    """将竞品报告 HTML 转为 Word 文档，保留标题、表格、段落；内嵌 base64 图片。"""
    import base64
    doc = Document()
    soup = BeautifulSoup(html_content, 'html.parser')
    # 按文档顺序处理：先所有 h3，再所有 table，再带 base64 的 img；段落用 p 或 div 内文本
    for el in soup.find_all(['h3', 'table', 'div', 'p', 'img']):
        if el.name == 'h3':
            p = doc.add_paragraph()
            run = p.add_run(el.get_text(strip=True))
            run.bold = True
            run.font.size = Pt(14)
        elif el.name == 'table':
            rows = el.find_all('tr')
            if not rows:
                continue
            cols = max(len(r.find_all(['th', 'td'])) for r in rows)
            if cols == 0:
                continue
            table = doc.add_table(rows=len(rows), cols=cols)
            table.style = 'Table Grid'
            for ri, tr in enumerate(rows):
                cells = tr.find_all(['th', 'td'])
                for ci, cell in enumerate(cells):
                    if ci < cols:
                        table.rows[ri].cells[ci].text = cell.get_text(strip=True).replace('\n', ' ')
            doc.add_paragraph()
        elif el.name == 'div' and (el.get('style') or '').find('padding') >= 0:
            for c in el.find_all('p'):
                t = c.get_text(strip=True)
                if t:
                    doc.add_paragraph(t)
            for img in el.find_all('img'):
                src = img.get('src') or ''
                if 'base64,' in src:
                    try:
                        b64 = src.split('base64,', 1)[1]
                        blob = base64.b64decode(b64)
                        doc.add_paragraph()
                        para = doc.add_paragraph()
                        run = para.add_run()
                        run.add_picture(BytesIO(blob), width=Inches(3.5))
                    except Exception:
                        pass
        elif el.name == 'p' and not el.find_parent('table'):
            t = el.get_text(strip=True)
            if t and len(t) > 1:
                doc.add_paragraph(t)
        elif el.name == 'img':
            src = el.get('src') or ''
            if 'base64,' in src:
                try:
                    b64 = src.split('base64,', 1)[1]
                    blob = base64.b64decode(b64)
                    doc.add_paragraph()
                    para = doc.add_paragraph()
                    run = para.add_run()
                    run.add_picture(BytesIO(blob), width=Inches(3.5))
                except Exception:
                    pass
    return doc


@app.route('/export_competitor_word/<int:record_id>')
@login_required
def export_competitor_word(record_id):
    """导出竞品报告为 Word（.docx）。仅限本人且 type=competitor 的记录。"""
    user_id = session.get('user_id')
    record = db.query_one(
        "SELECT id, title, result, type FROM analysis_results WHERE id = %s AND user_id = %s",
        (record_id, user_id)
    )
    if not record or record['type'] != 'competitor':
        return jsonify({'error': '记录不存在或无权导出'}), 404
    try:
        doc = build_competitor_docx(record['result'] or '')
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        filename = f"竞品报告_{record['title'][:20]}_{datetime.datetime.now().strftime('%Y%m%d%H%M')}.docx"
        return send_file(
            buf,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )
    except Exception as e:
        logger.error(f"❌ 导出 Word 失败: {e}")
        return jsonify({'error': str(e)}), 500


def build_competitor_excel(video_results):
    """构建竞品分析 Excel，按情绪分类分组并合并单元格，嵌入关键帧截图。

    表头：情绪 | 目标用户 | 简述 | 发布账号 | 参考图 | 链接
    """
    from openpyxl.drawing.image import Image as XlImage
    import tempfile

    wb = Workbook()
    ws = wb.active
    ws.title = "竞品视频分析"

    # 表头样式
    header_fill = PatternFill(start_color='D32F2F', end_color='D32F2F', fill_type='solid')
    header_font = Font(bold=True, color='FFFFFF', size=11)
    header_align = Alignment(horizontal='center', vertical='center', wrap_text=True)

    headers = ['情绪', '目标用户', '简述', '发布账号', '参考图', '链接']
    col_widths = [14, 14, 40, 16, 18, 40]

    for col_idx, (header, width) in enumerate(zip(headers, col_widths), 1):
        cell = ws.cell(row=1, column=col_idx, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = header_align
        ws.column_dimensions[chr(64 + col_idx)].width = width

    ws.row_dimensions[1].height = 28

    # 按 emotion 排序
    sorted_results = sorted(video_results, key=lambda x: x.get('emotion', '其他'))

    # 写入数据行
    cell_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
    link_font = Font(color='0563C1', underline='single', size=10)
    normal_font = Font(size=10)

    tmp_files = []  # 跟踪临时文件以便清理
    row = 2
    for item in sorted_results:
        ws.cell(row=row, column=1, value=item.get('emotion', '其他')).alignment = cell_align
        ws.cell(row=row, column=1).font = normal_font
        ws.cell(row=row, column=2, value=item.get('target_user', '未知')).alignment = cell_align
        ws.cell(row=row, column=2).font = normal_font
        ws.cell(row=row, column=3, value=item.get('summary', '')).alignment = Alignment(vertical='center', wrap_text=True)
        ws.cell(row=row, column=3).font = normal_font
        ws.cell(row=row, column=4, value=item.get('author', '未知')).alignment = cell_align
        ws.cell(row=row, column=4).font = normal_font

        # 嵌入关键帧截图
        thumb_b64 = item.get('thumbnail_b64', '')
        if thumb_b64:
            try:
                import base64 as _b64
                img_data = _b64.b64decode(thumb_b64)
                fd, tmp_path = tempfile.mkstemp(suffix='.jpg')
                with os.fdopen(fd, 'wb') as f:
                    f.write(img_data)
                tmp_files.append(tmp_path)

                img = XlImage(tmp_path)
                img.width = 120
                img.height = 80
                ws.add_image(img, f'E{row}')
                ws.row_dimensions[row].height = 65
            except Exception as img_err:
                logger.warning(f"⚠️ 嵌入缩略图失败: {img_err}")
                ws.cell(row=row, column=5, value='(图片加载失败)').alignment = cell_align
                ws.row_dimensions[row].height = 30
        else:
            ws.cell(row=row, column=5, value='(无截图)').alignment = cell_align
            ws.row_dimensions[row].height = 30

        # 链接列
        url = item.get('url', '')
        cell = ws.cell(row=row, column=6, value=url)
        cell.alignment = Alignment(vertical='center', wrap_text=True)
        cell.font = link_font
        if url:
            cell.hyperlink = url

        row += 1

    # 合并同情绪分类的第一列单元格
    if len(sorted_results) > 1:
        merge_start = 2
        for i in range(2, row):
            current_emotion = ws.cell(row=i, column=1).value
            next_emotion = ws.cell(row=i + 1, column=1).value if i + 1 < row else None
            if current_emotion != next_emotion:
                if i > merge_start:
                    ws.merge_cells(start_row=merge_start, start_column=1, end_row=i, end_column=1)
                merge_start = i + 1

    # 保存到 BytesIO
    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)

    # 清理临时图片文件
    for tmp_path in tmp_files:
        try:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)
        except Exception:
            pass

    return buf


@app.route('/api/competitor/dashboard/<int:record_id>')
@login_required
def competitor_dashboard_data(record_id):
    """返回竞品看板所需的聚合数据（4 个图表）"""
    user_id = session.get('user_id')
    record = db.query_one(
        "SELECT id, result_json FROM analysis_results WHERE id = %s AND user_id = %s AND type = 'competitor'",
        (record_id, user_id)
    )
    if not record or not record.get('result_json'):
        return jsonify({'error': '无数据'}), 404

    try:
        data = json.loads(record['result_json'])
        # 兼容三种格式：雷达 v2 / 旧版 dict / 旧版 list
        if isinstance(data, dict) and data.get('version') == competitor_radar.SCHEMA_VERSION:
            cleaned = []
            video_analysis = []
            for prof in data.get('profiles', []) or []:
                for v in prof.get('videos', []) or []:
                    cleaned.append({
                        'date': v.get('post_date') or (v.get('post_time') or '')[:10],
                        'views': v.get('views', 0),
                        'likes': v.get('likes', 0),
                        'comments': v.get('comments', 0),
                        'collects': v.get('collects', 0),
                        'shares': v.get('shares', 0),
                        'desc': v.get('caption', ''),
                        'url': v.get('video_url', ''),
                    })
                    if v.get('marketing_type') or v.get('vision_summary'):
                        video_analysis.append({
                            'emotion': v.get('marketing_type', '') or '其他',
                            'summary': v.get('vision_summary', ''),
                            'url': v.get('video_url', ''),
                        })
        elif isinstance(data, dict):
            cleaned = data.get("cleaned", [])
            video_analysis = data.get("video_analysis", [])
        else:
            cleaned = []
            video_analysis = data if isinstance(data, list) else []

        # 1. 发布趋势（按日期聚合）
        from collections import Counter, OrderedDict
        date_counts = Counter()
        date_views = Counter()
        for v in cleaned:
            d = v.get("date", "未知")
            date_counts[d] += 1
            date_views[d] += v.get("views", 0)
        sorted_dates = sorted(date_counts.keys())
        trend = [{"date": d, "count": date_counts[d], "views": date_views[d]} for d in sorted_dates]

        # 2. 互动指标汇总
        engagement = {
            "views": sum(v.get("views", 0) for v in cleaned),
            "likes": sum(v.get("likes", 0) for v in cleaned),
            "comments": sum(v.get("comments", 0) for v in cleaned),
            "collects": sum(v.get("collects", 0) for v in cleaned),
            "shares": sum(v.get("shares", 0) for v in cleaned),
        }

        # 3. 情绪分布
        emotion_counts = Counter(v.get("emotion", "未知") for v in video_analysis)
        emotions = [{"label": k, "count": v} for k, v in emotion_counts.items()]

        # 4. Top 10 视频（按播放量）
        sorted_videos = sorted(cleaned, key=lambda x: x.get("views", 0), reverse=True)[:10]
        top_videos = [{"desc": v.get("desc", "")[:30], "views": v.get("views", 0), "url": v.get("url", "")} for v in sorted_videos]

        return jsonify({"trend": trend, "engagement": engagement, "emotions": emotions, "top_videos": top_videos})
    except Exception as e:
        logger.error(f"❌ 看板数据聚合失败: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/export_radar_excel/<int:record_id>')
@login_required
def export_radar_excel(record_id):
    """导出竞品雷达 Excel：12 列表头，按主页分 sheet。"""
    user_id = session.get('user_id')
    record = db.query_one(
        "SELECT id, title, result_json, type FROM analysis_results WHERE id = %s AND user_id = %s",
        (record_id, user_id)
    )
    if not record or record['type'] != 'competitor':
        return jsonify({'error': '记录不存在或无权导出'}), 404
    result_json = record.get('result_json')
    if not result_json:
        return jsonify({'error': '该记录不含结构化数据，无法导出 Excel'}), 400
    try:
        structured = json.loads(result_json)
        if not isinstance(structured, dict) or structured.get('version') != competitor_radar.SCHEMA_VERSION:
            return jsonify({'error': '该记录为旧版数据，请使用「旧版导出」按钮'}), 400
        wb = competitor_radar.build_excel(structured)
        buf = BytesIO()
        wb.save(buf)
        buf.seek(0)
        filename = f"竞品雷达_{datetime.datetime.now().strftime('%Y%m%d_%H%M')}.xlsx"
        return send_file(
            buf,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=filename,
        )
    except Exception as e:
        logger.error(f"❌ 导出雷达 Excel 失败: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return jsonify({'error': str(e)}), 500


@app.route('/export_competitor_excel/<int:record_id>')
@login_required
def export_competitor_excel(record_id):
    """导出竞品分析 Excel（含情绪分类+关键帧截图）"""
    user_id = session.get('user_id')
    record = db.query_one(
        "SELECT id, title, result_json, type FROM analysis_results WHERE id = %s AND user_id = %s",
        (record_id, user_id)
    )
    if not record or record['type'] != 'competitor':
        return jsonify({'error': '记录不存在或无权导出'}), 404

    result_json = record.get('result_json')
    if not result_json:
        return jsonify({'error': '该报告无视频分析数据，无法导出 Excel。请重新生成报告并开启「看视频分析」。'}), 400

    try:
        video_results = json.loads(result_json)
        # 新版雷达：转发到雷达导出
        if isinstance(video_results, dict) and video_results.get('version') == competitor_radar.SCHEMA_VERSION:
            return export_radar_excel(record_id)
        # 兼容新格式（dict 含 cleaned + video_analysis）和旧格式（直接是 list）
        if isinstance(video_results, dict):
            video_results = video_results.get("video_analysis", [])
        if not video_results:
            return jsonify({'error': '视频分析数据为空'}), 400

        buf = build_competitor_excel(video_results)
        filename = f"竞品Excel_{record['title'][:20]}_{datetime.datetime.now().strftime('%Y%m%d%H%M')}.xlsx"
        return send_file(
            buf,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=filename
        )
    except Exception as e:
        logger.error(f"❌ 导出竞品 Excel 失败: {e}")
        return jsonify({'error': str(e)}), 500


# ============================================
# Excel 导出功能
# ============================================

def create_excel_by_language(results):
    """按语言分类生成 Excel"""
    wb = Workbook()
    wb.remove(wb.active)  # 删除默认 sheet

    # 按语言分组
    language_groups = {}
    for item in results:
        lang = item.get('language', '其他')
        if lang not in language_groups:
            language_groups[lang] = []
        language_groups[lang].append(item)

    # 为每个语言创建 Sheet
    for lang, items in sorted(language_groups.items()):
        ws = wb.create_sheet(title=lang)

        # 表头
        headers = ['序号', '原始评论', '归类', '情感倾向', '语言', '简要分析']
        ws.append(headers)

        # 设置表头样式
        header_fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
        header_font = Font(bold=True, color='FFFFFF')
        for cell in ws[1]:
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center', vertical='center')

        # 添加数据
        for idx, item in enumerate(items, 1):
            ws.append([
                idx,
                item.get('text', ''),
                item.get('category', ''),
                item.get('sentiment', ''),
                item.get('language', ''),
                item.get('analysis', '')
            ])

        # 设置列宽
        ws.column_dimensions['A'].width = 8
        ws.column_dimensions['B'].width = 50
        ws.column_dimensions['C'].width = 20
        ws.column_dimensions['D'].width = 12
        ws.column_dimensions['E'].width = 10
        ws.column_dimensions['F'].width = 40

        # 冻结首行
        ws.freeze_panes = 'A2'

    return wb


def create_excel_by_category(results):
    """按分类生成 Excel"""
    wb = Workbook()
    wb.remove(wb.active)

    # 按分类分组
    category_groups = {}
    for item in results:
        cat = item.get('category', '其他')
        if cat not in category_groups:
            category_groups[cat] = []
        category_groups[cat].append(item)

    # 分类顺序和 Sheet 名称映射（去掉特殊字符）
    category_mapping = {
        '外挂作弊': '外挂作弊',
        '游戏优化': '游戏优化',
        '游戏Bug': '游戏Bug',
        '充值退款': '充值退款',
        '新模式/地图/平衡性建议': '新模式地图平衡性建议'  # 去掉斜杠
    }

    # 为每个分类创建 Sheet
    for cat, sheet_name in category_mapping.items():
        if cat not in category_groups:
            continue

        items = category_groups[cat]
        ws = wb.create_sheet(title=sheet_name)

        # 表头
        headers = ['序号', '原始评论', '归类', '情感倾向', '语言', '简要分析']
        ws.append(headers)

        # 设置表头样式
        header_fill = PatternFill(start_color='D32F2F', end_color='D32F2F', fill_type='solid')
        header_font = Font(bold=True, color='FFFFFF')
        for cell in ws[1]:
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center', vertical='center')

        # 添加数据
        for idx, item in enumerate(items, 1):
            ws.append([
                idx,
                item.get('text', ''),
                item.get('category', ''),
                item.get('sentiment', ''),
                item.get('language', ''),
                item.get('analysis', '')
            ])

        # 设置列宽
        ws.column_dimensions['A'].width = 8
        ws.column_dimensions['B'].width = 50
        ws.column_dimensions['C'].width = 20
        ws.column_dimensions['D'].width = 12
        ws.column_dimensions['E'].width = 10
        ws.column_dimensions['F'].width = 40

        # 冻结首行
        ws.freeze_panes = 'A2'

    return wb


@app.route('/export_insight')
@login_required
def export_insight():
    """v2 舆情洞察 Excel 导出（按图中 12 列表头）。

    需要前端传入 record_id；记录里应包含 v2 结构化数据（_schema=insight_v2）。
    若发现是旧版 6 列数据，自动重定向到 /export_by_category。
    """
    user_id = session.get('user_id')
    record_id = request.args.get('record_id', type=int)
    if not record_id:
        return jsonify({'error': '缺少 record_id 参数'}), 400

    try:
        record = db.query_one(
            """
            SELECT result_json FROM analysis_results
            WHERE id = %s AND user_id = %s AND type = %s
            """,
            (record_id, user_id, 'sentiment'),
        )
        if not record:
            return jsonify({'error': '记录不存在或无权限访问'}), 404
        if not (ANALYSIS_RESULTS_HAS_JSON and record.get('result_json')):
            return jsonify({'error': '该历史记录无结构化数据，无法导出新版表格'}), 400

        try:
            results = json.loads(record['result_json'])
        except Exception as e:
            logger.error(f"❌ 解析 result_json 失败: {e}")
            return jsonify({'error': '该记录原始数据已损坏'}), 500

        if not isinstance(results, list) or not results:
            return jsonify({'error': '没有可导出的数据'}), 400

        is_v2 = isinstance(results[0], dict) and (
            results[0].get('_schema') == sentiment_insight.SCHEMA_VERSION
            or 'platform' in results[0]
        )
        if not is_v2:
            return jsonify({
                'error': '该记录是旧版（6 列）格式，请使用「按分类导出」或「按语言导出」'
            }), 400

        wb = sentiment_insight.build_excel(results)
        output = BytesIO()
        wb.save(output)
        output.seek(0)
        ts = datetime.datetime.now().strftime('%Y%m%d_%H%M')
        filename = f'舆情洞察_{ts}.xlsx'
        logger.info(f"📥 导出舆情洞察 v2: {len(results)} 行")
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=filename,
        )
    except Exception as e:
        logger.error(f"❌ 导出舆情洞察失败: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/export_by_language')
@login_required
def export_by_language():
    """按语言导出 Excel

    优先根据前端传入的 record_id 导出「当前查看的那一条」历史记录；
    若未提供 record_id，则退回到旧逻辑：导出当前会话最近一次分析结果。
    """
    user_id = session.get('user_id')
    record_id = request.args.get('record_id', type=int)

    results = []

    try:
        if record_id:
            # 新逻辑：按记录 ID 精确导出当前查看的历史记录
            logger.info(f"📥 按记录ID导出语言分类报告: record_id={record_id}, user_id={user_id}")
            record = db.query_one("""
                SELECT result_json
                FROM analysis_results
                WHERE id = %s AND user_id = %s AND type = %s
            """, (record_id, user_id, 'sentiment'))

            if not record:
                return jsonify({'error': '记录不存在或无权限访问'}), 404

            if ANALYSIS_RESULTS_HAS_JSON and record.get('result_json'):
                try:
                    results = json.loads(record['result_json'])
                except Exception as e:
                    logger.error(f"❌ 解析 result_json 失败: {e}")
                    return jsonify({'error': '该记录的原始数据已损坏，无法导出'}), 500
            else:
                return jsonify({'error': '该历史记录生成于旧版本，暂不支持导出，请重新分析一次'}), 400
        else:
            # 兼容旧逻辑：按当前会话最近一次分析导出
            session_id = session.get('session_id', 'default')
            results = LATEST_ANALYSIS_RESULTS.get(session_id, [])
            logger.info(f"📥 按会话导出语言分类报告: session_id={session_id}, count={len(results)}")

        if not results:
            return jsonify({'error': '没有可导出的数据'}), 400

        wb = create_excel_by_language(results)

        # 生成文件
        output = BytesIO()
        wb.save(output)
        output.seek(0)

        # 生成文件名
        timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M')
        filename = f'语言分类报告_{timestamp}.xlsx'

        logger.info(f"📥 导出语言分类报告: {len(results)} 条数据")

        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=filename
        )
    except Exception as e:
        logger.error(f"❌ 导出失败: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/export_by_category')
@login_required
def export_by_category():
    """按分类导出 Excel

    优先根据前端传入的 record_id 导出「当前查看的那一条」历史记录；
    若未提供 record_id，则退回到旧逻辑：导出当前会话最近一次分析结果。
    """
    user_id = session.get('user_id')
    record_id = request.args.get('record_id', type=int)

    results = []

    try:
        if record_id:
            # 新逻辑：按记录 ID 精确导出当前查看的历史记录
            logger.info(f"📥 按记录ID导出问题分类报告: record_id={record_id}, user_id={user_id}")
            record = db.query_one("""
                SELECT result_json
                FROM analysis_results
                WHERE id = %s AND user_id = %s AND type = %s
            """, (record_id, user_id, 'sentiment'))

            if not record:
                return jsonify({'error': '记录不存在或无权限访问'}), 404

            if ANALYSIS_RESULTS_HAS_JSON and record.get('result_json'):
                try:
                    results = json.loads(record['result_json'])
                except Exception as e:
                    logger.error(f"❌ 解析 result_json 失败: {e}")
                    return jsonify({'error': '该记录的原始数据已损坏，无法导出'}), 500
            else:
                return jsonify({'error': '该历史记录生成于旧版本，暂不支持导出，请重新分析一次'}), 400
        else:
            # 兼容旧逻辑：按当前会话最近一次分析导出
            session_id = session.get('session_id', 'default')
            results = LATEST_ANALYSIS_RESULTS.get(session_id, [])
            logger.info(f"📥 按会话导出问题分类报告: session_id={session_id}, count={len(results)}")

        if not results:
            return jsonify({'error': '没有可导出的数据'}), 400

        wb = create_excel_by_category(results)

        # 生成文件
        output = BytesIO()
        wb.save(output)
        output.seek(0)

        # 生成文件名
        timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M')
        filename = f'问题分类报告_{timestamp}.xlsx'

        logger.info(f"📥 导出问题分类报告: {len(results)} 条数据")

        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=filename
        )
    except Exception as e:
        logger.error(f"❌ 导出失败: {e}")
        return jsonify({'error': str(e)}), 500

# ============================================
# 统计功能
# ============================================

@app.route('/my-stats')
@login_required
def my_stats():
    """个人统计页面"""
    user_id = session.get('user_id')

    # 获取本月统计数据
    current_month = datetime.datetime.now().strftime('%Y-%m')

    stats_data = db.query_one("""
        SELECT
            COUNT(*) as count,
            COALESCE(SUM(comments_count), 0) as comments,
            COALESCE(SUM(total_cost), 0) as cost
        FROM usage_logs
        WHERE user_id = %s
          AND TO_CHAR(created_at, 'YYYY-MM') = %s
    """, (user_id, current_month))

    stats = {
        'count': stats_data['count'] if stats_data else 0,
        'comments': stats_data['comments'] if stats_data else 0,
        'cost': float(stats_data['cost']) if stats_data else 0.0,
        'avg_cost': float(stats_data['cost']) / stats_data['count'] if stats_data and stats_data['count'] > 0 else 0.0
    }

    # 获取最近20条使用记录
    logs = db.query_all("""
        SELECT *
        FROM usage_logs
        WHERE user_id = %s
        ORDER BY created_at DESC
        LIMIT 20
    """, (user_id,))

    return render_template('my_stats.html', stats=stats, logs=logs, user=session)


@app.route('/admin')
@admin_required
def admin_panel():
    """管理后台页面"""
    current_month = datetime.datetime.now().strftime('%Y-%m')

    # 全局统计
    global_data = db.query_one("""
        SELECT
            COALESCE(SUM(total_cost), 0) as total_cost,
            COUNT(DISTINCT user_id) as active_users,
            COUNT(*) as total_count
        FROM usage_logs
        WHERE TO_CHAR(created_at, 'YYYY-MM') = %s
    """, (current_month,))

    global_stats = {
        'total_cost': float(global_data['total_cost']) if global_data else 0.0,
        'active_users': global_data['active_users'] if global_data else 0,
        'total_count': global_data['total_count'] if global_data else 0,
        'avg_cost': float(global_data['total_cost']) / global_data['total_count'] if global_data and global_data['total_count'] > 0 else 0.0
    }

    # 部门统计
    dept_stats = db.query_all("""
        SELECT
            department,
            COALESCE(SUM(total_cost), 0) as cost,
            COUNT(*) as count
        FROM usage_logs
        WHERE TO_CHAR(created_at, 'YYYY-MM') = %s
        GROUP BY department
        ORDER BY cost DESC
    """, (current_month,))

    # 用户统计（Top 10）
    user_stats = db.query_all("""
        SELECT
            u.real_name,
            u.department,
            COALESCE(SUM(l.total_cost), 0) as cost,
            COUNT(l.id) as count
        FROM users u
        LEFT JOIN usage_logs l ON u.id = l.user_id
            AND TO_CHAR(l.created_at, 'YYYY-MM') = %s
        GROUP BY u.id, u.real_name, u.department
        ORDER BY cost DESC
        LIMIT 10
    """, (current_month,))

    # 所有用户列表
    all_users = db.query_all("""
        SELECT * FROM users ORDER BY created_at DESC
    """)

    return render_template('admin.html',
                         global_stats=global_stats,
                         dept_stats=dept_stats,
                         user_stats=user_stats,
                         all_users=all_users,
                         user=session)


@app.route('/admin/add_user', methods=['POST'])
@admin_required
def add_user():
    """添加新用户"""
    try:
        data = request.json
        username = data.get('username')
        password = data.get('password')
        real_name = data.get('real_name')
        department = data.get('department')
        role = data.get('role', 'user')

        # 检查用户名是否已存在
        existing = db.query_one("SELECT id FROM users WHERE username = %s", (username,))
        if existing:
            return jsonify({'error': '用户名已存在'}), 400

        # 加密密码
        password_hash = bcrypt.generate_password_hash(password).decode('utf-8')

        # 插入用户
        db.execute("""
            INSERT INTO users (username, password_hash, real_name, department, role)
            VALUES (%s, %s, %s, %s, %s)
        """, (username, password_hash, real_name, department, role))

        logger.info(f"✅ 管理员添加新用户: {username} ({real_name})")

        return jsonify({'success': True})

    except Exception as e:
        logger.error(f"❌ 添加用户失败: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/admin/delete_user/<int:user_id>', methods=['DELETE'])
@admin_required
def delete_user(user_id):
    """删除用户"""
    try:
        # 不允许删除管理员账号
        user = db.query_one("SELECT username FROM users WHERE id = %s", (user_id,))
        if user and user['username'] == 'admin':
            return jsonify({'error': '不能删除管理员账号'}), 403

        # 删除用户
        db.execute("DELETE FROM users WHERE id = %s", (user_id,))

        logger.info(f"✅ 管理员删除用户: ID={user_id}")

        return jsonify({'success': True})

    except Exception as e:
        logger.error(f"❌ 删除用户失败: {e}")
        return jsonify({'error': str(e)}), 500


# ============================================
# FB 舆情分析新接口
# ============================================

@app.route('/fb_dashboard')
@login_required
def fb_dashboard():
    """FB 舆情看板页面"""
    return render_template('fb_dashboard.html', user=session)


@app.route('/fb_schedule', methods=['POST'])
@login_required
def fb_schedule():
    """手动触发 FB 评论抓取（异步执行）"""
    try:
        if not USE_DB_WORKER:
            # 进程内线程执行（旧模式）
            # 创建任务记录
            task_id = db.execute_and_fetch_id(
                "INSERT INTO scrape_tasks (task_type, status) VALUES (%s, %s) RETURNING id",
                ('fb_scrape', 'pending')
            )

            # 在后台线程中执行抓取任务
            def run_scrape():
                try:
                    result = tasks.scrape_fb_comments(
                        task_id=task_id,
                        results_limit=2500,
                        enable_ai_analysis=True,
                        max_ai_comments=1200
                    )
                    logger.info(f"✅ 后台抓取完成 (task_id={task_id}): {result}")
                except Exception as e:
                    logger.error(f"❌ 后台抓取失败 (task_id={task_id}): {e}")
                    try:
                        db.execute(
                            "UPDATE scrape_tasks SET status = 'failed', completed_at = NOW(), error_message = %s WHERE id = %s",
                            (str(e), task_id)
                        )
                    except:
                        pass

            thread = threading.Thread(target=run_scrape, daemon=True)
            thread.start()
        else:
            # DB worker 模式：写入 task_queue，由独立 worker 拾取
            task_id_str = str(uuid.uuid4())
            user_id = session.get('user_id')
            session_id = session.get('session_id', 'default')
            create_task(task_id_str, user_id, session_id, function_type='fb_scrape')

            # 同时在 scrape_tasks 表也创建一条记录（保持前端轮询兼容）
            task_id = db.execute_and_fetch_id(
                "INSERT INTO scrape_tasks (task_type, status) VALUES (%s, %s) RETURNING id",
                ('fb_scrape', 'pending')
            )
            # 把 scrape_task_id 与参数存入 task_params 方便 worker 更新
            try:
                task_params = {
                    'source': 'fb_schedule',
                    'scrape_task_id': task_id,
                    'days_back': 7,
                    'results_limit': 2500,
                    'enable_ai_analysis': True,
                    'max_ai_comments': 1200
                }
                db.execute(
                    "UPDATE task_queue SET task_params = %s WHERE task_id = %s",
                    (json.dumps(task_params, ensure_ascii=False), task_id_str)
                )
            except Exception as e:
                logger.error(f"❌ 更新 task_params 失败: {e}")

            logger.info(f"✅ FB 抓取任务已创建 (task_queue={task_id_str}, scrape_tasks={task_id})，等待 worker 处理")

        return jsonify({
            'status': 'success',
            'message': '抓取任务已启动，请稍后刷新页面查看结果',
            'task_id': task_id
        })
    except Exception as e:
        logger.error(f"❌ FB 抓取启动失败: {e}")
        return jsonify({'status': 'error', 'message': str(e)}), 500


@app.route('/api/spd_schedule', methods=['POST'])
@login_required
def spd_schedule():
    """触发 SPD 专项抓取任务（支持试跑参数控制）"""
    try:
        def _normalize_string_list(raw_value):
            if raw_value is None:
                return []
            if isinstance(raw_value, str):
                raw_value = [raw_value]
            if not isinstance(raw_value, list):
                return []
            cleaned = []
            seen = set()
            for item in raw_value:
                if not isinstance(item, str):
                    continue
                val = item.strip()
                if not val:
                    continue
                key = val.lower()
                if key in seen:
                    continue
                seen.add(key)
                cleaned.append(val)
            return cleaned

        data = request.get_json(silent=True) or {}
        try:
            days_back = int(data.get('days_back', 14))
            results_limit = int(data.get('results_limit', 10000))
            max_ai_comments = int(data.get('max_ai_comments', 5000))
            discover_max_posts = int(data.get('discover_max_posts', 5000))
        except (TypeError, ValueError):
            return jsonify({'status': 'error', 'message': '数值参数格式错误'}), 400
        enable_ai_raw = data.get('enable_ai_analysis', True)
        if isinstance(enable_ai_raw, str):
            enable_ai_analysis = enable_ai_raw.strip().lower() in ('1', 'true', 'yes', 'y', 'on')
        else:
            enable_ai_analysis = bool(enable_ai_raw)
        post_urls = _normalize_string_list(data.get('post_urls')) or None
        # 兼容旧字段 seed_tags（默认视为 SPD 任务标签）
        seed_tags = _normalize_string_list(data.get('seed_tags'))
        mlbb_seed_tags = _normalize_string_list(data.get('mlbb_seed_tags'))
        spd_seed_tags = _normalize_string_list(data.get('spd_seed_tags'))
        mlbb_rule = (data.get('mlbb_rule') or '').strip()
        spd_rule = (data.get('spd_rule') or '').strip()
        if seed_tags and not spd_seed_tags:
            spd_seed_tags = list(seed_tags)
        platforms = _normalize_string_list(data.get('platforms')) or ['facebook', 'instagram']
        crawl_scope = str(data.get('crawl_scope', 'both')).strip().lower()
        if crawl_scope not in ('both', 'mlbb', 'spd'):
            return jsonify({'status': 'error', 'message': 'crawl_scope 仅支持 both/mlbb/spd'}), 400

        task_queries = []
        if not post_urls:
            if crawl_scope in ('both', 'mlbb'):
                task_queries.append({
                    'task_name': 'MLBB',
                    'seed_tags': mlbb_seed_tags or list(MLBB_DISCOVER_KEYWORDS),
                    'boolean_rule': mlbb_rule or TASK_BOOLEAN_RULES.get('MLBB', '')
                })
            if crawl_scope in ('both', 'spd'):
                task_queries.append({
                    'task_name': 'SPD',
                    'seed_tags': spd_seed_tags or list(SPD_KEYWORDS),
                    'boolean_rule': spd_rule or TASK_BOOLEAN_RULES.get('SPD', '')
                })

        if days_back < 1 or days_back > 60:
            return jsonify({'status': 'error', 'message': 'days_back 需在 1~60 之间'}), 400
        if results_limit < 100 or results_limit > 10000:
            return jsonify({'status': 'error', 'message': 'results_limit 需在 100~10000 之间'}), 400
        if max_ai_comments < 0 or max_ai_comments > 5000:
            return jsonify({'status': 'error', 'message': 'max_ai_comments 需在 0~5000 之间'}), 400
        if discover_max_posts < 20 or discover_max_posts > 5000:
            return jsonify({'status': 'error', 'message': 'discover_max_posts 需在 20~5000 之间'}), 400

        task_id = db.execute_and_fetch_id(
            "INSERT INTO scrape_tasks (task_type, status) VALUES (%s, %s) RETURNING id",
            ('spd_scrape', 'pending')
        )

        skip_discover = bool(data.get('skip_discover', False))
        top_n = max(1, min(int(data.get('top_n', 5)), 20))

        if not USE_DB_WORKER:
            def run_spd_scrape():
                try:
                    resolved_urls = post_urls
                    resolved_posts = []

                    if skip_discover and not resolved_urls:
                        end_dt = datetime.datetime.now()
                        start_dt = end_dt - datetime.timedelta(days=days_back)
                        top_rows = db.query_all("""
                            SELECT post_url, engagement, comments_count
                            FROM fb_post_metrics
                            WHERE post_date >= %s AND post_date <= %s
                              AND COALESCE(comments_count, 0) > 0
                            ORDER BY comments_count DESC
                            LIMIT %s
                        """, (start_dt.strftime('%Y-%m-%d'), end_dt.strftime('%Y-%m-%d'), top_n))
                        if not top_rows:
                            raise RuntimeError(f'fb_post_metrics 中无 {start_dt.date()}~{end_dt.date()} 的帖子，请先运行 discover')
                        resolved_urls = [r['post_url'] for r in top_rows]
                        logger.info(f"📊 skip_discover: Top{top_n} by comments_count: {[(r.get('post_url','')[-30:], r.get('comments_count'), r.get('engagement')) for r in top_rows]}")

                    elif (not resolved_urls) and task_queries:
                        merged = []
                        merged_posts = []
                        seen = set()
                        for query in task_queries:
                            discover_result = tasks.discover_posts_by_tags(
                                seed_tags=query.get('seed_tags') or [],
                                platforms=platforms,
                                days_back=days_back,
                                max_posts=discover_max_posts,
                                boolean_rule=query.get('boolean_rule')
                            )
                            if discover_result.get('status') != 'success':
                                raise RuntimeError(f"{query.get('task_name')} discover failed: {discover_result.get('message')}")
                            merged_posts.extend(discover_result.get('posts') or [])
                            for url in (discover_result.get('post_urls') or []):
                                key = str(url).strip().lower()
                                if not key or key in seen:
                                    continue
                                seen.add(key)
                                merged.append(url)
                        resolved_urls = merged
                        resolved_posts = merged_posts
                        if not resolved_urls:
                            raise RuntimeError('MLBB/SPD 分任务 discover 均未找到可抓取帖子')

                    tasks.scrape_fb_comments(
                        post_urls=resolved_urls,
                        discovered_posts=resolved_posts,
                        days_back=days_back,
                        task_id=task_id,
                        results_limit=results_limit,
                        enable_ai_analysis=enable_ai_analysis,
                        max_ai_comments=max_ai_comments,
                        allow_fallback_to_config=False
                    )
                except Exception as e:
                    logger.error(f"❌ SPD 抓取失败(task_id={task_id}): {e}")
                    try:
                        db.execute(
                            "UPDATE scrape_tasks SET status = 'failed', completed_at = NOW(), error_message = %s WHERE id = %s",
                            (str(e)[:500], task_id)
                        )
                    except Exception:
                        pass

            threading.Thread(target=run_spd_scrape, daemon=True).start()
        else:
            queue_task_id = str(uuid.uuid4())
            create_task(
                queue_task_id,
                session.get('user_id'),
                session.get('session_id', 'default'),
                function_type='fb_scrape'
            )
            task_params = {
                'source': 'spd_schedule',
                'scrape_task_id': task_id,
                'post_urls': post_urls,
                'seed_tags': seed_tags,
                'task_queries': task_queries,
                'platforms': platforms,
                'crawl_scope': crawl_scope,
                'discover_max_posts': discover_max_posts,
                'days_back': days_back,
                'results_limit': results_limit,
                'enable_ai_analysis': enable_ai_analysis,
                'max_ai_comments': max_ai_comments,
                'skip_discover': skip_discover,
                'top_n': top_n
            }
            db.execute(
                "UPDATE task_queue SET task_params = %s WHERE task_id = %s",
                (json.dumps(task_params, ensure_ascii=False), queue_task_id)
            )

        return jsonify({
            'status': 'success',
            'message': 'SPD 抓取任务已启动',
            'task_id': task_id,
            'config': {
                'crawl_scope': crawl_scope,
                'days_back': days_back,
                'results_limit': results_limit,
                'enable_ai_analysis': enable_ai_analysis,
                'max_ai_comments': max_ai_comments,
                'discover_max_posts': discover_max_posts,
                'mlbb_seed_tag_count': len(mlbb_seed_tags),
                'spd_seed_tag_count': len(spd_seed_tags or seed_tags),
                'mlbb_rule_set': bool(mlbb_rule),
                'spd_rule_set': bool(spd_rule),
                'task_query_count': len(task_queries),
                'custom_post_count': len(post_urls) if post_urls else 0
            }
        })
    except Exception as e:
        logger.error(f"❌ SPD 抓取任务启动失败: {e}")
        return jsonify({'status': 'error', 'message': '内部错误，请稍后重试'}), 500


@app.route('/fb_task_status/<int:task_id>', methods=['GET'])
@login_required
def fb_task_status(task_id):
    """查询抓取任务状态"""
    try:
        task = db.query_one("SELECT * FROM scrape_tasks WHERE id = %s", (task_id,))

        if not task:
            return jsonify({'error': '任务不存在'}), 404

        return jsonify({
            'status': task['status'],
            'started_at': task['started_at'].isoformat() if task.get('started_at') else None,
            'completed_at': task['completed_at'].isoformat() if task.get('completed_at') else None,
            'result_summary': task.get('result_summary'),
            'error_message': task.get('error_message')
        })

    except Exception as e:
        logger.error(f"❌ 查询任务状态失败: {e}")
        return jsonify({'error': str(e)}), 500


FB_QUERY_SEPARATOR_PATTERN = re.compile(r"[\s,，、/|]+")
FB_CHINESE_TOKEN_PATTERN = re.compile(r"[\u4e00-\u9fff]{2,}")
FB_LATIN_TOKEN_PATTERN = re.compile(r"[A-Za-z][A-Za-z0-9_+-]{1,}")
FB_TRANSLATE_CACHE = {}
FB_MAX_TRANSLATE_TERMS = 40
FB_WORDCLOUD_MAX_ROWS = 1500
FB_WORDCLOUD_MAX_TERMS = 100
FB_WORD_MIN_CHARS = 2
FB_QUERY_MAX_TERMS = 25
FB_QUERY_SQL_MAX_TERMS = 20
FB_TRANSLATE_CACHE_MAX_SIZE = 3000
FB_LATIN_MIN_CHARS = 3
FB_CHINESE_MAX_CHARS = 8
FB_WORD_MIN_COUNT_DEFAULT = 2

# 业务同义词词典（可按运营反馈持续扩充）
FB_SYNONYM_GROUPS = [
    {
        "canonical": "版本更新",
        "terms": ["版本更新", "新版本", "版更", "版本迭代", "更新后", "update", "new patch", "patch", "latest version"]
    },
    {
        "canonical": "外挂",
        "terms": ["外挂", "开挂", "作弊", "脚本", "hack", "hacker", "hacks", "cheat", "cheater", "cheating"]
    },
    {
        "canonical": "bug",
        "terms": ["bug", "bugs", "错误", "异常", "故障", "闪退", "crash", "crashes", "glitch", "glitches"]
    },
    {
        "canonical": "卡顿",
        "terms": ["卡顿", "掉帧", "延迟", "lag", "stutter", "latency", "fps drop", "freeze"]
    }
]

FB_STOP_WORDS = {
    "这个", "那个", "我们", "你们", "他们", "就是", "真的", "感觉", "还是", "然后", "已经", "现在", "可以",
    "非常", "比较", "不是", "没有", "一个", "一下", "因为", "所以", "什么", "怎么", "the", "and", "for",
    "with", "this", "that", "you", "your", "are", "was", "were", "from", "have", "has", "had", "but", "not",
    "too", "very", "pls", "please", "game", "player", "players",
    "ng", "na", "th", "ko", "mo", "sa", "pa", "ba", "nga", "ako", "lang", "cfl",
    "年轻", "哈哈", "记录", "队伍", "时间", "所有", "只是", "还有", "一下", "一个", "一些", "已经"
}

FB_WORD_NORMALIZATION_MAP = {
    "update": "版本更新",
    "patch": "版本更新",
    "new patch": "版本更新",
    "latest version": "版本更新",
    "bug": "bug",
    "bugs": "bug",
    "glitch": "bug",
    "glitches": "bug",
    "hack": "外挂",
    "hacker": "外挂",
    "hacks": "外挂",
    "cheat": "外挂",
    "cheating": "外挂",
    "cheater": "外挂",
    "lag": "卡顿",
    "stutter": "卡顿",
    "latency": "卡顿",
    "crash": "闪退",
    "crashes": "闪退",
}


def _dedupe_preserve_order(items):
    seen = set()
    result = []
    for item in items:
        if not item:
            continue
        value = item.strip()
        if not value:
            continue
        key = value.lower()
        if key in seen:
            continue
        seen.add(key)
        result.append(value)
    return result


def _put_translate_cache(key, value):
    if len(FB_TRANSLATE_CACHE) >= FB_TRANSLATE_CACHE_MAX_SIZE:
        # 轻量淘汰：删除最早插入的一批键，避免缓存无界增长
        stale_keys = list(FB_TRANSLATE_CACHE.keys())[:300]
        for stale_key in stale_keys:
            FB_TRANSLATE_CACHE.pop(stale_key, None)
    FB_TRANSLATE_CACHE[key] = value


def _escape_like_term(term):
    return term.replace("\\", "\\\\").replace("%", r"\%").replace("_", r"\_")


def _split_query_terms(query):
    if not query:
        return []
    raw_parts = FB_QUERY_SEPARATOR_PATTERN.split(query.strip())
    terms = _dedupe_preserve_order(raw_parts)
    return terms[:FB_QUERY_MAX_TERMS]


def _expand_synonyms(terms):
    if not terms:
        return []
    expanded = list(terms)
    lowered_terms = {t.lower() for t in terms}
    for group in FB_SYNONYM_GROUPS:
        group_terms = group.get("terms", [])
        lowered_group = {g.lower() for g in group_terms}
        if lowered_terms & lowered_group:
            expanded.extend(group_terms)
            expanded.append(group.get("canonical", ""))
    return _dedupe_preserve_order(expanded)


def _translate_terms_to_multilang(terms, target_languages=None):
    if not terms:
        return []
    if target_languages is None:
        target_languages = ["en", "id", "th", "vi"]

    cache_key = ("multi", "|".join(sorted([t.lower() for t in terms])))
    cached = FB_TRANSLATE_CACHE.get(cache_key)
    if cached:
        return cached

    if not qwen_client:
        return []

    try:
        term_list = terms[:FB_MAX_TRANSLATE_TERMS]
        prompt = (
            "你是多语言查询扩展助手。请把输入词列表分别翻译成目标语言并返回 JSON。\n"
            "要求：\n"
            "1) 只返回 JSON，不要解释；\n"
            "2) JSON 结构：{\"items\":[{\"source\":\"原词\",\"translations\":[\"词1\",\"词2\"]}]}\n"
            "3) 仅保留适合社媒评论检索的常见表达；\n"
            "4) 去掉无意义词；\n"
            f"目标语言: {', '.join(target_languages)}\n"
            f"输入词: {json.dumps(term_list, ensure_ascii=False)}"
        )
        response = qwen_client.chat.completions.create(
            model="qwen-turbo",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2
        )
        content = (response.choices[0].message.content or "").strip()
        if "```json" in content:
            content = content.split("```json")[1].split("```")[0].strip()
        elif content.startswith("```"):
            content = content.split("```")[1].split("```")[0].strip()
        data = json.loads(content)
        translations = []
        for item in data.get("items", []):
            translations.extend(item.get("translations", []))
        result = _dedupe_preserve_order(translations)
        _put_translate_cache(cache_key, result)
        return result
    except Exception as e:
        logger.warning(f"⚠️ 多语言扩展失败，降级为原词检索: {e}")
        return []


def _expand_query_terms(query):
    base_terms = _split_query_terms(query)
    if not base_terms:
        return []
    synonym_terms = _expand_synonyms(base_terms)
    translated_terms = _translate_terms_to_multilang(synonym_terms)
    all_terms = _dedupe_preserve_order(synonym_terms + translated_terms)
    return all_terms[: (FB_QUERY_MAX_TERMS * 3)]


def _build_fb_search_filters(start_date, end_date, sentiment, post_urls, expanded_terms=None):
    where_clauses = []
    params = []

    if start_date:
        where_clauses.append("created_at >= %s")
        params.append(start_date)
    if end_date:
        where_clauses.append("created_at <= %s")
        params.append(end_date + ' 23:59:59')

    if expanded_terms:
        term_clauses = []
        for term in expanded_terms[:FB_QUERY_SQL_MAX_TERMS]:
            term_clauses.append("(content ILIKE %s ESCAPE '\\' OR brief_analysis ILIKE %s ESCAPE '\\')")
            escaped = _escape_like_term(term)
            like_value = f"%{escaped}%"
            params.extend([like_value, like_value])
        where_clauses.append("(" + " OR ".join(term_clauses) + ")")

    if sentiment:
        if sentiment == 'positive':
            where_clauses.append("sentiment_score > 0.3")
        elif sentiment == 'negative':
            where_clauses.append("sentiment_score < -0.2")
        elif sentiment == 'neutral':
            where_clauses.append("sentiment_score BETWEEN -0.2 AND 0.3")

    if post_urls:
        placeholders = ','.join(['%s'] * len(post_urls))
        where_clauses.append(f"post_url IN ({placeholders})")
        params.extend(post_urls)

    where_sql = " AND ".join(where_clauses) if where_clauses else "1=1"
    return where_sql, params


def _extract_tokens_for_wordcloud(text):
    if not text:
        return []
    chinese_tokens = FB_CHINESE_TOKEN_PATTERN.findall(text)
    latin_tokens = FB_LATIN_TOKEN_PATTERN.findall(text.lower())
    return chinese_tokens + latin_tokens


def _resolve_sentiment_filter(sentiment, tone, default=''):
    valid_sentiments = {'positive', 'neutral', 'negative'}
    tone_normalized = (tone or '').strip().lower()
    if tone_normalized == 'all':
        return ''
    if tone_normalized in valid_sentiments:
        return tone_normalized
    sentiment_normalized = (sentiment or '').strip().lower()
    if sentiment_normalized == 'all':
        return ''
    if sentiment_normalized in valid_sentiments:
        return sentiment_normalized
    return default


def _public_tone_value(effective_sentiment):
    return effective_sentiment if effective_sentiment else 'all'


def _normalize_word(word):
    if not word:
        return ""
    w = word.strip().lower()
    if not w:
        return ""
    mapped = FB_WORD_NORMALIZATION_MAP.get(w)
    if mapped:
        return mapped
    if FB_CHINESE_TOKEN_PATTERN.fullmatch(word):
        return word
    return w


def _is_noise_token(token):
    if not token:
        return True
    t = token.strip().lower()
    if not t:
        return True
    if t in FB_STOP_WORDS:
        return True
    # 中文/英文笑声噪声
    if re.fullmatch(r"哈{2,}", token):
        return True
    if re.fullmatch(r"(ha){2,}", t):
        return True
    if t in {"lol", "lmao", "rofl"}:
        return True
    return False


def _count_keyword_stats_from_rows(rows):
    """
    返回:
      - tf_counter: 词项总出现次数
      - df_counter: 出现在多少条评论中（更适合舆情风向）
    """
    tf_counter = Counter()
    df_counter = Counter()

    for row in rows:
        original_comment = row.get('content', '')
        tokens_in_comment = set()
        for token in _extract_tokens_for_wordcloud(original_comment):
            normalized = _normalize_word(token)
            if len(normalized) < FB_WORD_MIN_CHARS:
                continue
            if normalized.isdigit():
                continue
            if FB_CHINESE_TOKEN_PATTERN.fullmatch(normalized) and len(normalized) > FB_CHINESE_MAX_CHARS:
                continue
            if not FB_CHINESE_TOKEN_PATTERN.fullmatch(normalized) and len(normalized) < FB_LATIN_MIN_CHARS:
                continue
            if _is_noise_token(normalized):
                continue
            tf_counter[normalized] += 1
            tokens_in_comment.add(normalized)

        for token in tokens_in_comment:
            df_counter[token] += 1

    return tf_counter, df_counter


def _translate_words_to_chinese(words):
    if not words:
        return {}
    uncached = []
    result = {}
    for word in words:
        cache_key = ("zh", word.lower())
        if cache_key in FB_TRANSLATE_CACHE:
            result[word] = FB_TRANSLATE_CACHE[cache_key]
        else:
            uncached.append(word)

    if not uncached or not qwen_client:
        return result

    try:
        batch = uncached[:FB_MAX_TRANSLATE_TERMS]
        prompt = (
            "请将以下词汇翻译为简体中文检索词，返回 JSON，禁止解释。\n"
            "格式: {\"items\":[{\"source\":\"word\",\"zh\":\"中文词\"}]}\n"
            f"词汇: {json.dumps(batch, ensure_ascii=False)}"
        )
        response = qwen_client.chat.completions.create(
            model="qwen-turbo",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2
        )
        content = (response.choices[0].message.content or "").strip()
        if "```json" in content:
            content = content.split("```json")[1].split("```")[0].strip()
        elif content.startswith("```"):
            content = content.split("```")[1].split("```")[0].strip()
        data = json.loads(content)
        for item in data.get("items", []):
            source = (item.get("source") or "").strip()
            zh = (item.get("zh") or "").strip()
            if source and zh:
                result[source] = zh
                _put_translate_cache(("zh", source.lower()), zh)
    except Exception as e:
        logger.warning(f"⚠️ 词云翻译降级: {e}")
    return result


@app.route('/fb_search', methods=['GET'])
@login_required
def fb_search():
    """
    FB 评论语义检索 + 日期过滤 + 关键词过滤

    Query params:
        - query: 搜索关键词（可选）
        - start_date: 开始日期 YYYY-MM-DD（可选）
        - end_date: 结束日期 YYYY-MM-DD（可选）
        - limit: 返回数量（默认 200）
    """
    try:
        query = request.args.get('query', '').strip()
        start_date = request.args.get('start_date')
        end_date = request.args.get('end_date')
        sentiment = request.args.get('sentiment', '').strip()
        post_urls = request.args.getlist('post_url')  # 支持多选
        post_urls = [url.strip() for url in post_urls if url.strip()]
        limit = int(request.args.get('limit', 200))
        limit = min(max(limit, 1), 500)

        expanded_terms = _expand_query_terms(query)
        where_sql, params = _build_fb_search_filters(
            start_date=start_date,
            end_date=end_date,
            sentiment=sentiment,
            post_urls=post_urls,
            expanded_terms=expanded_terms
        )

        # 查询评论
        sql = f"""
            SELECT id, post_url, comment_id, author, created_at, content,
                   sentiment_score, category, language, post_link, embedding, brief_analysis
            FROM fb_comments
            WHERE {where_sql}
            ORDER BY created_at DESC
            LIMIT %s
        """
        params.append(limit)

        rows = db.query_all(sql, tuple(params))

        # 如果有搜索词且有 embedding，进行相似度排序
        if query and rows:
            query_embedding = rag.get_embedding(query)
            if query_embedding:
                for row in rows:
                    if row.get('embedding'):
                        try:
                            comment_embedding = json.loads(row['embedding'])
                            similarity = rag._cosine_similarity(query_embedding, comment_embedding)
                            row['similarity'] = similarity
                        except (ValueError, TypeError, json.JSONDecodeError):
                            row['similarity'] = 0.0
                    else:
                        row['similarity'] = 0.0
                # 按相似度排序
                rows.sort(key=lambda x: x.get('similarity', 0), reverse=True)

        # 格式化返回
        results = []
        for r in rows:
            row_dict = dict(r)
            if 'created_at' in row_dict and row_dict['created_at']:
                row_dict['created_at'] = row_dict['created_at'].isoformat()
            if 'scraped_at' in row_dict and row_dict['scraped_at']:
                row_dict['scraped_at'] = row_dict['scraped_at'].isoformat()
            # 移除 embedding 字段（太大）
            row_dict.pop('embedding', None)
            results.append(row_dict)

        return jsonify({
            'status': 'success',
            'count': len(results),
            'expanded_terms': expanded_terms[:20],
            'results': results
        })

    except Exception as e:
        logger.error(f"❌ FB 搜索失败: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/fb_wordcloud', methods=['GET'])
@login_required
def fb_wordcloud():
    """获取 FB 评论词云（统一中文展示）"""
    try:
        query = request.args.get('query', '').strip()
        start_date = request.args.get('start_date')
        end_date = request.args.get('end_date')
        sentiment = request.args.get('sentiment', '').strip()
        tone = request.args.get('tone', '').strip()
        post_urls = request.args.getlist('post_url')
        post_urls = [url.strip() for url in post_urls if url.strip()]
        limit = int(request.args.get('limit', FB_WORDCLOUD_MAX_ROWS))
        limit = min(max(limit, 100), FB_WORDCLOUD_MAX_ROWS)
        effective_sentiment = _resolve_sentiment_filter(sentiment, tone, default='negative')

        expanded_terms = _expand_query_terms(query)
        where_sql, params = _build_fb_search_filters(
            start_date=start_date,
            end_date=end_date,
            sentiment=effective_sentiment,
            post_urls=post_urls,
            expanded_terms=expanded_terms
        )

        sql = f"""
            SELECT content
            FROM fb_comments
            WHERE {where_sql}
            ORDER BY created_at DESC
            LIMIT %s
        """
        params.append(limit)
        rows = db.query_all(sql, tuple(params))

        tf_counter, df_counter = _count_keyword_stats_from_rows(rows)

        if not tf_counter:
            return jsonify({
                'status': 'success',
                'items': [],
                'expanded_terms': expanded_terms[:20],
                'tone': _public_tone_value(effective_sentiment)
            })

        foreign_top_words = [
            w for w, _ in tf_counter.most_common(FB_MAX_TRANSLATE_TERMS)
            if not FB_CHINESE_TOKEN_PATTERN.fullmatch(w)
        ]
        translated_map = _translate_words_to_chinese(foreign_top_words)

        zh_tf_counter = Counter()
        zh_df_counter = Counter()
        for word, count in tf_counter.items():
            normalized = _normalize_word(word)
            if FB_CHINESE_TOKEN_PATTERN.fullmatch(normalized):
                zh_word = normalized
            else:
                zh_word = translated_map.get(word, normalized)
                zh_word = _normalize_word(zh_word)
            if len(zh_word) < FB_WORD_MIN_CHARS:
                continue
            if _is_noise_token(zh_word):
                continue
            # 词云展示只保留中文词，避免出现大量外语碎词影响可读性
            if not FB_CHINESE_TOKEN_PATTERN.fullmatch(zh_word):
                continue
            zh_tf_counter[zh_word] += count

        for word, count in df_counter.items():
            normalized = _normalize_word(word)
            if FB_CHINESE_TOKEN_PATTERN.fullmatch(normalized):
                zh_word = normalized
            else:
                zh_word = translated_map.get(word, normalized)
                zh_word = _normalize_word(zh_word)
            if len(zh_word) < FB_WORD_MIN_CHARS:
                continue
            if _is_noise_token(zh_word):
                continue
            if not FB_CHINESE_TOKEN_PATTERN.fullmatch(zh_word):
                continue
            zh_df_counter[zh_word] += count

        min_count = 1 if len(rows) < 50 else FB_WORD_MIN_COUNT_DEFAULT
        ranked = [
            (word, df, zh_tf_counter.get(word, 0))
            for word, df in sorted(zh_df_counter.items(), key=lambda x: (x[1], zh_tf_counter.get(x[0], 0)), reverse=True)
            if df >= min_count
        ][:40]

        max_df = max((item[1] for item in ranked), default=1)
        max_tf = max((item[2] for item in ranked), default=1)
        items = []
        total = len(ranked) if ranked else 1
        for idx, (word, df, tf) in enumerate(ranked):
            # 强制梯度：即使 df 接近，也保证头部词明显更大
            rank_factor = max(0.0, 1.0 - (idx / total))
            score = int(round(20 + rank_factor * 85 + (df / max_df) * 25 + (tf / max_tf) * 10))
            score = max(score, 18)
            items.append({'name': word, 'value': score})

        return jsonify({
            'status': 'success',
            'items': items,
            'expanded_terms': expanded_terms[:20],
            'tone': _public_tone_value(effective_sentiment),
            'min_count': min_count,
            'score_mode': 'document_frequency'
        })
    except Exception as e:
        logger.error(f"❌ FB 词云失败: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/fb_topic_top10', methods=['GET'])
@login_required
def fb_topic_top10():
    """获取关键词 Top10，默认负面视角（基于评论原文）"""
    try:
        query = request.args.get('query', '').strip()
        start_date = request.args.get('start_date')
        end_date = request.args.get('end_date')
        sentiment = request.args.get('sentiment', '').strip()
        tone = request.args.get('tone', '').strip()
        post_urls = request.args.getlist('post_url')
        post_urls = [url.strip() for url in post_urls if url.strip()]
        effective_sentiment = _resolve_sentiment_filter(sentiment, tone, default='negative')

        expanded_terms = _expand_query_terms(query)
        where_sql, params = _build_fb_search_filters(
            start_date=start_date,
            end_date=end_date,
            sentiment=effective_sentiment,
            post_urls=post_urls,
            expanded_terms=expanded_terms
        )

        sql = f"""
            SELECT content
            FROM fb_comments
            WHERE {where_sql}
            ORDER BY created_at DESC
            LIMIT %s
        """
        params.append(1000)
        rows = db.query_all(sql, tuple(params))

        tf_counter, df_counter = _count_keyword_stats_from_rows(rows)

        foreign_top_words = [
            w for w, _ in tf_counter.most_common(FB_MAX_TRANSLATE_TERMS)
            if not FB_CHINESE_TOKEN_PATTERN.fullmatch(w)
        ]
        translated_map = _translate_words_to_chinese(foreign_top_words)

        zh_tf_counter = Counter()
        zh_df_counter = Counter()

        for word, count in tf_counter.items():
            normalized = _normalize_word(word)
            if FB_CHINESE_TOKEN_PATTERN.fullmatch(normalized):
                zh_word = normalized
            else:
                zh_word = translated_map.get(word, normalized)
                zh_word = _normalize_word(zh_word)
            if len(zh_word) < FB_WORD_MIN_CHARS:
                continue
            if _is_noise_token(zh_word):
                continue
            if not FB_CHINESE_TOKEN_PATTERN.fullmatch(zh_word):
                continue
            zh_tf_counter[zh_word] += count

        for word, count in df_counter.items():
            normalized = _normalize_word(word)
            if FB_CHINESE_TOKEN_PATTERN.fullmatch(normalized):
                zh_word = normalized
            else:
                zh_word = translated_map.get(word, normalized)
                zh_word = _normalize_word(zh_word)
            if len(zh_word) < FB_WORD_MIN_CHARS:
                continue
            if _is_noise_token(zh_word):
                continue
            if not FB_CHINESE_TOKEN_PATTERN.fullmatch(zh_word):
                continue
            zh_df_counter[zh_word] += count

        items = [
            {'name': word, 'value': df}
            for word, df in sorted(zh_df_counter.items(), key=lambda x: (x[1], zh_tf_counter.get(x[0], 0)), reverse=True)[:10]
            if df >= 2
        ]

        return jsonify({
            'status': 'success',
            'items': items,
            'tone': _public_tone_value(effective_sentiment)
        })
    except Exception as e:
        logger.error(f"❌ FB 主题Top10失败: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/fb_export', methods=['POST'])
@login_required
def fb_export():
    """导出 FB 评论为 Excel"""
    try:
        data = request.json
        comment_ids = data.get('comment_ids', [])

        if not comment_ids:
            return jsonify({'error': '未选择评论'}), 400

        # 查询评论
        placeholders = ','.join(['%s'] * len(comment_ids))
        sql = f"SELECT * FROM fb_comments WHERE id IN ({placeholders}) ORDER BY created_at DESC"
        rows = db.query_all(sql, tuple(comment_ids))

        if not rows:
            return jsonify({'error': '未找到评论'}), 404

        # 创建 Excel
        wb = Workbook()
        ws = wb.active
        ws.title = "FB评论导出"

        # 表头
        headers = ['ID', '作者', '评论时间', '内容', '简要分析', '情感分数', '分类', '语言', '原帖链接']
        ws.append(headers)

        # 样式
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        for cell in ws[1]:
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal='center', vertical='center')

        # 数据行
        for row in rows:
            ws.append([
                row['id'],
                row.get('author', ''),
                row.get('created_at').strftime('%Y-%m-%d %H:%M:%S') if row.get('created_at') else '',
                row.get('content', ''),
                row.get('brief_analysis', ''),
                row.get('sentiment_score', 0),
                row.get('category', ''),
                row.get('language', ''),
                row.get('post_link', '')
            ])

        # 调整列宽
        ws.column_dimensions['A'].width = 8
        ws.column_dimensions['B'].width = 15
        ws.column_dimensions['C'].width = 20
        ws.column_dimensions['D'].width = 50
        ws.column_dimensions['E'].width = 50
        ws.column_dimensions['F'].width = 12
        ws.column_dimensions['G'].width = 15
        ws.column_dimensions['H'].width = 10
        ws.column_dimensions['I'].width = 40

        # 保存到内存
        output = BytesIO()
        wb.save(output)
        output.seek(0)

        filename = f"FB评论导出_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"

        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=filename
        )

    except Exception as e:
        logger.error(f"❌ FB 导出失败: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/fb_stats', methods=['GET'])
@login_required
def fb_stats():
    """获取 FB 评论统计数据（用于看板图表）"""
    try:
        # 今日新增
        today_count = db.query_one("""
            SELECT COUNT(*) as count FROM fb_comments
            WHERE DATE(scraped_at) = CURRENT_DATE
        """)

        # 近 7 天趋势
        trend_data = db.query_all("""
            SELECT DATE(scraped_at) as date, COUNT(*) as count
            FROM fb_comments
            WHERE scraped_at >= CURRENT_DATE - INTERVAL '7 days'
            GROUP BY DATE(scraped_at)
            ORDER BY date
        """)

        # 情感分布
        sentiment_dist = db.query_all("""
            SELECT
                CASE
                    WHEN sentiment_score > 0.3 THEN 'positive'
                    WHEN sentiment_score < -0.3 THEN 'negative'
                    ELSE 'neutral'
                END as sentiment,
                COUNT(*) as count
            FROM fb_comments
            GROUP BY sentiment
        """)

        # 高频关键词（简化版，从分类统计）
        category_dist = db.query_all("""
            SELECT category, COUNT(*) as count
            FROM fb_comments
            WHERE category IS NOT NULL AND category != 'unknown'
            GROUP BY category
            ORDER BY count DESC
            LIMIT 10
        """)

        return jsonify({
            'status': 'success',
            'today_count': today_count['count'] if today_count else 0,
            'trend': [{'date': str(r['date']), 'count': r['count']} for r in trend_data],
            'sentiment': [{'sentiment': r['sentiment'], 'count': r['count']} for r in sentiment_dist],
            'categories': [{'category': r['category'], 'count': r['count']} for r in category_dist]
        })

    except Exception as e:
        logger.error(f"❌ FB 统计失败: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/fb_config', methods=['GET'])
@login_required
def fb_config():
    """获取 FB 监控配置列表"""
    try:
        configs = db.query_all("""
            SELECT id, post_url, post_title, is_active, last_scraped_at, created_at
            FROM fb_monitor_config
            ORDER BY created_at DESC
        """)

        result = []
        for config in configs:
            result.append({
                'id': config['id'],
                'post_url': config['post_url'],
                'post_title': config.get('post_title', ''),
                'is_active': config['is_active'],
                'last_scraped_at': config['last_scraped_at'].isoformat() if config.get('last_scraped_at') else None,
                'created_at': config['created_at'].isoformat() if config.get('created_at') else None
            })

        return jsonify({
            'status': 'success',
            'configs': result
        })

    except Exception as e:
        logger.error(f"❌ 获取配置失败: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/fb_config', methods=['POST'])
@login_required
def add_fb_config():
    """添加 FB 监控配置"""
    try:
        data = request.json
        post_url = data.get('post_url', '').strip()
        post_title = data.get('post_title', '').strip()

        if not post_url:
            return jsonify({'error': '帖子链接不能为空'}), 400

        # 检查是否已存在
        existing = db.query_one("SELECT id FROM fb_monitor_config WHERE post_url = %s", (post_url,))
        if existing:
            return jsonify({'error': '该帖子已在监控列表中'}), 400

        # 插入配置
        db.execute("""
            INSERT INTO fb_monitor_config (post_url, post_title, created_by)
            VALUES (%s, %s, %s)
        """, (post_url, post_title, session.get('user_id')))

        logger.info(f"✅ 添加 FB 监控配置: {post_url}")

        return jsonify({'status': 'success', 'message': '添加成功'})

    except Exception as e:
        logger.error(f"❌ 添加配置失败: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/fb_config/<int:config_id>', methods=['DELETE'])
@login_required
def delete_fb_config(config_id):
    """删除 FB 监控配置（同时删除该帖子的所有评论数据）"""
    try:
        # 先查询配置的 post_url
        config = db.query_one("SELECT post_url FROM fb_monitor_config WHERE id = %s", (config_id,))
        if not config:
            return jsonify({'error': '配置不存在'}), 404

        post_url = config['post_url']

        # 删除该帖子的所有评论
        result = db.execute("DELETE FROM fb_comments WHERE post_url = %s", (post_url,))
        deleted_comments = result if isinstance(result, int) else 0
        logger.info(f"🗑️ 删除帖子评论: post_url={post_url}, 删除 {deleted_comments} 条评论")

        # 删除监控配置
        db.execute("DELETE FROM fb_monitor_config WHERE id = %s", (config_id,))
        logger.info(f"✅ 删除 FB 监控配置: ID={config_id}")

        return jsonify({
            'status': 'success',
            'message': f'删除成功（配置 + {deleted_comments} 条评论）'
        })

    except Exception as e:
        logger.error(f"❌ 删除配置失败: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/fb_cleanup_orphan_comments', methods=['POST'])
@login_required
def cleanup_orphan_comments():
    """清理孤立评论数据（监控配置已删除但评论还在的数据）"""
    try:
        # 查询所有有效的监控配置 post_url
        valid_urls = db.query_all("SELECT post_url FROM fb_monitor_config")
        valid_url_set = {row['post_url'] for row in valid_urls}

        logger.info(f"📋 当前有效监控配置: {len(valid_url_set)} 个")

        if not valid_url_set:
            # 如果没有任何监控配置，删除所有评论
            logger.warning("⚠️  没有任何监控配置，将删除所有评论数据")
            result = db.execute("DELETE FROM fb_comments")
            deleted_count = result if isinstance(result, int) else 0
            logger.info(f"✅ 删除了 {deleted_count} 条评论")
            return jsonify({
                'status': 'success',
                'message': f'已删除所有评论数据（{deleted_count} 条）'
            })

        # 查询所有评论的 post_url
        all_comments = db.query_all("SELECT DISTINCT post_url FROM fb_comments")
        comment_urls = {row['post_url'] for row in all_comments}

        logger.info(f"📊 数据库中有 {len(comment_urls)} 个不同的帖子评论")

        # 找出孤立的 URL（评论存在但配置不存在）
        orphan_urls = comment_urls - valid_url_set

        if not orphan_urls:
            logger.info("✅ 没有孤立评论，数据库干净")
            return jsonify({
                'status': 'success',
                'message': '没有孤立评论，数据库已是干净状态'
            })

        logger.info(f"🗑️  发现 {len(orphan_urls)} 个孤立帖子的评论，准备删除...")

        total_deleted = 0
        for url in orphan_urls:
            result = db.execute("DELETE FROM fb_comments WHERE post_url = %s", (url,))
            deleted = result if isinstance(result, int) else 0
            total_deleted += deleted
            logger.info(f"   删除: {url[:80]}... ({deleted} 条评论)")

        logger.info(f"✅ 清理完成，共删除 {total_deleted} 条孤立评论")

        return jsonify({
            'status': 'success',
            'message': f'清理完成，删除了 {len(orphan_urls)} 个帖子的 {total_deleted} 条孤立评论'
        })

    except Exception as e:
        logger.error(f"❌ 清理失败: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/import_manual_comments', methods=['POST'])
@login_required
def import_manual_comments():
    """一次性导入手动复制的评论（用完即删）"""
    try:
        MANUAL_COMMENTS = {
            "https://www.instagram.com/p/DWli1ACClRg/": [
                "Glowing min glowing",
                "udah beliii min, lucu bgttt zetiannya wkwk",
                "Trinity maksudnya apa? 🙏🏻",
                "Abis 100 dm ga dapet",
                "Aduh min, item Trinity nya sih oke , tapi tambahin glowing wand lah , bnyak player yang mikir mage ga perlu glowing, karena langsung dari moonton nya😢",
                "Sesat, gk pakek glowing kwkwk",
                "Egk pake glowing gpp kah?",
                "build nya @rrq_clayyy",
                "HAMA YG PAKE NGGA BELI GLOWING",
                "gmn bang lgsg dr muntunnya ini wkwk",
                "Minn, itu pharsa bukannya justru bisa d counter Zetian yakk??",
                "by one ga min kalah post ig wkwkwk",
                "Benerin darksytem lo min buat solo player parag banget masa prosentase nya 200 pertandingan cuma 40% WR nya",
                "Min buat starlight lebih effort dan bervariasi napa kayak starlight yang dulu2. Ini Zetian menurut gue parah asli gk tau temannya apa, plek ketiplek kayak skin basic beda warna dan ditambah aksesoris tipis2 buat pantes2an doang.",
                "MAKIN BORING SKIN STARLIGHT!",
                "build versiku, talent insipire, sepatu sungai, impure rage. magic boots, enchanted, glowing, fleeting time, concentrad, terakhir kadang blood wings, kadang divine glave",
                "Bootnya aja aku ngga pake itu apalagi yg lainnya, aku yg penting Holy cristal sama Glowing aja sama fleeting time biar cpt gunain Ulti",
                "inspire, sepatu sungai, impure rage. magic/tough/arcane boots, liat musuhnya, lalu glowing, trinity item, tengah\"nya taro fleeting time. cadangan imortal sm winter. done.",
                "Kalo gua ttep pake glowing sih",
                "Atmin darksistem gamau bikin glowing",
                "zetian tuh kan kalo erly , susah di lawan sama herly soal nya sheeld nya tebel, kalo late game enak",
                "Dah beli Gweh skin sl Zetian, bagus min tinggal matchnya aja jngn kek tai",
                "Enak ny Cetian Ulti ny kyk maphack 😂 .. Counter utama ny Teman Buta Map 😤",
            ],
            "https://www.instagram.com/p/DWV4cqlILcH/": [
                "Aku nak 🔥🙌",
                "😍😍 I'm ready for both Minato & Itachi 🔥",
                "Koleb jjk kapan abangkuh",
                "Delete Sky Piercer item >>>>",
                "What's the Valir skin doin on the Naruto Banner??",
                "Day 85 asking for Xavier Gojo skin",
                "What about a guide on how you guys plan to fix your flawed reporting system next? :D Fix your reporting system. I'm going to write what somebody else commented previously which I think best summarises it: You can get muted for 14 matches for simply saying \"noob\" but someone can intentionally throw the game and be 0-13-2 and have no clear violation while trolling. But when you express your anger you get reported for inappropriate chatting while the person intentionally throwing the game gets away completely scot free. @mobilelegendsgame, you guys don't have any idea what actually goes on during gameplays, do you? Talk about \"ensuring fair gameplay\" that you guys just love to shove at our faces. Your OWN reporting system isn't even supporting fair gameplay.",
                "Stùpid, usèless developers! Solo Ranked matchmaking is atrocious! I wish your families would diè! Narrow-eyed animals.",
                "_!_ for admins !!!",
                "I got itachi sasuke and naruto",
                "Only 3 free beeds exchange tokens this time?",
                "Stop deleting comments and listen to your community! Where is JJK?",
                "Ml tau lose streak terus",
                "The closest we can get of Valir and Vale coupley skin HAHAHAHAHA",
                "im getting that sasuke back 😍",
                "Когда выйдет",
                "berapa cras oi",
                "So do yall just look at the seasons ban list and decide to stick it up your ass?",
            ],
            "https://www.instagram.com/p/DWs8m3eiBfE/": [
                "tgl 10 udah pasti punya itachi😍",
                "valir itu dapet dri mna?",
                "41 draw cm 400an crest🗿",
                "Udh kebeli alhmdulillah",
                "Lagi mode hemat buat minato 🫠🫠",
                "one piece kapan keluar, dan plis gosah yg aneh dan maksa kek Julian Itachi😂 beda di anime beda di skill hero",
                "Modal Pase 1 langsung dapet Itachi di 20× draw😍 thank bgt❤️❤️❤️❤️",
                "Modal kouta udah cukup",
                "Klo ambil 2 skin, Itachi sama apa ya enaknya... Sasuke udah punya",
                "Kenapa gratis? Karna efeknya tipis 😂",
                "Udh nabung 450 lumayan dikit-dikit tembus 1.200 demi julian 😆",
                "Udh gacha 49 kali gk dapat dapat itu skin,, kayak nya emang sama monton,, udh diseting gk bakalan muncul di gacha,, harus ditukar pake chris baru bisa dapat itu skin😢",
                "Kapan colleb ama one Piece min, keburu balek pondok gw min😢😢",
                "jangan ngejek min,makanya gratisin semua skin",
                "@realmobilelegendsid tolong ya skin Julian di perbaiki efeknya radak kurang gak sesuai animenya",
                "Ya ampun sinyal masih lag gak di perbaikin bnerin servernya min",
                "Skill 2 itachi kureng",
                "Ngasih skin gratis nya yg agak bagusan kek,biar yg gak mampu beli ngerasa puas juga",
            ],
            "https://www.instagram.com/p/DWn7KGJEweD/": [
                "Aku Vale, kamu Estes, gas bang?",
                "Maaf ini siapa ? Muncul di timeline saya 🙏salam interaksi",
                "Super cecep ready terus 🙌",
                "guru alis tebal❌ guru alis empat ✅",
                "suganteh @ae.ninoo",
                "mandi ga sih mas",
                "jangan terlalu di anime animekan mas",
                "Jan dianime\"in banget mas 🙂\u200d↔️",
                "ibarat rambo",
                "el mrangkak",
                "cctv lantai3 rumah inara ada distory😱",
                "Misi menyelamatkan murid yg star syndrome , nyelamatin diri sendiri dari bola dan angin aja ga bisa",
                "stress",
                "awasnya bijinya kena😂",
            ],
        }

        import hashlib
        now = datetime.datetime.now(tasks.BEIJING_TZ)
        total_inserted = 0
        skipped = 0

        for post_url, comments in MANUAL_COMMENTS.items():
            for i, content in enumerate(comments):
                cid = hashlib.md5(f"{post_url}:{content}:{i}".encode()).hexdigest()
                existing = db.query_one(
                    "SELECT 1 FROM fb_comments WHERE comment_id = %s", (cid,)
                )
                if existing:
                    skipped += 1
                    continue
                db.execute(
                    """INSERT INTO fb_comments
                       (post_url, comment_id, author, created_at, content,
                        sentiment_score, category, language, post_link)
                       VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)""",
                    (post_url, cid, 'ig_user', now, content,
                     None, None, None, post_url)
                )
                total_inserted += 1

        logger.info(f"✅ 手动评论导入完成: 插入 {total_inserted}, 跳过 {skipped}")
        return jsonify({
            'status': 'success',
            'message': f'导入完成: 新增 {total_inserted} 条, 跳过 {skipped} 条重复',
            'total_inserted': total_inserted,
            'skipped': skipped
        })
    except Exception as e:
        logger.error(f"❌ 手动评论导入失败: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/fb_get_recent_datasets', methods=['GET'])
@login_required
def get_recent_datasets():
    """获取最近的 Apify runs 和 datasets"""
    try:
        if not APIFY_TOKEN:
            return jsonify({'error': 'Apify token 未配置'}), 500

        headers = {
            "Authorization": f"Bearer {APIFY_TOKEN}",
            "Content-Type": "application/json"
        }

        # 获取最近的 runs
        api_url = "https://api.apify.com/v2/acts/apify~facebook-comments-scraper/runs?limit=10&status=SUCCEEDED"
        response = requests.get(api_url, headers=headers, timeout=30)

        if response.status_code != 200:
            return jsonify({'error': f'Apify API 错误: {response.status_code}'}), 500

        data = response.json()
        runs = data.get('data', {}).get('items', [])

        results = []
        for run in runs:
            dataset_id = run.get('defaultDatasetId')
            started_at = run.get('startedAt')
            finished_at = run.get('finishedAt')

            # 尝试从 input 获取 post URL
            run_input = run.get('buildInput', {}) or run.get('input', {})
            start_urls = run_input.get('startUrls', [])
            post_url = start_urls[0].get('url') if start_urls else 'Unknown'

            results.append({
                'run_id': run.get('id'),
                'dataset_id': dataset_id,
                'post_url': post_url,
                'started_at': started_at,
                'finished_at': finished_at
            })

        return jsonify({
            'status': 'success',
            'runs': results
        })

    except Exception as e:
        logger.error(f"❌ 获取 runs 失败: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/fb_process_existing_data', methods=['POST'])
@login_required
def process_existing_data():
    """处理已抓取但未分析的 Apify 数据"""
    try:
        data = request.json
        run_ids = data.get('run_ids', [])  # 改为接收 run_ids
        datasets = data.get('datasets', [])

        if not run_ids and not datasets:
            return jsonify({'error': '未提供 run_id 或 dataset 信息'}), 400

        # 如果提供的是 run_ids，先转换为 datasets
        if run_ids:
            headers = {
                "Authorization": f"Bearer {APIFY_TOKEN}",
                "Content-Type": "application/json"
            }

            # 已知的帖子 URL 映射（按 run_id 顺序）
            known_post_urls = data.get('post_urls', [])

            datasets = []
            for idx, run_id in enumerate(run_ids):
                try:
                    # 从 run_id 获取 dataset_id 和 input
                    run_api_url = f"https://api.apify.com/v2/actor-runs/{run_id}"
                    response = requests.get(run_api_url, headers=headers, timeout=10)

                    if response.status_code == 200:
                        run_data = response.json()['data']
                        dataset_id = run_data.get('defaultDatasetId')

                        # 打印完整数据结构帮助调试
                        logger.info(f"📋 Run {run_id} keys: {list(run_data.keys())}")
                        logger.info(f"📋 Run {run_id} defaultDatasetId: {dataset_id}")

                        # 尝试多种方式获取 post_url
                        post_url = None

                        # 方式1：从 input 获取
                        for input_key in ['input', 'buildInput', 'options']:
                            run_input = run_data.get(input_key)
                            if run_input and isinstance(run_input, dict):
                                start_urls = run_input.get('startUrls', [])
                                if start_urls:
                                    post_url = start_urls[0].get('url') if isinstance(start_urls[0], dict) else start_urls[0]
                                    break

                        # 方式2：从已知的 URL 列表匹配
                        if not post_url and known_post_urls and idx < len(known_post_urls):
                            post_url = known_post_urls[idx]
                            logger.info(f"📋 Using known post_url: {post_url}")

                        if dataset_id and post_url:
                            datasets.append({
                                'dataset_id': dataset_id,
                                'post_url': post_url
                            })
                            logger.info(f"✅ Run {run_id} -> Dataset {dataset_id}, URL: {post_url[:60]}")
                        else:
                            logger.warning(f"⚠️ Run {run_id} 缺少 dataset_id({dataset_id}) 或 post_url({post_url})")
                    else:
                        logger.error(f"❌ 获取 run {run_id} 失败: {response.status_code}")
                except Exception as e:
                    logger.error(f"❌ 处理 run {run_id} 失败: {e}")
                    continue

        if not datasets:
            return jsonify({'error': '未找到有效的 dataset'}), 400

        logger.info(f"📦 准备处理 {len(datasets)} 个 datasets")

        # 在后台线程处理
        def run_process():
            total_new = 0
            total_updated = 0

            for item in datasets:
                dataset_id = item.get('dataset_id')
                post_url = item.get('post_url')

                if not dataset_id or not post_url:
                    continue

                logger.info(f"📦 处理 dataset: {dataset_id}")

                try:
                    # 获取数据
                    headers = {
                        "Authorization": f"Bearer {APIFY_TOKEN}",
                        "Content-Type": "application/json"
                    }

                    dataset_api_url = f"https://api.apify.com/v2/datasets/{dataset_id}/items"
                    response = requests.get(dataset_api_url, headers=headers, timeout=30)

                    if response.status_code != 200:
                        logger.error(f"❌ Failed to fetch dataset: {response.status_code}")
                        continue

                    items = response.json()
                    logger.info(f"✅ 获取到 {len(items)} 条评论")

                    # 批量查询已存在的 comment_id
                    comment_ids = [i.get('id') for i in items if i.get('id')]
                    existing_ids = set()

                    if comment_ids:
                        placeholders = ','.join(['%s'] * len(comment_ids))
                        sql = f"SELECT comment_id FROM fb_comments WHERE comment_id IN ({placeholders})"
                        rows = db.query_all(sql, tuple(comment_ids))
                        existing_ids = {row['comment_id'] for row in rows}
                        logger.info(f"📊 已存在 {len(existing_ids)} 条评论")

                    # 处理每条评论
                    cutoff_date = datetime.datetime.now(tasks.BEIJING_TZ) - datetime.timedelta(days=7)

                    for idx, comment_item in enumerate(items, 1):
                        comment_id = comment_item.get('id')
                        author = comment_item.get('author', {}).get('name', 'Unknown')
                        content = comment_item.get('text', '')
                        created_at_str = comment_item.get('createdTime')

                        if not content or not comment_id:
                            continue

                        # 解析时间
                        try:
                            created_at = datetime.datetime.fromisoformat(created_at_str.replace('Z', '+00:00'))
                            created_at = created_at.astimezone(tasks.BEIJING_TZ)
                        except:
                            created_at = datetime.datetime.now(tasks.BEIJING_TZ)

                        # 跳过旧评论
                        if created_at < cutoff_date:
                            continue

                        # 跳过已存在的评论
                        if comment_id in existing_ids:
                            total_updated += 1
                            continue

                        if idx % 10 == 0:
                            logger.info(f"🔄 处理进度: {idx}/{len(items)}")

                        # 分析情感
                        sentiment_score, category, language, brief_analysis = tasks.analyze_comment_sentiment(content)

                        # 生成 embedding
                        embedding = rag.get_embedding(content)
                        embedding_json = json.dumps(embedding) if embedding else None

                        # 插入数据库
                        db.execute(
                            """
                            INSERT INTO fb_comments
                            (post_url, comment_id, author, created_at, content, sentiment_score,
                             category, language, post_link, embedding, brief_analysis)
                            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                            """,
                            (post_url, comment_id, author, created_at, content, sentiment_score,
                             category, language, post_url, embedding_json, brief_analysis)
                        )
                        total_new += 1

                    logger.info(f"✅ Dataset {dataset_id} 处理完成")

                except Exception as e:
                    logger.error(f"❌ 处理 dataset {dataset_id} 失败: {e}")
                    continue

            logger.info(f"✅ 全部处理完成: {total_new} 条新增, {total_updated} 条已存在")

        thread = threading.Thread(target=run_process, daemon=True)
        thread.start()

        return jsonify({
            'status': 'success',
            'message': f'开始处理 {len(datasets)} 个 dataset，请稍后刷新页面查看结果'
        })

    except Exception as e:
        logger.error(f"❌ 启动处理失败: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/fb_config/<int:config_id>/toggle', methods=['POST'])
@login_required
def toggle_fb_config(config_id):
    """启用/禁用 FB 监控配置"""
    try:
        config = db.query_one("SELECT is_active FROM fb_monitor_config WHERE id = %s", (config_id,))
        if not config:
            return jsonify({'error': '配置不存在'}), 404

        new_status = not config['is_active']
        db.execute("UPDATE fb_monitor_config SET is_active = %s WHERE id = %s", (new_status, config_id))

        logger.info(f"✅ 切换 FB 监控配置状态: ID={config_id}, active={new_status}")

        return jsonify({'status': 'success', 'is_active': new_status})

    except Exception as e:
        logger.error(f"❌ 切换配置状态失败: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/fb_scheduler_status', methods=['GET'])
@login_required
def fb_scheduler_status():
    """查询 FB 抓取调度器状态"""
    try:
        job = scheduler.get_job('fb_scrape_job')
        if not job:
            return jsonify({'error': '任务未找到'}), 404

        return jsonify({
            'status': 'success',
            'job_id': job.id,
            'next_run_time': job.next_run_time.isoformat() if job.next_run_time else None,
            'trigger': str(job.trigger),
            'timezone': str(scheduler.timezone)
        })
    except Exception as e:
        logger.error(f"❌ 查询调度器状态失败: {e}")
        return jsonify({'error': str(e)}), 500


# ============================================
# 泰国专题报告
# ============================================

# Aliases for backward compat within this file
_thai_datasets_config = thai_datasets_config
_thai_matching_datasets = thai_matching_datasets
_contains_any_tag_or_term = contains_any_tag_or_term


@app.route('/thai-report-tool')
@login_required
def thai_report_tool():
    return render_template('thai_report.html')


@app.route('/api/thai_reset', methods=['POST'])
@login_required
def thai_reset():
    """Full reset: clear dataset tags, recalculate engagement, optionally purge stale sentiment."""
    try:
        payload = request.get_json(silent=True) or {}
        confirm = (payload.get('confirm') or '').strip()
        if confirm != 'RESET_THAI':
            return jsonify({'status': 'error', 'message': '需传 confirm=RESET_THAI 才能执行重置'}), 400

        # 1) Recalculate engagement for posts that belong to Thai datasets
        recalc_row = db.execute_and_fetch_one("""
            WITH recalced AS (
                UPDATE fb_post_metrics m
                SET engagement = COALESCE(m.likes,0) + COALESCE(m.shares,0) + COALESCE(m.comments_count,0),
                    updated_at = NOW()
                WHERE EXISTS (
                    SELECT 1 FROM thai_report_datasets d WHERE d.post_url = m.post_url
                )
                RETURNING 1
            )
            SELECT COUNT(*) AS cnt FROM recalced
        """)
        recalc_cnt = int((recalc_row or {}).get('cnt') or 0)

        # 2) Clear all dataset tags
        del_row = db.execute_and_fetch_one("""
            WITH deleted AS (
                DELETE FROM thai_report_datasets RETURNING 1
            )
            SELECT COUNT(*) AS cnt FROM deleted
        """)
        del_cnt = int((del_row or {}).get('cnt') or 0)

        # 3) Purge stale sentiment records (placeholder analyses)
        purge_row = db.execute_and_fetch_one("""
            WITH purged AS (
                DELETE FROM fb_comments
                WHERE sentiment_score = 0 AND category = 'unknown'
                RETURNING 1
            )
            SELECT COUNT(*) AS cnt FROM purged
        """)
        purge_cnt = int((purge_row or {}).get('cnt') or 0)

        return jsonify({
            'status': 'success',
            'recalculated_engagement': recalc_cnt,
            'deleted_dataset_tags': del_cnt,
            'purged_stale_sentiment': purge_cnt,
            'message': f'重置完成：重算 {recalc_cnt} 条 engagement，清空 {del_cnt} 条标签，清理 {purge_cnt} 条无效情感记录',
        })
    except Exception as e:
        logger.error(f"❌ thai_reset 失败: {e}")
        return jsonify({'status': 'error', 'message': str(e)}), 500


@app.route('/api/thai_running_tasks', methods=['GET'])
@login_required
def thai_running_tasks():
    """Return currently running/pending Thai scrape tasks for frontend recovery."""
    try:
        rows = db.query_all("""
            SELECT id, task_type, status, result_summary
            FROM scrape_tasks
            WHERE task_type LIKE 'thai_%%_scrape'
              AND status IN ('pending', 'running')
            ORDER BY id DESC
            LIMIT 10
        """)
        tasks = []
        for r in rows:
            tt = r.get('task_type') or ''
            game_type = tt.replace('thai_', '').replace('_scrape', '').upper()
            tasks.append({
                'task_id': r['id'],
                'game_type': game_type,
                'status': r.get('status'),
                'result_summary': r.get('result_summary') or '',
            })
        return jsonify({'status': 'success', 'tasks': tasks})
    except Exception as e:
        logger.error(f"❌ thai_running_tasks 失败: {e}")
        return jsonify({'status': 'error', 'tasks': []})


@app.route('/api/thai_schedule', methods=['POST'])
@login_required
def thai_schedule():
    """触发泰国专题抓取任务（MLBB / SPD / ROV）"""
    try:
        def _norm_list(v):
            if v is None: return []
            if isinstance(v, str): v = [v]
            return [x.strip() for x in v if isinstance(x, str) and x.strip()]

        data = request.get_json(silent=True) or {}
        game_type = (data.get('game_type') or 'MLBB').strip().upper()
        dataset_name = (data.get('dataset_name') or '').strip()
        skip_discover_raw = data.get('skip_discover', False)
        if isinstance(skip_discover_raw, str):
            skip_discover = skip_discover_raw.strip().lower() in ('1', 'true', 'yes', 'y', 'on')
        else:
            skip_discover = bool(skip_discover_raw)
        if game_type not in ('MLBB', 'SPD', 'ROV'):
            return jsonify({'status': 'error', 'message': 'game_type 仅支持 MLBB/SPD/ROV'}), 400
        if not dataset_name:
            return jsonify({'status': 'error', 'message': 'dataset_name 必填'}), 400

        try:
            results_limit = int(data.get('results_limit', 5000))
            max_ai_comments = int(data.get('max_ai_comments', 5000))
            discover_max_posts = int(data.get('discover_max_posts', 3000))
            _min_raw = data.get('min_comments_for_actor', 1)
            min_comments_for_actor = int(_min_raw) if _min_raw is not None and str(_min_raw).strip() != '' else 0
        except (TypeError, ValueError):
            return jsonify({'status': 'error', 'message': '数值参数格式错误'}), 400
        if min_comments_for_actor < 0:
            min_comments_for_actor = 0

        seed_tags = _norm_list(data.get('seed_tags'))
        platforms = _norm_list(data.get('platforms')) or ['facebook', 'instagram']

        if not skip_discover and not seed_tags:
            if game_type == 'MLBB':
                seed_tags = list(MLBB_DISCOVER_KEYWORDS)
            elif game_type == 'SPD':
                seed_tags = list(SPD_KEYWORDS)
            else:
                seed_tags = list(ROV_DISCOVER_KEYWORDS)

        source_dataset_name = None
        dataset_cfg = _thai_datasets_config().get(dataset_name, {})
        dataset_start = dataset_cfg.get('start')
        dataset_end = dataset_cfg.get('end')

        # days_back: cover the full dataset date range so no comments get dropped
        if dataset_start:
            delta = (datetime.date.today() - datetime.date.fromisoformat(dataset_start)).days
            days_back = max(delta + 7, 30)
        else:
            days_back = int(data.get('days_back', 30))

        # SPD 统一走 MLBB 母池切片：不再独立 discover
        if game_type == 'SPD':
            skip_discover = True
            source_dataset_name = '泰国区域'

        # Duplicate prevention: reject if same game_type already has a running/pending task
        existing = db.query_one(
            """SELECT id FROM scrape_tasks
               WHERE task_type = %s AND status IN ('pending','running')
               LIMIT 1""",
            (f'thai_{game_type.lower()}_scrape',)
        )
        if existing:
            return jsonify({
                'status': 'duplicate',
                'task_id': existing['id'],
                'message': f'{game_type} 已有运行中的采集任务（ID: {existing["id"]}），请等待完成后再提交',
            }), 409

        task_id = db.execute_and_fetch_id(
            "INSERT INTO scrape_tasks (task_type, status) VALUES (%s, %s) RETURNING id",
            (f'thai_{game_type.lower()}_scrape', 'pending')
        )

        job_kw = dict(
            scrape_task_id=task_id,
            game_type=game_type,
            dataset_name=dataset_name,
            skip_discover=skip_discover,
            seed_tags=seed_tags,
            platforms=platforms,
            days_back=days_back,
            results_limit=results_limit,
            max_ai_comments=max_ai_comments,
            discover_max_posts=discover_max_posts,
            min_comments_for_actor=min_comments_for_actor,
            source_dataset_name=source_dataset_name,
            dataset_start=dataset_start,
            dataset_end=dataset_end,
        )

        if USE_DB_WORKER:
            queue_task_id = str(uuid.uuid4())
            user_id = session.get('user_id')
            session_id = session.get('session_id', 'default')
            create_task(queue_task_id, user_id, session_id, function_type='thai_scrape')
            task_params = {
                'source': 'thai_schedule',
                'scrape_task_id': task_id,
                'game_type': game_type,
                'dataset_name': dataset_name,
                'skip_discover': skip_discover,
                'seed_tags': seed_tags,
                'platforms': platforms,
                'days_back': days_back,
                'results_limit': results_limit,
                'max_ai_comments': max_ai_comments,
                'discover_max_posts': discover_max_posts,
                'min_comments_for_actor': min_comments_for_actor,
                'source_dataset_name': source_dataset_name,
                'dataset_start': dataset_start,
                'dataset_end': dataset_end,
                'user_id': user_id,
                'session_id': session_id,
            }
            try:
                db.execute(
                    "UPDATE task_queue SET task_params = %s WHERE task_id = %s",
                    (json.dumps(task_params, ensure_ascii=False), queue_task_id),
                )
            except Exception as e:
                logger.error(f"❌ 泰国专题 task_params 写入失败: {e}")
                try:
                    db.execute(
                        "UPDATE scrape_tasks SET status='failed', completed_at=NOW(), error_message=%s WHERE id=%s",
                        (f'队列参数写入失败: {e}'[:500], task_id),
                    )
                except Exception:
                    pass
                return jsonify({'status': 'error', 'message': '任务入队失败，请稍后重试'}), 500
            logger.info(f"✅ 泰国专题任务已入队 (task_queue={queue_task_id}, scrape_tasks={task_id})")
        else:

            def _run_in_web():
                tasks.run_thai_scrape_job(**job_kw, re_raise=False)

            threading.Thread(target=_run_in_web, daemon=True).start()

        return jsonify({
            'status': 'success',
            'message': f'{game_type} 泰国抓取任务已启动' + ('（已入队，由 Worker 执行）' if USE_DB_WORKER else ''),
            'task_id': task_id,
            'dataset_name': dataset_name,
            'seed_tag_count': len(seed_tags),
            'skip_discover': skip_discover,
            'queued': bool(USE_DB_WORKER),
        })
    except Exception as e:
        logger.error(f"❌ 泰国专题任务启动失败: {e}")
        return jsonify({'status': 'error', 'message': str(e)}), 500


import time as _time

_thai_report_cache = {}
_THAI_CACHE_TTL = 300


def _thai_cache_get(key):
    entry = _thai_report_cache.get(key)
    if entry and (_time.time() - entry['ts']) < _THAI_CACHE_TTL:
        return entry['data']
    return None


def _thai_cache_set(key, data):
    _thai_report_cache[key] = {'data': data, 'ts': _time.time()}


from datetime import date as _date, timedelta as _timedelta


def _fetch_daily_engagement_fn(dataset_name, start, end):
            rows = db.query_all_with_timeout("""
                SELECT m.post_date::text AS day,
                       COALESCE(SUM(m.likes),0) + COALESCE(SUM(m.shares),0) + COALESCE(SUM(m.comments_count),0) AS total
                FROM fb_post_metrics m
                JOIN thai_report_datasets d ON d.post_url = m.post_url
                WHERE d.dataset_name = %s
                  AND m.post_date >= %s AND m.post_date <= %s
                GROUP BY m.post_date
                ORDER BY m.post_date
            """, (dataset_name, start, end))
            return {r['day']: int(r['total']) for r in rows}


def _fetch_daily_new_comments_fn(dataset_name, start, end):
            rows = db.query_all_with_timeout("""
                SELECT DATE(c.created_at AT TIME ZONE 'Asia/Shanghai')::text AS day,
                       COUNT(*) AS cnt
                FROM fb_comments c
                JOIN thai_report_datasets d ON d.post_url = c.post_url
                WHERE d.dataset_name = %s
                  AND c.created_at >= %s::date::timestamp AT TIME ZONE 'Asia/Shanghai'
                  AND c.created_at <  (%s::date + INTERVAL '1 day')::timestamp AT TIME ZONE 'Asia/Shanghai'
                GROUP BY 1
                ORDER BY 1
            """, (dataset_name, start, end))
            return {r['day']: int(r['cnt']) for r in rows}


def _peak_post_fn(dataset_name, peak_day):
            row = db.query_one_with_timeout("""
                SELECT m.author, m.post_content,
                       COALESCE(m.likes,0) + COALESCE(m.shares,0) + COALESCE(m.comments_count,0) AS engagement,
                       COUNT(c.id) AS comment_cnt
                FROM fb_comments c
                JOIN fb_post_metrics m ON m.post_url = c.post_url
                JOIN thai_report_datasets d ON d.post_url = c.post_url
                WHERE d.dataset_name = %s
                  AND c.created_at >= %s::date::timestamp AT TIME ZONE 'Asia/Shanghai'
                  AND c.created_at <  (%s::date + INTERVAL '1 day')::timestamp AT TIME ZONE 'Asia/Shanghai'
                GROUP BY m.post_url, m.author, m.post_content, m.likes, m.shares, m.comments_count
                ORDER BY comment_cnt DESC
                LIMIT 1
            """, (dataset_name, peak_day, peak_day))
            if not row:
                return None
            eng = int(row.get('engagement') or 0)
            eng_k = f"{eng/1000:.1f}k" if eng >= 1000 else str(eng)
            return {
                'author': row.get('author') or '',
                'summary': (row.get('post_content') or '')[:60],
                'topic': '',
                'engagement': eng,
                'engagement_label': eng_k,
            }


def _date_range_fn(start, end):
    s = _date.fromisoformat(start)
    e = _date.fromisoformat(end)
    days = []
    cur = s
    while cur <= e:
        days.append(cur.isoformat())
        cur += _timedelta(days=1)
    return days


@app.route('/api/thai_report_data', methods=['GET'])
@login_required
def thai_report_data():
    """Charts + peaks (lightweight, cached 5 min). Top5 moved to /api/thai_report_top5."""
    try:
        cached = _thai_cache_get('report_charts')
        if cached:
            return jsonify(cached)

        # -- keep old local names for minimal diff below --
        _fetch_daily_engagement = _fetch_daily_engagement_fn
        _fetch_daily_new_comments = _fetch_daily_new_comments_fn
        _peak_post = _peak_post_fn
        _date_range = _date_range_fn
        datasets_out = {}
        for ds_name, cfg in _thai_datasets_config().items():
            s, e = cfg['start'], cfg['end']
            days = _date_range_fn(s, e)
            eng_map = _fetch_daily_engagement_fn(ds_name, s, e)
            cmt_map = _fetch_daily_new_comments_fn(ds_name, s, e)
            datasets_out[ds_name] = {
                'game': cfg['game'],
                'start': s,
                'end': e,
                'days': days,
                'engagement': [eng_map.get(d, 0) for d in days],
                'new_comments': [cmt_map.get(d, 0) for d in days],
            }

        chart_c_datasets = ["26Y SPD泰国2", "115泰国热点", "25Y SPD泰国"]
        chart_c_max_days = max(
            (len(datasets_out[n]['days']) for n in chart_c_datasets if n in datasets_out),
            default=0,
        )
        chart_c_labels = [f"D{i+1}" for i in range(chart_c_max_days)]

        peak_chart_groups = [
            ("泰国区域", "26Y SPD泰国1"),
            ("泰国区域", "ROV泰国"),
            ("26Y SPD泰国2", "115泰国热点", "25Y SPD泰国"),
        ]
        peak_data = []
        for group in peak_chart_groups:
            group_peaks = {}
            for ds_name in group:
                if ds_name not in datasets_out:
                    continue
                cmt_list = datasets_out[ds_name]['new_comments']
                days_list = datasets_out[ds_name]['days']
                if not cmt_list or max(cmt_list) == 0:
                    group_peaks[ds_name] = None
                    continue
                peak_idx = cmt_list.index(max(cmt_list))
                peak_day = days_list[peak_idx]
                post = _peak_post_fn(ds_name, peak_day)
                group_peaks[ds_name] = {
                    'peak_day': peak_day,
                    'peak_count': cmt_list[peak_idx],
                    'post': post,
                }
            peak_data.append(group_peaks)

        result = {
            'status': 'success',
            'datasets': datasets_out,
            'chart_c_labels': chart_c_labels,
            'chart_c_datasets': chart_c_datasets,
            'peak_data': peak_data,
        }
        _thai_cache_set('report_charts', result)
        return jsonify(result)

    except Exception as e:
        logger.error(f"\u274c thai_report_data 失败: {e}")
        return jsonify({'status': 'error', 'message': str(e)}), 500


@app.route('/api/thai_report_top5', methods=['GET'])
@login_required
def thai_report_top5():
    """Top5 posts with sentiment analysis (heavier, cached 5 min)."""
    try:
        cached = _thai_cache_get('report_top5')
        if cached:
            return jsonify(cached)

        spd1_cfg = _thai_datasets_config().get("26Y SPD泰国1", {})
        ds_name = "26Y SPD泰国1"
        start = spd1_cfg.get('start', '')
        end = spd1_cfg.get('end', '')

        rows = db.query_all_with_timeout("""
            SELECT m.author, m.post_content, m.post_url, m.platform,
                   COALESCE(m.likes,0) AS likes, COALESCE(m.shares,0) AS shares,
                   COALESCE(m.comments_count,0) AS comments_count,
                   COALESCE(m.likes,0) + COALESCE(m.shares,0) + COALESCE(m.comments_count,0) AS engagement
            FROM fb_post_metrics m
            JOIN thai_report_datasets d ON d.post_url = m.post_url
            WHERE d.dataset_name = %s
              AND m.post_date >= %s AND m.post_date <= %s
            ORDER BY COALESCE(m.comments_count,0) DESC, engagement DESC
            LIMIT 20
        """, (ds_name, start, end))
        preferred = [r for r in rows if int(r.get('comments_count') or 0) > 0]
        chosen_rows = (preferred[:5] if len(preferred) >= 5 else (preferred + rows)[:5])
        top5 = []
        for r in chosen_rows:
            eng = int(r.get('engagement') or 0)
            post_url = r.get('post_url') or ''
            raw_content = r.get('post_content') or ''

            sentiment_ratio = {'positive': 0, 'neutral': 0, 'negative': 0}
            top_sentiments = []
            comments = []
            try:
                comments = db.query_all_with_timeout("""
                    SELECT content, sentiment_score, category, brief_analysis
                    FROM fb_comments
                    WHERE post_url = %s AND sentiment_score IS NOT NULL AND sentiment_score != 0
                    ORDER BY ABS(sentiment_score) DESC
                    LIMIT 200
                """, (post_url,))
                if comments:
                    pos = sum(1 for c in comments if (c.get('sentiment_score') or 0) > 0)
                    neg = sum(1 for c in comments if (c.get('sentiment_score') or 0) < 0)
                    neu = len(comments) - pos - neg
                    total = len(comments)
                    sentiment_ratio = {
                        'positive': round(pos / total * 100),
                        'neutral': round(neu / total * 100),
                        'negative': round(neg / total * 100),
                    }
                    seen_analyses = set()
                    for c in comments:
                        analysis = (c.get('brief_analysis') or '').strip()
                        if analysis and analysis not in seen_analyses:
                            seen_analyses.add(analysis)
                            score = c.get('sentiment_score') or 0
                            label = 'positive' if score > 0 else ('negative' if score < 0 else 'neutral')
                            top_sentiments.append({'text': analysis[:100], 'type': label})
                        if len(top_sentiments) >= 5:
                            break
            except Exception as e:
                logger.warning(f"\u26a0\ufe0f 获取 Top5 情感数据失败: {e}")

            if not top_sentiments and comments:
                try:
                    ai_result = _ai_analyze_single_post(comments, len(comments))
                    if ai_result:
                        if ai_result.get('sentiment_ratio'):
                            sentiment_ratio = ai_result['sentiment_ratio']
                        for s in (ai_result.get('sentiments') or [])[:5]:
                            top_sentiments.append({
                                'text': (s.get('opinion') or s.get('text') or '')[:100],
                                'type': s.get('type', 'neutral'),
                            })
                except Exception:
                    pass

            top5.append({
                'author': r.get('author') or '',
                'platform': (r.get('platform') or 'FACEBOOK').upper(),
                'region': ds_name,
                'post_content': {'topic': '', 'text': raw_content[:120]},
                'url': post_url,
                'likes': int(r.get('likes') or 0),
                'shares': int(r.get('shares') or 0),
                'comments': int(r.get('comments_count') or 0),
                'engagement': eng,
                'engagement_label': f"{eng/1000:.1f}k" if eng >= 1000 else str(eng),
                'sentiment_ratio': sentiment_ratio,
                'top_sentiments': top_sentiments,
            })

        result = {'status': 'success', 'top5': top5}
        _thai_cache_set('report_top5', result)
        return jsonify(result)

    except Exception as e:
        logger.error(f"\u274c thai_report_top5 失败: {e}")
        return jsonify({'status': 'error', 'message': str(e)}), 500


# ============================================
# Excel 数据工具（ETL）
# ============================================

ETL_COMMENTS_MAX_URLS = int(os.environ.get('ETL_COMMENTS_MAX_URLS', '200'))
ETL_PROFILE_VIDEO_MAX_PROFILES = int(os.environ.get('ETL_PROFILE_VIDEO_MAX_PROFILES', '200'))
ETL_PROFILE_VIDEO_MAX_PER_PROFILE = int(os.environ.get('ETL_PROFILE_VIDEO_MAX_PER_PROFILE', '300'))
ETL_VIDEO_METRICS_MAX_UPLOAD_BYTES = int(os.environ.get('ETL_VIDEO_METRICS_MAX_UPLOAD_BYTES', str(20 * 1024 * 1024)))
DEFAULT_VIDEO_METRIC_FIELDS = list(video_metrics_etl.DEFAULT_VIDEO_METRIC_FIELDS)


@app.route('/data-etl')
@login_required
def data_etl_tool():
    return render_template(
        'data_etl.html',
        max_urls=ETL_COMMENTS_MAX_URLS,
    )


@app.route('/api/etl/preview_columns', methods=['POST'])
@login_required
def etl_preview_columns():
    """上传 Excel 首行，返回列名（功能2 下拉）。"""
    try:
        f = request.files.get('file')
        if not f or not f.filename:
            return jsonify({'status': 'error', 'message': '请上传文件'}), 400
        raw = f.read()
        cols = etl_tools.list_text_columns_preview(raw)
        return jsonify({'status': 'success', 'columns': cols})
    except Exception as e:
        logger.error(f"etl_preview_columns: {e}")
        return jsonify({'status': 'error', 'message': str(e)}), 500


@app.route('/api/etl/filter_thai', methods=['POST'])
@login_required
def etl_filter_thai():
    """功能2：仅保留泰语行，同步返回 xlsx。"""
    try:
        f = request.files.get('file')
        if not f or not f.filename:
            return jsonify({'status': 'error', 'message': '请上传文件'}), 400
        col = (request.form.get('text_column') or '').strip() or None
        raw = f.read()
        out_bytes, kept, dropped = etl_tools.filter_thai_rows_excel(raw, col)
        buf = BytesIO(out_bytes)
        buf.seek(0)
        return send_file(
            buf,
            as_attachment=True,
            download_name='thai_filtered.xlsx',
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        )
    except Exception as e:
        logger.error(f"etl_filter_thai: {e}")
        return jsonify({'status': 'error', 'message': str(e)}), 400


@app.route('/api/etl/sum_engagement', methods=['POST'])
@login_required
def etl_sum_engagement():
    """功能4：多文件按日期加总互动量。"""
    try:
        files = request.files.getlist('files')
        if not files:
            return jsonify({'status': 'error', 'message': '请至少上传一个 xlsx'}), 400
        blobs = []
        for f in files:
            if f and f.filename:
                blobs.append(f.read())
        if not blobs:
            return jsonify({'status': 'error', 'message': '无有效文件'}), 400
        out_bytes = etl_tools.sum_daily_engagement_from_excels(blobs)
        buf = BytesIO(out_bytes)
        buf.seek(0)
        return send_file(
            buf,
            as_attachment=True,
            download_name='daily_engagement_sum.xlsx',
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        )
    except Exception as e:
        logger.error(f"etl_sum_engagement: {e}")
        return jsonify({'status': 'error', 'message': str(e)}), 400


@app.route('/api/etl/hashtag_start', methods=['POST'])
@login_required
def etl_hashtag_start():
    """功能1：入队 Hashtag 发现任务。"""
    try:
        data = request.get_json(silent=True) or {}
        seed_tags = [s.strip() for s in (data.get('seed_tags') or []) if s and str(s).strip()]
        if isinstance(data.get('seed_tags'), str):
            seed_tags = [x.strip() for x in data['seed_tags'].replace('\r', '').split('\n') if x.strip()]
        platforms = data.get('platforms') or ['facebook', 'instagram']
        start_date = (data.get('start_date') or '')[:10]
        end_date = (data.get('end_date') or '')[:10]
        max_posts = int(data.get('max_posts') or 500)
        if not seed_tags:
            return jsonify({'status': 'error', 'message': '请填写至少一个 hashtag/关键词'}), 400
        if not start_date or not end_date:
            return jsonify({'status': 'error', 'message': '请填写开始与结束日期'}), 400
    except (TypeError, ValueError) as e:
        return jsonify({'status': 'error', 'message': f'参数错误: {e}'}), 400

    user_id = session.get('user_id')
    session_id = session.get('session_id', 'default')
    queue_task_id = str(uuid.uuid4())
    params = {
        'seed_tags': seed_tags,
        'platforms': platforms,
        'start_date': start_date,
        'end_date': end_date,
        'max_posts': max_posts,
        'user_id': user_id,
        'session_id': session_id,
    }

    create_task(queue_task_id, user_id, session_id, function_type='etl_hashtag')
    if USE_DB_WORKER:
        db.execute(
            "UPDATE task_queue SET task_params = %s WHERE task_id = %s",
            (json.dumps({'source': 'etl_hashtag', **params}, ensure_ascii=False), queue_task_id),
        )
    else:

        def _run():
            etl_jobs.run_etl_hashtag_task(queue_task_id, params, update_task)

        threading.Thread(target=_run, daemon=True).start()

    return jsonify({'status': 'queued', 'task_id': queue_task_id})


@app.route('/api/etl/comments_start', methods=['POST'])
@login_required
def etl_comments_start():
    """功能3：上传含链接的 Excel，入队批量抓评论。"""
    try:
        f = request.files.get('file')
        if not f or not f.filename:
            return jsonify({'status': 'error', 'message': '请上传 Excel'}), 400
        url_col = (request.form.get('url_column') or '').strip() or None
        results_limit = int(request.form.get('results_limit') or 2500)
        days_back = int(request.form.get('days_back') or 365)
        raw = f.read()
        urls = etl_tools.read_urls_from_excel(raw, url_col)
        if len(urls) > ETL_COMMENTS_MAX_URLS:
            return jsonify({
                'status': 'error',
                'message': f'链接数超过上限 {ETL_COMMENTS_MAX_URLS}，请分批上传',
            }), 400
        if not urls:
            return jsonify({'status': 'error', 'message': '未解析到有效 http(s) 链接'}), 400
    except (TypeError, ValueError) as e:
        return jsonify({'status': 'error', 'message': str(e)}), 400

    user_id = session.get('user_id')
    session_id = session.get('session_id', 'default')
    queue_task_id = str(uuid.uuid4())
    params = {
        'post_urls': urls,
        'results_limit': results_limit,
        'days_back': days_back,
        'user_id': user_id,
        'session_id': session_id,
    }

    create_task(queue_task_id, user_id, session_id, function_type='etl_comments')
    if USE_DB_WORKER:
        db.execute(
            "UPDATE task_queue SET task_params = %s WHERE task_id = %s",
            (json.dumps({'source': 'etl_comments', **params}, ensure_ascii=False), queue_task_id),
        )
    else:

        def _run():
            etl_jobs.run_etl_comments_task(queue_task_id, params, update_task)

        threading.Thread(target=_run, daemon=True).start()

    return jsonify({'status': 'queued', 'task_id': queue_task_id, 'url_count': len(urls)})


@app.route('/api/etl/video_metrics_start', methods=['POST'])
@login_required
def etl_video_metrics_start():
    """功能5：批量拉取视频播放量等指标并写回 Excel。"""
    try:
        f = request.files.get('file')
        url_col = (request.form.get('url_column') or '').strip() or None
        manual_urls_text = request.form.get('manual_urls') or ''
        manual_urls = video_metrics_etl.parse_manual_urls(manual_urls_text)

        raw = None
        if f and f.filename:
            raw = f.read()
            if len(raw) > ETL_VIDEO_METRICS_MAX_UPLOAD_BYTES:
                max_mb = round(ETL_VIDEO_METRICS_MAX_UPLOAD_BYTES / 1024 / 1024, 1)
                return jsonify({
                    'status': 'error',
                    'message': f'Excel 文件过大，请控制在 {max_mb}MB 以内或拆分后上传',
                }), 413

        if raw is None and not manual_urls:
            return jsonify({'status': 'error', 'message': '未解析到有效 http(s) 链接'}), 400

        selected_raw = (request.form.get('selected_fields') or '').strip()
        if selected_raw:
            selected_fields = [x.strip() for x in selected_raw.split(',') if x.strip()]
        else:
            selected_fields = list(DEFAULT_VIDEO_METRIC_FIELDS)
    except (TypeError, ValueError) as e:
        return jsonify({'status': 'error', 'message': f'参数错误: {e}'}), 400

    user_id = session.get('user_id')
    session_id = session.get('session_id', 'default')
    queue_task_id = str(uuid.uuid4())

    input_file_id = None
    if raw is not None:
        input_row = db.execute_and_fetch_one(
            """
            INSERT INTO etl_file_outputs (task_id, user_id, filename, content)
            VALUES (%s, %s, %s, %s)
            RETURNING id
            """,
            (queue_task_id, user_id, '_input_video_metrics.xlsx', raw),
        )
        input_file_id = input_row['id'] if input_row else None
        if not input_file_id:
            return jsonify({'status': 'error', 'message': '保存输入文件失败'}), 500

    params = {
        'input_file_id': input_file_id,
        'url_column': url_col,
        'urls': [] if raw is not None else manual_urls,
        'excel_urls': [],
        'manual_urls': manual_urls,
        'extra_manual_urls': None if raw is not None else [],
        'selected_fields': selected_fields,
        'user_id': user_id,
        'session_id': session_id,
    }

    create_task(queue_task_id, user_id, session_id, function_type='etl_video_metrics')
    try:
        set_task_params(queue_task_id, {'source': 'etl_video_metrics', **params})
    except Exception as e:
        logger.exception(f"❌ 写入 task_params 失败: {e}")
        try:
            update_task(queue_task_id, status='failed', error=str(e)[:500], progress='任务参数写入失败')
        except Exception:
            pass
        return jsonify({'status': 'error', 'message': '任务参数写入失败，请稍后重试'}), 500
    if USE_DB_WORKER:
        pass
    else:

        def _run():
            etl_jobs.run_etl_video_metrics_task(queue_task_id, params, update_task)

        threading.Thread(target=_run, daemon=True).start()

    return jsonify({
        'status': 'queued',
        'task_id': queue_task_id,
        'url_count': len(manual_urls) if raw is None else None,
        'message': '任务已排队，后台将解析 Excel 链接' if raw is not None else '任务已排队',
    })


@app.route('/api/etl/profile_video_export_start', methods=['POST'])
@login_required
def etl_profile_video_export_start():
    """按主页链接拉取日期范围内的视频数据，并导出一行一个视频。"""
    try:
        f = request.files.get('file')
        url_col = (request.form.get('url_column') or '').strip() or None
        manual_urls_text = request.form.get('profile_urls') or request.form.get('manual_urls') or ''
        manual_urls = video_metrics_etl.parse_manual_urls(manual_urls_text)

        excel_urls = []
        if f and f.filename:
            parsed = etl_tools.parse_excel_urls(f.read(), url_col)
            excel_urls = parsed.urls

        profile_urls = []
        seen_urls = set()
        for raw_url in [*excel_urls, *manual_urls]:
            key = video_metrics_etl.normalize_url(raw_url)
            if not key or key in seen_urls:
                continue
            if video_metrics_etl.detect_platform(key) == 'UNKNOWN':
                continue
            seen_urls.add(key)
            profile_urls.append(raw_url)

        if not profile_urls:
            return jsonify({'status': 'error', 'message': '未解析到有效主页链接'}), 400
        if len(profile_urls) > ETL_PROFILE_VIDEO_MAX_PROFILES:
            return jsonify({
                'status': 'error',
                'message': f'主页链接数超过上限 {ETL_PROFILE_VIDEO_MAX_PROFILES}，请分批上传',
            }), 400

        start_date = (request.form.get('start_date') or '')[:10]
        end_date = (request.form.get('end_date') or '')[:10]
        if not start_date or not end_date:
            return jsonify({'status': 'error', 'message': '请选择开始与结束日期'}), 400
        if start_date > end_date:
            return jsonify({'status': 'error', 'message': '开始日期不能晚于结束日期'}), 400

        max_videos = int(request.form.get('max_videos') or 100)
        max_videos = max(1, min(max_videos, ETL_PROFILE_VIDEO_MAX_PER_PROFILE))
        hashtag_enabled = _truthy_form_value(request.form.get('hashtag_enabled'))
        hashtag = (request.form.get('hashtag') or '').strip()
        if hashtag_enabled and not video_metrics_etl.parse_hashtag_terms(hashtag):
            return jsonify({'status': 'error', 'message': '启用 hashtag 匹配时请填写 hashtag'}), 400
    except (TypeError, ValueError) as e:
        return jsonify({'status': 'error', 'message': f'参数错误: {e}'}), 400

    user_id = session.get('user_id')
    session_id = session.get('session_id', 'default')
    queue_task_id = str(uuid.uuid4())
    params = {
        'source': 'etl_profile_video_export',
        'mode': 'profile_videos',
        'profile_urls': profile_urls,
        'start_date': start_date,
        'end_date': end_date,
        'max_videos': max_videos,
        'hashtag_enabled': hashtag_enabled,
        'hashtag': hashtag,
        'user_id': user_id,
        'session_id': session_id,
    }

    create_task(queue_task_id, user_id, session_id, function_type='etl_video_metrics')
    try:
        set_task_params(queue_task_id, params)
    except Exception as e:
        logger.error(f"❌ 写入主页视频导出 task_params 失败: {e}")
    if not USE_DB_WORKER:

        def _run():
            etl_jobs.run_etl_video_metrics_task(queue_task_id, params, update_task)

        threading.Thread(target=_run, daemon=True).start()

    return jsonify({
        'status': 'queued',
        'task_id': queue_task_id,
        'profile_count': len(profile_urls),
    })


@app.route('/api/etl/profile-video/configs')
@login_required
def profile_video_configs():
    try:
        rows = profile_video_scheduler.list_configs(session.get('user_id'))
        return jsonify({'status': 'success', 'items': _json_safe_rows(rows)})
    except Exception as e:
        logger.error(f"profile_video_configs failed: {e}")
        return jsonify({'status': 'error', 'message': str(e)}), 500


@app.route('/api/etl/profile-video/feishu_fields')
@login_required
def profile_video_feishu_fields():
    return jsonify({
        'status': 'success',
        'fields': list(profile_video_scheduler.FEISHU_FIELD_MAP.values()),
    })


@app.route('/api/etl/profile-video/configs', methods=['POST'])
@login_required
def profile_video_configs_create():
    try:
        data = request.get_json(silent=True) or {}
        raw_urls = data.get('profile_urls') or data.get('urls') or ''
        if isinstance(raw_urls, str):
            urls = [x.strip() for x in raw_urls.replace('\r', '').split('\n') if x.strip()]
        else:
            urls = [str(x).strip() for x in raw_urls if str(x).strip()]
        if not urls:
            return jsonify({'status': 'error', 'message': '请填写至少一个主页链接'}), 400
        result = profile_video_scheduler.upsert_configs(
            urls,
            user_id=session.get('user_id'),
            enabled=bool(data.get('enabled', True)),
            sync_scope=data.get('sync_scope') or 'recent',
            start_date=data.get('start_date') or None,
            end_date=data.get('end_date') or None,
            max_videos=int(data.get('max_videos') or profile_video_scheduler.default_max_videos_per_profile()),
            schedule_hour=int(data.get('schedule_hour') or profile_video_scheduler.default_sync_hour()),
            feishu_app_token=data.get('feishu_app_token') or None,
            feishu_table_id=data.get('feishu_table_id') or None,
        )
        return jsonify({'status': 'success', **result})
    except Exception as e:
        logger.error(f"profile_video_configs_create failed: {e}")
        return jsonify({'status': 'error', 'message': str(e)}), 400


@app.route('/api/etl/profile-video/configs/<int:config_id>', methods=['PATCH'])
@login_required
def profile_video_configs_update(config_id):
    try:
        data = request.get_json(silent=True) or {}
        row = profile_video_scheduler.update_config(config_id, session.get('user_id'), data)
        if not row:
            return jsonify({'status': 'error', 'message': '配置不存在'}), 404
        return jsonify({'status': 'success', 'item': _json_safe(row)})
    except Exception as e:
        logger.error(f"profile_video_configs_update failed: {e}")
        return jsonify({'status': 'error', 'message': str(e)}), 400


@app.route('/api/etl/profile-video/sync_start', methods=['POST'])
@login_required
def profile_video_sync_start():
    if profile_video_scheduler.profile_video_sync_hard_disabled():
        return jsonify({'status': 'error', 'message': '主页视频同步已被硬禁用，当前不允许创建主页抓取任务'}), 403
    if not profile_video_scheduler.profile_video_sync_enabled():
        return jsonify({'status': 'error', 'message': '主页视频同步未启用，请先配置 PROFILE_VIDEO_SYNC_ENABLED=true'}), 403
    try:
        data = request.get_json(silent=True) or {}
        config_ids = data.get('config_ids') or []
        if isinstance(config_ids, str):
            config_ids = [int(x) for x in config_ids.split(',') if x.strip()]
        config_ids = [int(x) for x in config_ids if str(x).strip()]
        profile_urls = data.get('profile_urls') or []
        if isinstance(profile_urls, str):
            profile_urls = [x.strip() for x in profile_urls.replace('\r', '').split('\n') if x.strip()]
        if not config_ids and not profile_urls:
            enabled = profile_video_scheduler.list_configs(session.get('user_id'))
            config_ids = [r['id'] for r in enabled if r.get('enabled')]
        if not config_ids and not profile_urls:
            return jsonify({'status': 'error', 'message': '没有可同步的主页配置'}), 400
    except Exception as e:
        return jsonify({'status': 'error', 'message': f'参数错误: {e}'}), 400

    user_id = session.get('user_id')
    session_id = session.get('session_id', 'default')
    queue_task_id = str(uuid.uuid4())
    params = {
        'source': 'profile_video_sync',
        'trigger_type': data.get('trigger_type') or 'manual',
        'config_ids': config_ids,
        'profile_urls': profile_urls,
        'user_id': user_id,
        'session_id': session_id,
    }

    create_task(queue_task_id, user_id, session_id, function_type='profile_video_sync')
    try:
        set_task_params(queue_task_id, params)
    except Exception as e:
        logger.error(f"❌ 写入 task_params 失败: {e}")
    if USE_DB_WORKER:
        pass
    else:

        def _run():
            profile_video_scheduler.run_profile_video_sync_task(queue_task_id, params, update_task)

        threading.Thread(target=_run, daemon=True).start()

    return jsonify({'status': 'queued', 'task_id': queue_task_id, 'profile_count': len(config_ids) + len(profile_urls)})


@app.route('/api/etl/feishu-profile-video/config_status')
@login_required
def feishu_profile_video_config_status():
    ok, missing = profile_video_scheduler.validate_feishu_video_table_config()
    config = profile_video_scheduler.feishu_video_table_config()
    return jsonify({
        'status': 'success',
        'configured': ok,
        'missing': missing,
        'base_token': config.get('base_token'),
        'tables': {
            'config': config.get('config_table_id'),
            'latest': config.get('latest_table_id'),
            'snapshot': config.get('snapshot_table_id'),
            'log': config.get('log_table_id'),
        },
    })


@app.route('/api/etl/feishu-profile-video/stop_all', methods=['POST'])
@admin_required
def feishu_profile_video_stop_all():
    """Hard-stop queued/running homepage video automation tasks without calling Apify."""
    reason = ''
    if request.is_json:
        data = request.get_json(silent=True) or {}
        reason = str(data.get('reason') or '').strip()
    if not reason:
        reason = '主页视频同步已手动停止，请重新启用后再执行'
    try:
        result = profile_video_scheduler.stop_pending_profile_video_tasks(reason)
        result.update({
            'status': 'success',
            'message': reason,
            'env': {
                'PROFILE_VIDEO_SYNC_ENABLED': os.environ.get('PROFILE_VIDEO_SYNC_ENABLED', 'false'),
                'FEISHU_PROFILE_VIDEO_SYNC_ENABLED': os.environ.get('FEISHU_PROFILE_VIDEO_SYNC_ENABLED', 'false'),
            },
        })
        return jsonify(result)
    except Exception as e:
        logger.error(f"停止主页视频同步任务失败: {e}")
        return jsonify({'status': 'error', 'message': str(e)[:500]}), 500


@app.route('/api/etl/feishu-profile-video/sync_start', methods=['POST'])
@login_required
def feishu_profile_video_sync_start():
    if profile_video_scheduler.profile_video_sync_hard_disabled():
        return jsonify({'status': 'error', 'message': '飞书主页视频同步已被硬禁用，当前不允许创建主页抓取任务'}), 403
    if not profile_video_scheduler.feishu_profile_video_sync_enabled():
        return jsonify({'status': 'error', 'message': '飞书主页视频同步未启用，请先配置 FEISHU_PROFILE_VIDEO_SYNC_ENABLED=true'}), 403
    ok, missing = profile_video_scheduler.validate_feishu_video_table_config()
    if not ok:
        return jsonify({'status': 'error', 'message': f'缺少环境变量: {", ".join(missing)}'}), 400
    try:
        data = request.get_json(silent=True) or {}
        record_ids = data.get('config_record_ids') or []
        if isinstance(record_ids, str):
            record_ids = [x.strip() for x in record_ids.split(',') if x.strip()]
        profile_urls = data.get('profile_urls') or []
        if isinstance(profile_urls, str):
            profile_urls = [x.strip() for x in profile_urls.replace('\r', '').split('\n') if x.strip()]
    except Exception as e:
        return jsonify({'status': 'error', 'message': f'参数错误: {e}'}), 400

    user_id = session.get('user_id')
    session_id = session.get('session_id', 'default')
    queue_task_id = str(uuid.uuid4())
    params = {
        'source': 'feishu_profile_video_sync',
        'trigger_type': data.get('trigger_type') or 'manual',
        'config_record_ids': record_ids,
        'profile_urls': profile_urls,
        'sync_scope': data.get('sync_scope') or 'recent',
        'recent_days': data.get('recent_days'),
        'max_videos': data.get('max_videos'),
        'project': data.get('project'),
        'user_id': user_id,
        'session_id': session_id,
    }
    create_task(queue_task_id, user_id, session_id, function_type='feishu_profile_video_sync')
    try:
        set_task_params(queue_task_id, params)
    except Exception as e:
        logger.error(f"❌ 写入 task_params 失败: {e}")
    if not USE_DB_WORKER:

        def _run():
            profile_video_scheduler.run_feishu_profile_video_sync_task(queue_task_id, params, update_task)

        threading.Thread(target=_run, daemon=True).start()
    return jsonify({
        'status': 'queued',
        'task_id': queue_task_id,
        'profile_count': len(record_ids) + len(profile_urls) if (record_ids or profile_urls) else None,
    })


@app.route('/api/etl/task_status/<task_id>')
@login_required
def api_etl_task_status(task_id):
    """ETL 异步任务状态（result 为 JSON 字符串，不做 HTML 清洗）。"""
    row = db.query_one(
        """
        SELECT status, progress, result, error, user_id
        FROM task_queue WHERE task_id = %s
        """,
        (task_id,),
    )
    if not row:
        return jsonify({'error': '任务不存在'}), 404
    if row.get('user_id') is not None and row.get('user_id') != session.get('user_id'):
        return jsonify({'error': '无权访问'}), 403
    return jsonify({
        'status': row.get('status'),
        'progress': row.get('progress'),
        'result': row.get('result') or '',
        'error': row.get('error'),
    })


@app.route('/api/etl/video_metrics_tasks')
@login_required
def api_etl_video_metrics_tasks():
    """最近的拉视频数据任务，供轮询超时或刷新后找回下载文件。"""
    try:
        limit = max(1, min(int(request.args.get('limit') or 20), 100))
    except (TypeError, ValueError):
        limit = 20

    uid = session.get('user_id')
    rows = db.query_all(
        """
        WITH recent_task_ids AS (
            SELECT task_id
            FROM task_queue
            WHERE user_id = %s AND function_type = 'etl_video_metrics'
            UNION
            SELECT task_id
            FROM etl_file_outputs
            WHERE user_id = %s AND filename = 'video_metrics.xlsx'
        )
        SELECT q.task_id,
               q.status,
               q.progress,
               q.result,
               q.error,
               q.created_at,
               q.updated_at,
               q.finished_at,
               o.download_id AS fallback_download_id,
               o.filename AS fallback_filename,
               o.created_at AS fallback_created_at
        FROM recent_task_ids ids
        LEFT JOIN task_queue q
               ON q.task_id = ids.task_id
              AND (q.user_id = %s OR q.user_id IS NULL)
        LEFT JOIN LATERAL (
            SELECT id AS download_id, filename, created_at
            FROM etl_file_outputs
            WHERE user_id = %s
              AND task_id = ids.task_id
              AND filename <> '_input_video_metrics.xlsx'
            ORDER BY id DESC
            LIMIT 1
        ) o ON TRUE
        ORDER BY COALESCE(q.created_at, o.created_at) DESC
        LIMIT %s
        """,
        (uid, uid, uid, uid, limit),
    ) or []

    items = []
    for row in rows:
        result_payload = {}
        raw_result = row.get('result')
        if raw_result:
            try:
                result_payload = json.loads(raw_result) if isinstance(raw_result, str) else dict(raw_result)
            except Exception:
                result_payload = {}

        download_id = result_payload.get('download_id') or row.get('fallback_download_id')
        filename = result_payload.get('filename') or row.get('fallback_filename') or 'video_metrics.xlsx'
        items.append({
            'task_id': row.get('task_id'),
            'status': row.get('status') or ('completed' if download_id else 'unknown'),
            'progress': row.get('progress') or ('已生成文件' if download_id else None),
            'error': row.get('error'),
            'created_at': row.get('created_at') or row.get('fallback_created_at'),
            'updated_at': row.get('updated_at'),
            'finished_at': row.get('finished_at'),
            'download_id': download_id,
            'filename': filename,
            'mode': result_payload.get('mode'),
            'url_count': result_payload.get('url_count'),
            'profile_count': result_payload.get('profile_count'),
            'video_count': result_payload.get('video_count'),
            'success_count': result_payload.get('success_count'),
        })

    return jsonify({'status': 'success', 'items': _json_safe_rows(items)})


@app.route('/api/etl/download/<int:output_id>')
@login_required
def etl_download(output_id):
    uid = session.get('user_id')
    row = db.query_one(
        "SELECT id, filename, content, user_id FROM etl_file_outputs WHERE id = %s",
        (output_id,),
    )
    if not row:
        return jsonify({'error': '文件不存在'}), 404
    if row.get('user_id') is not None and row.get('user_id') != uid:
        return jsonify({'error': '无权下载'}), 403
    data = row.get('content')
    if data is None:
        return jsonify({'error': '空文件'}), 404
    if isinstance(data, memoryview):
        data = data.tobytes()
    buf = BytesIO(data)
    buf.seek(0)
    return send_file(
        buf,
        as_attachment=True,
        download_name=row.get('filename') or 'export.xlsx',
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    )


# ============================================
# TikTok 官号监控（官方 API）
# ============================================

@app.route('/tiktok-official')
@login_required
def tiktok_official_page():
    return render_template('tiktok_official.html')


@app.route('/api/tiktok-official/accounts')
@login_required
def api_tiktok_official_accounts():
    try:
        accounts = tiktok_official_service.sync_configured_accounts()
        return jsonify({'status': 'success', 'accounts': _json_safe_rows(accounts)})
    except Exception as e:
        logger.error(f"tiktok_official accounts failed: {e}")
        return jsonify({'status': 'error', 'message': str(e)}), 500


@app.route('/api/tiktok-official/auth-urls')
@login_required
def api_tiktok_official_auth_urls():
    app_id = (
        os.environ.get('TIKTOK_APP_ID')
        or os.environ.get('TIKTOK_CLIENT_KEY')
        or ''
    ).strip()
    if not app_id:
        return jsonify({'status': 'error', 'message': 'TIKTOK_APP_ID 未配置'}), 400
    public_base = _tiktok_public_base_url()
    account_redirect = f'{public_base}/tiktok/account/callback'
    business_redirect = f'{public_base}/tiktok/business/callback'
    scopes = [
        'user.info.basic',
        'user.info.username',
        'user.info.stats',
        'user.info.profile',
        'user.account.type',
        'user.insights',
        'video.list',
        'video.insights',
    ]
    account_params = {
        'client_key': app_id,
        'scope': ','.join(scopes),
        'response_type': 'code',
        'redirect_uri': account_redirect,
        'state': f'tiktok_account_{uuid.uuid4().hex[:12]}',
    }
    business_params = {
        'app_id': app_id,
        'state': f'tiktok_business_{uuid.uuid4().hex[:12]}',
        'redirect_uri': business_redirect,
    }
    return jsonify({
        'status': 'success',
        'account_url': 'https://www.tiktok.com/v2/auth/authorize?' + urlencode(account_params),
        'business_url': 'https://business-api.tiktok.com/portal/auth?' + urlencode(business_params),
        'account_redirect_uri': account_redirect,
        'business_redirect_uri': business_redirect,
    })


@app.route('/api/tiktok-official/refresh', methods=['POST'])
@login_required
def api_tiktok_official_refresh():
    try:
        data = request.get_json(silent=True) or {}
        business_ids = data.get('business_ids') or []
        if isinstance(business_ids, str):
            business_ids = [business_ids]
        profile_days = int(data.get('profile_days') or 30)
        max_pages = int(data.get('max_pages') or 5)
        profile_days = max(1, min(profile_days, 60))
        max_pages = max(1, min(max_pages, 50))
    except (TypeError, ValueError) as e:
        return jsonify({'status': 'error', 'message': f'参数错误: {e}'}), 400

    user_id = session.get('user_id')
    session_id = session.get('session_id', 'default')
    queue_task_id = str(uuid.uuid4())
    params = {
        'business_ids': business_ids,
        'profile_days': profile_days,
        'max_pages': max_pages,
        'user_id': user_id,
        'session_id': session_id,
    }

    create_task(queue_task_id, user_id, session_id, function_type='tiktok_official_refresh')
    if USE_DB_WORKER:
        db.execute(
            "UPDATE task_queue SET task_params = %s WHERE task_id = %s",
            (json.dumps({'source': 'tiktok_official_refresh', **params}, ensure_ascii=False), queue_task_id),
        )
    else:

        def _run():
            tiktok_official_service.run_refresh_task(queue_task_id, params, update_task)

        threading.Thread(target=_run, daemon=True).start()

    return jsonify({'status': 'queued', 'task_id': queue_task_id})


@app.route('/api/tiktok-official/refresh-status/<task_id>')
@login_required
def api_tiktok_official_refresh_status(task_id):
    return api_etl_task_status(task_id)


@app.route('/api/tiktok-official/videos')
@login_required
def api_tiktok_official_videos():
    try:
        business_id = (request.args.get('business_id') or '').strip() or None
        page = int(request.args.get('page') or 1)
        page_size = int(request.args.get('page_size') or 50)
        page = max(1, page)
        page_size = max(1, min(page_size, 200))
        result = tiktok_official_service.list_videos(business_id=business_id, page=page, page_size=page_size)
        return jsonify({'status': 'success', **_json_safe(result)})
    except Exception as e:
        logger.error(f"tiktok_official videos failed: {e}")
        return jsonify({'status': 'error', 'message': str(e)}), 500


@app.route('/api/tiktok-official/videos/<item_id>')
@login_required
def api_tiktok_official_video_detail(item_id):
    try:
        business_id = (request.args.get('business_id') or '').strip() or None
        row = tiktok_official_service.get_video(item_id, business_id)
        if not row:
            return jsonify({'status': 'error', 'message': '视频不存在'}), 404
        return jsonify({'status': 'success', 'video': _json_safe(row)})
    except Exception as e:
        logger.error(f"tiktok_official video detail failed: {e}")
        return jsonify({'status': 'error', 'message': str(e)}), 500


@app.route('/api/tiktok-official/profile-metrics')
@login_required
def api_tiktok_official_profile_metrics():
    try:
        business_id = (request.args.get('business_id') or '').strip() or None
        days = int(request.args.get('days') or 30)
        rows = tiktok_official_service.list_profile_metrics(business_id=business_id, days=days)
        return jsonify({'status': 'success', 'items': _json_safe_rows(rows)})
    except Exception as e:
        logger.error(f"tiktok_official profile metrics failed: {e}")
        return jsonify({'status': 'error', 'message': str(e)}), 500


@app.route('/api/tiktok-official/export', methods=['POST'])
@login_required
def api_tiktok_official_export():
    try:
        data = request.get_json(silent=True) or {}
        videos = data.get('videos') or []
        if not isinstance(videos, list):
            return jsonify({'status': 'error', 'message': 'videos 必须是数组'}), 400
        out = tiktok_official_service.build_export(videos=videos)
        buf = BytesIO(out)
        buf.seek(0)
        return send_file(
            buf,
            as_attachment=True,
            download_name='tiktok_official_monitor.xlsx',
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        )
    except Exception as e:
        logger.error(f"tiktok_official export failed: {e}")
        return jsonify({'status': 'error', 'message': str(e)}), 500


def _json_safe_rows(rows):
    return [_json_safe(dict(row)) for row in (rows or [])]


def _json_safe(value):
    if isinstance(value, dict):
        return {k: _json_safe(v) for k, v in value.items()}
    if isinstance(value, list):
        return [_json_safe(v) for v in value]
    if isinstance(value, (datetime.datetime, datetime.date)):
        return value.isoformat()
    return value


# ============================================
# 应用启动
# ============================================

if __name__ == '__main__':
    # 恢复被中断的任务
    logger.info("🔄 检查并恢复被中断的任务...")
    recover_interrupted_tasks()

    logger.info("\n" + "=" * 60)
    logger.info("🎉 Sailson AI 工作台已启动")
    logger.info(f"🌐 访问地址: http://0.0.0.0:{PORT}")
    logger.info("=" * 60 + "\n")

    app.run(debug=False, host='0.0.0.0', port=PORT)
